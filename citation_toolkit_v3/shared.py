"""
Citation Toolkit — shared functions and CSS.
Imported by every page module.
"""
"""
Citation Toolkit v3 — Complete
All 8 apps in one file.
Run: streamlit run app_v2_new.py
"""
import base64
import html as html_module
import io
import re
import sqlite3
import tempfile
import os
import xml.etree.ElementTree as ET
import zipfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path

import requests
import streamlit as st
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

st.set_page_config(page_title="Citation Toolkit", page_icon="📚",
                   layout="wide", initial_sidebar_state="expanded")

APP_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Libre+Baskerville:ital,wght@0,400;0,700;1,400&family=Source+Sans+3:wght@300;400;500;600&family=Source+Serif+4:ital,opsz,wght@0,8..60,300;0,8..60,400;1,8..60,300;1,8..60,400&display=swap');

:root {
  --bg:         #f5f0e8;
  --bg2:        #ede8dc;
  --surface:    #faf7f2;
  --border:     #c8b89a;
  --border-soft:#d8cbb5;
  --ink:        #1a1208;
  --ink-mid:    #2e2416;
  --ink-dim:    #5a4a38;
  --ink-faint:  #8a7a65;
  --accent:     #8b1a1a;
  --accent2:    #c0392b;
  --green:      #1e5c3a;
  --amber:      #8b4a10;
  --radius:     4px;
  --shadow:     0 1px 4px rgba(0,0,0,0.12);
}

html, body, [class*="css"] {
  font-family: 'Source Sans 3', Georgia, sans-serif;
  background: var(--bg) !important;
  color: var(--ink);
}
h1,h2,h3,h4 {
  font-family: 'Libre Baskerville', Georgia, serif;
  color: var(--ink);
  font-weight: 400;
}
h1, h2, h3, h4,
div[data-testid="stMarkdownContainer"] h1,
div[data-testid="stMarkdownContainer"] h2,
div[data-testid="stMarkdownContainer"] h3 {
  font-family: 'Libre Baskerville', Georgia, serif !important;
  font-weight: 400 !important;
  color: var(--ink) !important;
  letter-spacing: -0.02em;
}
h2, div[data-testid="stMarkdownContainer"] h2 {
  font-size: 2.1rem !important;
  line-height: 1.2 !important;
  margin-bottom: 0.3rem !important;
  border: none !important;
}
h3 { font-size: 1.25rem !important; color: var(--ink-mid) !important; }
p, div[data-testid="stMarkdownContainer"] p {
  color: var(--ink-mid) !important;
  line-height: 1.75 !important;
  font-size: 1.02rem !important;
  font-family: 'Source Sans 3', sans-serif !important;
}
/* Kill monospace on ALL text elements */
div[data-testid="stMarkdownContainer"] * {
  font-family: 'Source Sans 3', Georgia, sans-serif;
}
div[data-testid="stMarkdownContainer"] code,
div[data-testid="stMarkdownContainer"] pre {
  font-family: 'Source Code Pro', monospace !important;
}

/* Main layout */
.main { background: var(--bg) !important; }
.main .block-container {
  padding: 2.8rem 3.5rem 5rem;
  max-width: 980px;
  background: var(--bg);
}

/* Sidebar */
section[data-testid="stSidebar"] {
  background: var(--surface) !important;
  border-right: 1px solid var(--border) !important;
}
section[data-testid="stSidebar"] * { color: var(--ink-mid) !important; }
/* Sidebar nav buttons — visible, styled as nav items */
section[data-testid="stSidebar"] .stButton button {
  display: flex !important;
  width: 100% !important;
  background: transparent !important;
  border: none !important;
  border-left: 2px solid transparent !important;
  border-radius: 3px !important;
  padding: 7px 10px !important;
  margin-bottom: 1px !important;
  text-align: left !important;
  justify-content: flex-start !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.87rem !important;
  font-weight: 400 !important;
  color: var(--ink) !important;
  cursor: pointer !important;
  transition: background 0.12s, border-color 0.12s !important;
  box-shadow: none !important;
  letter-spacing: 0 !important;
  text-transform: none !important;
}
section[data-testid="stSidebar"] .stButton button:hover {
  background: var(--bg2) !important;
  border-left-color: var(--border) !important;
  color: var(--ink) !important;
}
section[data-testid="stSidebar"] .stButton {
  margin-bottom: 0 !important;
  padding: 0 2px !important;
}
/* Section rule headers */
section[data-testid="stSidebar"] .sidebar-rule {
  display: flex;
  align-items: center;
  gap: 8px;
  font-size: 0.58rem !important;
  letter-spacing: 0.18em !important;
  text-transform: uppercase !important;
  color: var(--ink-faint) !important;
  padding: 12px 10px 5px !important;
}
section[data-testid="stSidebar"] .sidebar-rule::after {
  content: '';
  flex: 1;
  height: 1px;
  background: var(--border-soft);
}
/* Ensure sidebar labels are never washed out */
section[data-testid="stSidebar"] label { color: var(--ink) !important; }

/* Sidebar logo area */
.sidebar-logo {
  font-family: 'Libre Baskerville', Georgia, serif;
  font-size: 1.25rem;
  color: var(--ink) !important;
  font-weight: 400;
  line-height: 1.2;
}
.sidebar-logo em { color: var(--accent) !important; font-style: italic; }
.sidebar-sub {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.72rem;
  letter-spacing: 0.14em;
  text-transform: uppercase;
  color: var(--ink-faint) !important;
  margin-top: 4px;
  margin-bottom: 1.2rem;
}
.sidebar-rule {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.65rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint) !important;
  display: flex;
  align-items: center;
  gap: 8px;
  margin: 1.2rem 0 0.5rem;
}
.sidebar-rule::after {
  content: '';
  flex: 1;
  height: 1px;
  background: var(--border);
}

/* App pill badge */
.app-label {
  display: inline-block;
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.68rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint);
  border-top: 2px solid var(--border);
  border-bottom: 2px solid var(--border);
  padding: 4px 0;
  margin-bottom: 0.6rem;
  display: flex;
  align-items: center;
  gap: 10px;
}
.app-label::before, .app-label::after {
  content: '—';
  color: var(--ink-faint);
  font-size: 0.8rem;
}

/* Instruction box */
.instruction-box {
  background: var(--surface);
  border: 1px solid var(--border-soft);
  border-left: 3px solid var(--accent);
  border-radius: var(--radius);
  padding: 1rem 1.3rem;
  margin: 0.8rem 0 1.6rem;
  font-size: 0.9rem;
  color: var(--ink-mid);
  line-height: 1.8;
}
.instruction-box b  { color: var(--ink); font-weight: 600; }
.instruction-box code {
  background: var(--bg2);
  padding: 1px 5px;
  border-radius: 3px;
  font-family: 'Source Code Pro', monospace;
  font-size: 0.82rem;
  color: var(--accent);
  border: 1px solid var(--border-soft);
}
.instruction-box ul { margin: 0.3rem 0 0; padding-left: 1.2rem; }
.instruction-box li { margin-bottom: 0.25rem; }

/* Cards */
.card {
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  padding: 1.2rem 1.5rem;
  margin-bottom: 0.9rem;
  box-shadow: var(--shadow);
}
.card-accent { border-left: 3px solid var(--accent); }
.card-green  { border-left: 3px solid var(--green); background: #f2f7f4; }
.card-red    { border-left: 3px solid #c0392b; background: #fdf5f5; }
.card-amber  { border-left: 3px solid var(--amber); background: #fdf8f2; }

/* Reference items */
.ref-item {
  padding: 0.5rem 0.85rem;
  border-radius: var(--radius);
  margin-bottom: 5px;
  font-size: 0.85rem;
  border-left: 3px solid;
  line-height: 1.5;
  color: var(--ink-mid);
}
.ref-item.error   { border-color: #c0392b; background: #fdf5f5; }
.ref-item.warning { border-color: var(--amber); background: #fdf8f2; }
.ref-item.ok      { border-color: var(--green); background: #f2f7f4; }

/* Section divider */
.section-rule {
  display: flex;
  align-items: center;
  gap: 10px;
  margin: 2rem 0 1.2rem;
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.68rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint);
}
.section-rule::before, .section-rule::after {
  content: ''; flex: 1; height: 1px; background: var(--border);
}

/* Match cards */
.match-card {
  background: var(--surface);
  border: 1px solid var(--border);
  border-left: 3px solid var(--accent);
  border-radius: var(--radius);
  padding: 1rem 1.2rem;
  margin-bottom: 0.7rem;
  box-shadow: var(--shadow);
}
.match-card.accepted { border-left-color: var(--green); background: #f2f7f4; }
.match-card.skipped  { border-left-color: #c0392b; background: #fdf5f5; }
.match-sentence { font-size: 0.87rem; color: var(--ink-dim); font-style: italic; margin-bottom: 0.5rem; line-height: 1.7; }
.match-marker { font-family: monospace; font-size: 0.75rem; background: var(--bg2); color: var(--accent); padding: 2px 7px; border-radius: 3px; border: 1px solid var(--border-soft); display: inline-block; margin-bottom: 6px; }
.score-pill { font-family: monospace; font-size: 0.7rem; padding: 2px 8px; border-radius: 3px; border: 1px solid; }
.score-high { background: #f2f7f4; color: var(--green); border-color: #a8d4be; }
.score-mid  { background: #fdf8f2; color: var(--amber); border-color: #e0c8a8; }
.score-low  { background: #fdf5f5; color: #c0392b; border-color: #e0b8b8; }

/* ── Step cards (App 2 stages) ───────────────────────────────────────── */
.step-card {
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  padding: 0.9rem 1.2rem;
  margin-bottom: 0.8rem;
  box-shadow: var(--shadow);
  border-left: 4px solid var(--border);
}
.step-card.active { border-left-color: var(--accent); }
.step-card.done   { border-left-color: var(--green); background: #f4fbf7; }
.step-card.waiting{ border-left-color: var(--border); opacity: 0.7; }
.step-header {
  display: flex;
  align-items: baseline;
  gap: 12px;
  margin-bottom: 4px;
}
.step-num {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.65rem;
  font-weight: 700;
  letter-spacing: 0.12em;
  text-transform: uppercase;
  color: var(--accent);
  white-space: nowrap;
  padding: 2px 7px;
  border: 1px solid var(--accent);
  border-radius: 3px;
  line-height: 1.6;
}
.step-num.done {
  color: var(--green);
  border-color: var(--green);
  background: rgba(45,106,79,0.08);
}
.step-title {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.95rem;
  font-weight: 600;
  color: var(--ink);
}
.step-desc {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.83rem;
  color: var(--ink-dim);
  margin-top: 2px;
  line-height: 1.5;
}

/* ── Progress bar (custom) ───────────────────────────────────────── */
.progress-outer {
  background: var(--bg2);
  border-radius: 4px;
  height: 8px;
  width: 100%;
  overflow: hidden;
  margin: 6px 0;
}
.progress-inner {
  background: var(--accent);
  height: 100%;
  border-radius: 4px;
  transition: width 0.3s ease;
}

/* ── Section label (inline divider text) ────────────────────────── */
.section-label {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.65rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint);
  display: flex;
  align-items: center;
  gap: 10px;
  margin: 1.2rem 0 0.8rem;
}
.section-label::before, .section-label::after {
  content: ''; flex: 1; height: 1px; background: var(--border-soft);
}

/* ── Streamlit global overrides ─────────────────────────────────────── */
.stApp, .main, section[data-testid="stMain"],
div[data-testid="stAppViewContainer"] {
  background: var(--bg) !important;
}

/* Metrics */
div[data-testid="stMetric"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  padding: 0.8rem 1rem !important;
  box-shadow: var(--shadow) !important;
}
div[data-testid="stMetric"] label {
  color: var(--ink-dim) !important;
  font-size: 0.75rem !important;
  font-family: 'Source Sans 3', sans-serif !important;
  letter-spacing: 0.08em !important;
  text-transform: uppercase !important;
}
div[data-testid="stMetric"] div[data-testid="stMetricValue"] {
  font-family: 'Libre Baskerville', serif !important;
  font-size: 1.9rem !important;
  font-weight: 400 !important;
  color: var(--ink) !important;
}

/* Buttons — warm, visible, never black-filled */
button[kind="primary"], .stButton > button[kind="primary"] {
  background: #8b1a1a !important;
  color: #faf7f2 !important;
  border: none !important;
  border-radius: var(--radius) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.84rem !important;
  font-weight: 600 !important;
  letter-spacing: 0.08em !important;
  text-transform: uppercase !important;
  padding: 0.55rem 1.5rem !important;
  box-shadow: 0 1px 4px rgba(139,26,26,0.25) !important;
  transition: background 0.15s, box-shadow 0.15s !important;
}
button[kind="primary"]:hover {
  background: #6b1212 !important;
  box-shadow: 0 2px 8px rgba(139,26,26,0.4) !important;
}
button[kind="secondary"], .stButton > button[kind="secondary"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.84rem !important;
}
button[kind="secondary"]:hover {
  border-color: var(--accent) !important;
  color: var(--accent) !important;
}

/* File uploader — cream, warm border, never dark */
div[data-testid="stFileUploader"] {
  background: var(--surface) !important;
  border: 1.5px dashed var(--border) !important;
  border-radius: var(--radius) !important;
}
div[data-testid="stFileUploader"] * {
  color: var(--ink-mid) !important;
  background: transparent !important;
}
div[data-testid="stFileUploader"]:hover {
  border-color: var(--accent) !important;
}
/* The inner Browse-files button */
div[data-testid="stFileUploader"] button {
  background: var(--bg2) !important;
  border: 1px solid var(--border) !important;
  color: var(--ink-mid) !important;
  border-radius: var(--radius) !important;
}
div[data-testid="stFileUploader"] button:hover {
  border-color: var(--accent) !important;
  color: var(--accent) !important;
}

/* Expanders */
div[data-testid="stExpander"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  box-shadow: var(--shadow) !important;
  margin-bottom: 0.6rem !important;
}
div[data-testid="stExpander"] {
  overflow: visible !important;
}
/* Expander — comprehensive fix for arrow text leak ("_arro", "_Step" etc) */
div[data-testid="stExpander"] {
  overflow: hidden !important;
  border-radius: var(--radius) !important;
}
div[data-testid="stExpander"] details > summary {
  display: flex !important;
  align-items: center !important;
  gap: 8px !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.92rem !important;
  font-weight: 500 !important;
  color: var(--ink) !important;
  background: var(--surface) !important;
  padding: 0.75rem 1rem !important;
  line-height: 1.5 !important;
  cursor: pointer !important;
  list-style: none !important;
  user-select: none !important;
  overflow: hidden !important;
}
div[data-testid="stExpander"] details > summary::-webkit-details-marker {
  display: none !important;
}
div[data-testid="stExpander"] details > summary::marker {
  display: none !important;
  content: "" !important;
}
div[data-testid="stExpander"] details > summary:hover {
  background: var(--bg2) !important;
}
/* Hide the internal collapse button that leaks "_arro" text */
div[data-testid="stExpander"] details > summary > div:first-child {
  display: flex !important;
  align-items: center !important;
  overflow: hidden !important;
  min-width: 0 !important;
}
/* The arrow/toggle button element — clip any overflowing text */
div[data-testid="stExpander"] button[kind="header"],
div[data-testid="stExpander"] [data-testid="stExpanderToggleIcon"] {
  overflow: hidden !important;
  font-size: 0 !important;
  width: 20px !important;
  flex-shrink: 0 !important;
}
/* Label text inside summary */
div[data-testid="stExpander"] details > summary p,
div[data-testid="stExpander"] details > summary span,
div[data-testid="stExpander"] details > summary div {
  color: var(--ink) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.92rem !important;
  font-weight: 500 !important;
  margin: 0 !important;
  overflow: hidden !important;
  text-overflow: ellipsis !important;
}
/* Arrow SVG — properly contained */
div[data-testid="stExpander"] details > summary svg {
  color: var(--ink-dim) !important;
  fill: var(--ink-dim) !important;
  min-width: 16px !important;
  width: 16px !important;
  height: 16px !important;
  flex-shrink: 0 !important;
  overflow: visible !important;
}
/* Body text inside expander */
div[data-testid="stExpander"] div[data-testid="stMarkdownContainer"] p,
div[data-testid="stExpander"] div[data-testid="stMarkdownContainer"] * {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}
/* Expander body text — dark and readable */
div[data-testid="stExpander"] div[data-testid="stMarkdownContainer"] p,
div[data-testid="stExpander"] div[data-testid="stMarkdownContainer"] * {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}

/* Text inputs, textareas, number inputs */
div[data-testid="stTextInput"] input,
div[data-testid="stTextArea"] textarea,
div[data-testid="stNumberInput"] input {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  color: var(--ink) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.9rem !important;
}
div[data-testid="stTextInput"] input:focus,
div[data-testid="stTextArea"] textarea:focus,
div[data-testid="stNumberInput"] input:focus {
  border-color: var(--accent) !important;
  box-shadow: 0 0 0 2px rgba(139,26,26,0.12) !important;
}
div[data-testid="stTextInput"] label,
div[data-testid="stTextArea"] label,
div[data-testid="stNumberInput"] label {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.85rem !important;
}

/* Selectbox / dropdown */
div[data-testid="stSelectbox"] > div > div,
div[data-testid="stSelectbox"] > div > div > div,
div[data-testid="stSelectbox"] span,
div[data-testid="stSelectbox"] p {
  background: var(--surface) !important;
  border-color: var(--border) !important;
  border-radius: var(--radius) !important;
  color: var(--ink) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.9rem !important;
}
div[data-testid="stSelectbox"] label {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.85rem !important;
}
/* Dropdown arrow container */
div[data-testid="stSelectbox"] svg {
  color: var(--ink-dim) !important;
  fill: var(--ink-dim) !important;
}
/* Dropdown option list */
div[role="listbox"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}
div[role="option"] {
  font-family: 'Source Sans 3', sans-serif !important;
  color: var(--ink) !important;
  font-size: 0.9rem !important;
  background: var(--surface) !important;
}
div[role="option"]:hover {
  background: var(--bg2) !important;
}

/* Radio buttons */
div[data-testid="stRadio"] label {
  color: var(--ink) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.9rem !important;
  font-weight: 400 !important;
}
div[data-testid="stRadio"] label:hover { color: var(--accent) !important; }
/* All markdown text — force dark */
div[data-testid="stMarkdownContainer"],
div[data-testid="stMarkdownContainer"] p,
div[data-testid="stMarkdownContainer"] li,
div[data-testid="stMarkdownContainer"] span {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}
div[data-testid="stMarkdownContainer"] strong,
div[data-testid="stMarkdownContainer"] b {
  color: var(--ink) !important;
  font-weight: 600 !important;
}
/* Caption */
div[data-testid="stCaptionContainer"],
.stCaption, small {
  color: var(--ink-dim) !important;
  font-size: 0.79rem !important;
}

/* Checkbox */
div[data-testid="stCheckbox"] label {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.88rem !important;
}

/* Toggle */
div[data-testid="stToggle"] label {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}

/* DataFrames / tables */
.stDataFrame {
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  background: var(--surface) !important;
}

/* Alerts / info boxes */
div[data-testid="stAlert"] {
  border-radius: var(--radius) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.88rem !important;
}

/* Tabs */
div[data-testid="stTabs"] button {
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.87rem !important;
  color: var(--ink-dim) !important;
  background: transparent !important;
}
div[data-testid="stTabs"] button[aria-selected="true"] {
  color: var(--accent) !important;
  border-bottom-color: var(--accent) !important;
}

/* Progress bar */
div[data-testid="stProgress"] > div { background: var(--bg2) !important; }
div[data-testid="stProgress"] > div > div { background: var(--accent) !important; }

/* Caption / small text */
.stCaption, small { color: var(--ink-dim) !important; font-size: 0.79rem !important; }

/* Spinner */
div[data-testid="stSpinner"] > div { border-top-color: var(--accent) !important; }

/* Download button inherits primary */
a[data-testid="stDownloadButton"] button,
div[data-testid="stDownloadButton"] button {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  color: var(--ink-mid) !important;
  border-radius: var(--radius) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.84rem !important;
}
a[data-testid="stDownloadButton"] button:hover,
div[data-testid="stDownloadButton"] button:hover {
  border-color: #3d5a6e !important;
  color: #3d5a6e !important;
}

/* Dividers */
hr { border-color: var(--border) !important; }

/* Columns gap */
div[data-testid="column"] { background: transparent !important; }

/* Toast / notifications */
div[data-testid="stToast"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  color: var(--ink) !important;
  border-radius: var(--radius) !important;
}
</style>
"""

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Libre+Baskerville:ital,wght@0,400;0,700;1,400&family=Source+Sans+3:wght@300;400;500;600&family=Source+Serif+4:ital,opsz,wght@0,8..60,300;0,8..60,400;1,8..60,300;1,8..60,400&display=swap');

:root {
  --bg:         #f5f0e8;
  --bg2:        #ede8dc;
  --surface:    #faf7f2;
  --border:     #c8b89a;
  --border-soft:#d8cbb5;
  --ink:        #1a1208;
  --ink-mid:    #2e2416;
  --ink-dim:    #5a4a38;
  --ink-faint:  #8a7a65;
  --accent:     #8b1a1a;
  --accent2:    #c0392b;
  --green:      #1e5c3a;
  --amber:      #8b4a10;
  --radius:     4px;
  --shadow:     0 1px 4px rgba(0,0,0,0.12);
}

html, body, [class*="css"] {
  font-family: 'Source Sans 3', Georgia, sans-serif;
  background: var(--bg) !important;
  color: var(--ink);
}
h1,h2,h3,h4 {
  font-family: 'Libre Baskerville', Georgia, serif;
  color: var(--ink);
  font-weight: 400;
}
h1, h2, h3, h4,
div[data-testid="stMarkdownContainer"] h1,
div[data-testid="stMarkdownContainer"] h2,
div[data-testid="stMarkdownContainer"] h3 {
  font-family: 'Libre Baskerville', Georgia, serif !important;
  font-weight: 400 !important;
  color: var(--ink) !important;
  letter-spacing: -0.02em;
}
h2, div[data-testid="stMarkdownContainer"] h2 {
  font-size: 2.1rem !important;
  line-height: 1.2 !important;
  margin-bottom: 0.3rem !important;
  border: none !important;
}
h3 { font-size: 1.25rem !important; color: var(--ink-mid) !important; }
p, div[data-testid="stMarkdownContainer"] p {
  color: var(--ink-mid) !important;
  line-height: 1.75 !important;
  font-size: 1.02rem !important;
  font-family: 'Source Sans 3', sans-serif !important;
}
/* Kill monospace on ALL text elements */
div[data-testid="stMarkdownContainer"] * {
  font-family: 'Source Sans 3', Georgia, sans-serif;
}
div[data-testid="stMarkdownContainer"] code,
div[data-testid="stMarkdownContainer"] pre {
  font-family: 'Source Code Pro', monospace !important;
}

/* Main layout */
.main { background: var(--bg) !important; }
.main .block-container {
  padding: 2.8rem 3.5rem 5rem;
  max-width: 980px;
  background: var(--bg);
}

/* Sidebar */
section[data-testid="stSidebar"] {
  background: var(--surface) !important;
  border-right: 1px solid var(--border) !important;
}
section[data-testid="stSidebar"] * { color: var(--ink-mid) !important; }
/* Sidebar nav buttons — visible, styled as nav items */
section[data-testid="stSidebar"] .stButton button {
  display: flex !important;
  width: 100% !important;
  background: transparent !important;
  border: none !important;
  border-left: 2px solid transparent !important;
  border-radius: 3px !important;
  padding: 7px 10px !important;
  margin-bottom: 1px !important;
  text-align: left !important;
  justify-content: flex-start !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.87rem !important;
  font-weight: 400 !important;
  color: var(--ink) !important;
  cursor: pointer !important;
  transition: background 0.12s, border-color 0.12s !important;
  box-shadow: none !important;
  letter-spacing: 0 !important;
  text-transform: none !important;
}
section[data-testid="stSidebar"] .stButton button:hover {
  background: var(--bg2) !important;
  border-left-color: var(--border) !important;
  color: var(--ink) !important;
}
section[data-testid="stSidebar"] .stButton {
  margin-bottom: 0 !important;
  padding: 0 2px !important;
}
/* Section rule headers */
section[data-testid="stSidebar"] .sidebar-rule {
  display: flex;
  align-items: center;
  gap: 8px;
  font-size: 0.58rem !important;
  letter-spacing: 0.18em !important;
  text-transform: uppercase !important;
  color: var(--ink-faint) !important;
  padding: 12px 10px 5px !important;
}
section[data-testid="stSidebar"] .sidebar-rule::after {
  content: '';
  flex: 1;
  height: 1px;
  background: var(--border-soft);
}
/* Ensure sidebar labels are never washed out */
section[data-testid="stSidebar"] label { color: var(--ink) !important; }

/* Sidebar logo area */
.sidebar-logo {
  font-family: 'Libre Baskerville', Georgia, serif;
  font-size: 1.25rem;
  color: var(--ink) !important;
  font-weight: 400;
  line-height: 1.2;
}
.sidebar-logo em { color: var(--accent) !important; font-style: italic; }
.sidebar-sub {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.72rem;
  letter-spacing: 0.14em;
  text-transform: uppercase;
  color: var(--ink-faint) !important;
  margin-top: 4px;
  margin-bottom: 1.2rem;
}
.sidebar-rule {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.65rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint) !important;
  display: flex;
  align-items: center;
  gap: 8px;
  margin: 1.2rem 0 0.5rem;
}
.sidebar-rule::after {
  content: '';
  flex: 1;
  height: 1px;
  background: var(--border);
}

/* App pill badge */
.app-label {
  display: inline-block;
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.68rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint);
  border-top: 2px solid var(--border);
  border-bottom: 2px solid var(--border);
  padding: 4px 0;
  margin-bottom: 0.6rem;
  display: flex;
  align-items: center;
  gap: 10px;
}
.app-label::before, .app-label::after {
  content: '—';
  color: var(--ink-faint);
  font-size: 0.8rem;
}

/* Instruction box */
.instruction-box {
  background: var(--surface);
  border: 1px solid var(--border-soft);
  border-left: 3px solid var(--accent);
  border-radius: var(--radius);
  padding: 1rem 1.3rem;
  margin: 0.8rem 0 1.6rem;
  font-size: 0.9rem;
  color: var(--ink-mid);
  line-height: 1.8;
}
.instruction-box b  { color: var(--ink); font-weight: 600; }
.instruction-box code {
  background: var(--bg2);
  padding: 1px 5px;
  border-radius: 3px;
  font-family: 'Source Code Pro', monospace;
  font-size: 0.82rem;
  color: var(--accent);
  border: 1px solid var(--border-soft);
}
.instruction-box ul { margin: 0.3rem 0 0; padding-left: 1.2rem; }
.instruction-box li { margin-bottom: 0.25rem; }

/* Cards */
.card {
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  padding: 1.2rem 1.5rem;
  margin-bottom: 0.9rem;
  box-shadow: var(--shadow);
}
.card-accent { border-left: 3px solid var(--accent); }
.card-green  { border-left: 3px solid var(--green); background: #f2f7f4; }
.card-red    { border-left: 3px solid #c0392b; background: #fdf5f5; }
.card-amber  { border-left: 3px solid var(--amber); background: #fdf8f2; }

/* Reference items */
.ref-item {
  padding: 0.5rem 0.85rem;
  border-radius: var(--radius);
  margin-bottom: 5px;
  font-size: 0.85rem;
  border-left: 3px solid;
  line-height: 1.5;
  color: var(--ink-mid);
}
.ref-item.error   { border-color: #c0392b; background: #fdf5f5; }
.ref-item.warning { border-color: var(--amber); background: #fdf8f2; }
.ref-item.ok      { border-color: var(--green); background: #f2f7f4; }

/* Section divider */
.section-rule {
  display: flex;
  align-items: center;
  gap: 10px;
  margin: 2rem 0 1.2rem;
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.68rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint);
}
.section-rule::before, .section-rule::after {
  content: ''; flex: 1; height: 1px; background: var(--border);
}

/* Match cards */
.match-card {
  background: var(--surface);
  border: 1px solid var(--border);
  border-left: 3px solid var(--accent);
  border-radius: var(--radius);
  padding: 1rem 1.2rem;
  margin-bottom: 0.7rem;
  box-shadow: var(--shadow);
}
.match-card.accepted { border-left-color: var(--green); background: #f2f7f4; }
.match-card.skipped  { border-left-color: #c0392b; background: #fdf5f5; }
.match-sentence { font-size: 0.87rem; color: var(--ink-dim); font-style: italic; margin-bottom: 0.5rem; line-height: 1.7; }
.match-marker { font-family: monospace; font-size: 0.75rem; background: var(--bg2); color: var(--accent); padding: 2px 7px; border-radius: 3px; border: 1px solid var(--border-soft); display: inline-block; margin-bottom: 6px; }
.score-pill { font-family: monospace; font-size: 0.7rem; padding: 2px 8px; border-radius: 3px; border: 1px solid; }
.score-high { background: #f2f7f4; color: var(--green); border-color: #a8d4be; }
.score-mid  { background: #fdf8f2; color: var(--amber); border-color: #e0c8a8; }
.score-low  { background: #fdf5f5; color: #c0392b; border-color: #e0b8b8; }

/* ── Step cards (App 2 stages) ───────────────────────────────────────── */
.step-card {
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  padding: 0.9rem 1.2rem;
  margin-bottom: 0.8rem;
  box-shadow: var(--shadow);
  border-left: 4px solid var(--border);
}
.step-card.active { border-left-color: var(--accent); }
.step-card.done   { border-left-color: var(--green); background: #f4fbf7; }
.step-card.waiting{ border-left-color: var(--border); opacity: 0.7; }
.step-header {
  display: flex;
  align-items: baseline;
  gap: 12px;
  margin-bottom: 4px;
}
.step-num {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.65rem;
  font-weight: 700;
  letter-spacing: 0.12em;
  text-transform: uppercase;
  color: var(--accent);
  white-space: nowrap;
  padding: 2px 7px;
  border: 1px solid var(--accent);
  border-radius: 3px;
  line-height: 1.6;
}
.step-num.done {
  color: var(--green);
  border-color: var(--green);
  background: rgba(45,106,79,0.08);
}
.step-title {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.95rem;
  font-weight: 600;
  color: var(--ink);
}
.step-desc {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.83rem;
  color: var(--ink-dim);
  margin-top: 2px;
  line-height: 1.5;
}

/* ── Progress bar (custom) ───────────────────────────────────────── */
.progress-outer {
  background: var(--bg2);
  border-radius: 4px;
  height: 8px;
  width: 100%;
  overflow: hidden;
  margin: 6px 0;
}
.progress-inner {
  background: var(--accent);
  height: 100%;
  border-radius: 4px;
  transition: width 0.3s ease;
}

/* ── Section label (inline divider text) ────────────────────────── */
.section-label {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.65rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint);
  display: flex;
  align-items: center;
  gap: 10px;
  margin: 1.2rem 0 0.8rem;
}
.section-label::before, .section-label::after {
  content: ''; flex: 1; height: 1px; background: var(--border-soft);
}

/* ── Streamlit global overrides ─────────────────────────────────────── */
.stApp, .main, section[data-testid="stMain"],
div[data-testid="stAppViewContainer"] {
  background: var(--bg) !important;
}

/* Metrics */
div[data-testid="stMetric"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  padding: 0.8rem 1rem !important;
  box-shadow: var(--shadow) !important;
}
div[data-testid="stMetric"] label {
  color: var(--ink-dim) !important;
  font-size: 0.75rem !important;
  font-family: 'Source Sans 3', sans-serif !important;
  letter-spacing: 0.08em !important;
  text-transform: uppercase !important;
}
div[data-testid="stMetric"] div[data-testid="stMetricValue"] {
  font-family: 'Libre Baskerville', serif !important;
  font-size: 1.9rem !important;
  font-weight: 400 !important;
  color: var(--ink) !important;
}

/* Buttons — warm, visible, never black-filled */
button[kind="primary"], .stButton > button[kind="primary"] {
  background: #8b1a1a !important;
  color: #faf7f2 !important;
  border: none !important;
  border-radius: var(--radius) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.84rem !important;
  font-weight: 600 !important;
  letter-spacing: 0.08em !important;
  text-transform: uppercase !important;
  padding: 0.55rem 1.5rem !important;
  box-shadow: 0 1px 4px rgba(139,26,26,0.25) !important;
  transition: background 0.15s, box-shadow 0.15s !important;
}
button[kind="primary"]:hover {
  background: #6b1212 !important;
  box-shadow: 0 2px 8px rgba(139,26,26,0.4) !important;
}
button[kind="secondary"], .stButton > button[kind="secondary"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.84rem !important;
}
button[kind="secondary"]:hover {
  border-color: var(--accent) !important;
  color: var(--accent) !important;
}

/* File uploader — cream, warm border, never dark */
div[data-testid="stFileUploader"] {
  background: var(--surface) !important;
  border: 1.5px dashed var(--border) !important;
  border-radius: var(--radius) !important;
}
div[data-testid="stFileUploader"] * {
  color: var(--ink-mid) !important;
  background: transparent !important;
}
div[data-testid="stFileUploader"]:hover {
  border-color: var(--accent) !important;
}
/* The inner Browse-files button */
div[data-testid="stFileUploader"] button {
  background: var(--bg2) !important;
  border: 1px solid var(--border) !important;
  color: var(--ink-mid) !important;
  border-radius: var(--radius) !important;
}
div[data-testid="stFileUploader"] button:hover {
  border-color: var(--accent) !important;
  color: var(--accent) !important;
}

/* Expanders */
div[data-testid="stExpander"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  box-shadow: var(--shadow) !important;
  margin-bottom: 0.6rem !important;
}
div[data-testid="stExpander"] {
  overflow: visible !important;
}
div[data-testid="stExpander"] details > summary {
  display: flex !important;
  align-items: center !important;
  gap: 8px !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.92rem !important;
  font-weight: 500 !important;
  color: var(--ink) !important;
  background: var(--surface) !important;
  padding: 0.75rem 1rem !important;
  line-height: 1.5 !important;
  cursor: pointer !important;
  list-style: none !important;
  user-select: none !important;
}
div[data-testid="stExpander"] details > summary::-webkit-details-marker {
  display: none !important;
}
div[data-testid="stExpander"] details > summary:hover {
  background: var(--bg2) !important;
}
/* Label text inside summary */
div[data-testid="stExpander"] details > summary p,
div[data-testid="stExpander"] details > summary span,
div[data-testid="stExpander"] details > summary div {
  color: var(--ink) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.92rem !important;
  font-weight: 500 !important;
  margin: 0 !important;
}
/* Arrow SVG */
div[data-testid="stExpander"] details > summary svg {
  color: var(--ink-dim) !important;
  fill: var(--ink-dim) !important;
  width: 16px !important;
  height: 16px !important;
  flex-shrink: 0 !important;
}
/* Body text inside expander */
div[data-testid="stExpander"] div[data-testid="stMarkdownContainer"] p,
div[data-testid="stExpander"] div[data-testid="stMarkdownContainer"] * {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}
/* Expander body text — dark and readable */
div[data-testid="stExpander"] div[data-testid="stMarkdownContainer"] p,
div[data-testid="stExpander"] div[data-testid="stMarkdownContainer"] * {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}

/* Text inputs, textareas, number inputs */
div[data-testid="stTextInput"] input,
div[data-testid="stTextArea"] textarea,
div[data-testid="stNumberInput"] input {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  color: var(--ink) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.9rem !important;
}
div[data-testid="stTextInput"] input:focus,
div[data-testid="stTextArea"] textarea:focus,
div[data-testid="stNumberInput"] input:focus {
  border-color: var(--accent) !important;
  box-shadow: 0 0 0 2px rgba(139,26,26,0.12) !important;
}
div[data-testid="stTextInput"] label,
div[data-testid="stTextArea"] label,
div[data-testid="stNumberInput"] label {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.85rem !important;
}

/* Selectbox / dropdown */
div[data-testid="stSelectbox"] > div > div,
div[data-testid="stSelectbox"] > div > div > div,
div[data-testid="stSelectbox"] span,
div[data-testid="stSelectbox"] p {
  background: var(--surface) !important;
  border-color: var(--border) !important;
  border-radius: var(--radius) !important;
  color: var(--ink) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.9rem !important;
}
div[data-testid="stSelectbox"] label {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.85rem !important;
}
/* Dropdown arrow container */
div[data-testid="stSelectbox"] svg {
  color: var(--ink-dim) !important;
  fill: var(--ink-dim) !important;
}
/* Dropdown option list */
div[role="listbox"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}
div[role="option"] {
  font-family: 'Source Sans 3', sans-serif !important;
  color: var(--ink) !important;
  font-size: 0.9rem !important;
  background: var(--surface) !important;
}
div[role="option"]:hover {
  background: var(--bg2) !important;
}

/* Radio buttons */
div[data-testid="stRadio"] label {
  color: var(--ink) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.9rem !important;
  font-weight: 400 !important;
}
div[data-testid="stRadio"] label:hover { color: var(--accent) !important; }
/* All markdown text — force dark */
div[data-testid="stMarkdownContainer"],
div[data-testid="stMarkdownContainer"] p,
div[data-testid="stMarkdownContainer"] li,
div[data-testid="stMarkdownContainer"] span {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}
div[data-testid="stMarkdownContainer"] strong,
div[data-testid="stMarkdownContainer"] b {
  color: var(--ink) !important;
  font-weight: 600 !important;
}
/* Caption */
div[data-testid="stCaptionContainer"],
.stCaption, small {
  color: var(--ink-dim) !important;
  font-size: 0.79rem !important;
}

/* Checkbox */
div[data-testid="stCheckbox"] label {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.88rem !important;
}

/* Toggle */
div[data-testid="stToggle"] label {
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}

/* DataFrames / tables */
.stDataFrame {
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  background: var(--surface) !important;
}

/* Alerts / info boxes */
div[data-testid="stAlert"] {
  border-radius: var(--radius) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.88rem !important;
}

/* Tabs */
div[data-testid="stTabs"] button {
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.87rem !important;
  color: var(--ink-dim) !important;
  background: transparent !important;
}
div[data-testid="stTabs"] button[aria-selected="true"] {
  color: var(--accent) !important;
  border-bottom-color: var(--accent) !important;
}

/* Progress bar */
div[data-testid="stProgress"] > div { background: var(--bg2) !important; }
div[data-testid="stProgress"] > div > div { background: var(--accent) !important; }

/* Caption / small text */
.stCaption, small { color: var(--ink-dim) !important; font-size: 0.79rem !important; }

/* Spinner */
div[data-testid="stSpinner"] > div { border-top-color: var(--accent) !important; }

/* Download button inherits primary */
a[data-testid="stDownloadButton"] button,
div[data-testid="stDownloadButton"] button {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  color: var(--ink-mid) !important;
  border-radius: var(--radius) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.84rem !important;
}
a[data-testid="stDownloadButton"] button:hover,
div[data-testid="stDownloadButton"] button:hover {
  border-color: #3d5a6e !important;
  color: #3d5a6e !important;
}

/* Dividers */
hr { border-color: var(--border) !important; }

/* Columns gap */
div[data-testid="column"] { background: transparent !important; }

/* Toast / notifications */
div[data-testid="stToast"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  color: var(--ink) !important;
  border-radius: var(--radius) !important;
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────
TOP_N = 5
TFIDF_THRESHOLD = 0.12
PUBMED_MAX = 5
MISSING_PATTERNS = [
    r'\[CITATION\]', r'\[REF\]', r'\[ref\]', r'\[\?\]',
    r'\[citation needed\]', r'\bXXX\b', r'\[#\]', r'<citation>', r'\[ *\]',
]
CITATION_MARKERS = re.compile('|'.join(MISSING_PATTERNS), re.IGNORECASE)
MATCH_THRESHOLD = 0.28
FUZZY_THRESHOLD = 0.08
PUBMED_ESEARCH = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
PUBMED_ESUM    = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"

# ─────────────────────────────────────────────────────────────────────────────
# SHARED HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def xml_text(elem, path):
    node = elem.find(path)
    if node is None: return ""
    return " ".join(p.strip() for p in ([node.text or ""] + [c.text or "" for c in node]) if p.strip())

def fmt_ref(ref, short=False):
    aa = ref.get("authors", [])
    if aa:
        a = (aa[0].split(",")[0] if len(aa)==1 else
             f"{aa[0].split(',')[0]} & {aa[1].split(',')[0]}" if len(aa)==2 else
             f"{aa[0].split(',')[0]} et al.")
    else: a = "Unknown"
    y = ref.get("year","n.d.")
    t = ref.get("title","")[:90] + ("..." if len(ref.get("title",""))>90 else "")
    return f"{a} ({y}) — {t}" if short else f"{a} ({y}). {t}. {ref.get('journal','')}"

def score_class(s):
    return "score-high" if s>=0.20 else "score-mid" if s>=0.10 else "score-low"

def doc_to_bytes(doc):
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

def _enl_table(cursor):
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
    tables = [t[0] for t in cursor.fetchall()]
    return 'refs' if 'refs' in tables else 'enl_refs'

# ─────────────────────────────────────────────────────────────────────────────
# APP 2 LOGIC — BROKEN CITATION FIXER
# ─────────────────────────────────────────────────────────────────────────────
def analyze_docx_citations(docx_bytes):
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        raw = z.read('word/document.xml').decode('utf-8')
    total    = raw.count('ADDIN EN.CITE')
    working  = len(re.findall(r'&lt;EndNote&gt;', raw))
    empty    = len(re.findall(r'<w:instrText[^>]*> ADDIN EN\.CITE </w:instrText>', raw))
    flddata  = raw.count('<w:fldData')
    db_ids   = list(set(re.findall(r'&lt;key[^&]*db-id=&quot;([^&]+)&quot;', raw)))

    # Count unique RecNums EndNote can currently see (working fields only)
    working_rns = set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', raw))

    # Count unique RecNums locked in broken fldData (invisible to EndNote)
    import base64 as _b64
    fld_pat = re.compile(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>')
    fld_rns = set()
    for b64r in fld_pat.findall(raw):
        b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
        pad = (4-len(b64)%4)%4
        try:
            dec = _b64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
            for rn in re.findall(r'<RecNum>(\d+)</RecNum>', dec):
                fld_rns.add(rn)
        except: pass
    hidden_rns = fld_rns - working_rns   # RecNums only in fldData, invisible to EndNote

    # Count bibliography entries
    try:
        _doc = Document(io.BytesIO(docx_bytes))
        _ref_pat = re.compile(r'^\s*(\d+)[\.)\s]\s+')
        bib_count = sum(1 for p in _doc.paragraphs if _ref_pat.match(p.text.strip()))
    except: bib_count = 0

    return dict(raw=raw, total_fields=total, working=working,
                broken_empty=empty, flddata_count=flddata, db_ids=db_ids,
                working_rns=working_rns, hidden_rns=hidden_rns,
                endnote_sees=len(working_rns),
                endnote_misses=len(hidden_rns),
                bib_count=bib_count)

def fix_broken_fields(raw_xml):
    """
    Fixes three patterns of broken EndNote citation field codes:
    A: instrText says 'ADDIN EN.CITE' but is empty — data only in fldData
    B: instrText says 'ADDIN EN.CITE.DATA' — data only in fldData
    C: no instrText at all — fldChar begin immediately before fldData
    All three cause EndNote to undercount references in the bibliography.
    """
    fixes = 0

    # Pattern A: empty ADDIN EN.CITE instrText with fldData nearby
    pat_a = re.compile(
        r'(<w:instrText[^>]*>) ADDIN EN\.CITE (</w:instrText>)'
        r'([\s\S]{0,2000}?)<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>', re.DOTALL)
    def rep_a(m):
        nonlocal fixes
        io_, ic, between, b64r = m.group(1), m.group(2), m.group(3), m.group(4)
        b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
        pad = (4-len(b64)%4)%4
        try:
            dec = base64.b64decode(b64+'='*pad).decode('utf-8', errors='replace')
            dec = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', dec)
            esc = dec.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
            fixes += 1
            return f'{io_} ADDIN EN.CITE {esc}{ic}{between}<w:fldData xml:space="preserve">{b64r}</w:fldData>'
        except: return m.group(0)
    result = pat_a.sub(rep_a, raw_xml)

    # Pattern B: instrText says "ADDIN EN.CITE.DATA" instead of full XML
    pat_b = re.compile(
        r'(<w:instrText[^>]*>)\s*ADDIN EN\.CITE\.DATA\s*(</w:instrText>)'
        r'([\s\S]{0,2000}?)<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>', re.DOTALL)
    def rep_b(m):
        nonlocal fixes
        io_, ic, between, b64r = m.group(1), m.group(2), m.group(3), m.group(4)
        b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
        pad = (4-len(b64)%4)%4
        try:
            dec = base64.b64decode(b64+'='*pad).decode('utf-8', errors='replace')
            dec = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', dec)
            esc = dec.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
            fixes += 1
            return f'{io_} ADDIN EN.CITE {esc}{ic}{between}<w:fldData xml:space="preserve">{b64r}</w:fldData>'
        except: return m.group(0)
    result = pat_b.sub(rep_b, result)

    # Pattern B (reverse): fldData appears BEFORE the instrText (nested field variant)
    # Word sometimes stores: fldChar(begin) + fldData(data) + instrText(ADDIN EN.CITE.DATA) + fldChar(end)
    # The cite data is in the fldData block that precedes the marker instrText.
    def _rep_b_rev(m):
        nonlocal fixes
        b64r = m.group(1)
        b64  = b64r.replace('\r','').replace('\n','').replace(' ','')
        pad  = (4 - len(b64) % 4) % 4
        try:
            dec = base64.b64decode(b64 + '=' * pad).decode('utf-8', errors='replace')
            dec = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', dec)
            if '<EndNote>' not in dec: return m.group(0)
            esc = dec.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
            fixes += 1
            return (f'<w:fldData xml:space="preserve">{b64r}</w:fldData>'
                    f'{m.group(2)}'
                    f'{m.group(3)} ADDIN EN.CITE {esc}{m.group(4)}')
        except: return m.group(0)
    _pat_b_rev = re.compile(
        r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>'
        r'([\s\S]{0,500}?)'
        r'(<w:instrText[^>]*>)\s*ADDIN EN\.CITE\.DATA\s*(</w:instrText>)',
        re.DOTALL)
    result = _pat_b_rev.sub(_rep_b_rev, result)


    # Pattern C: no instrText at all — insert clean field before the malformed structure
    fld_pat = re.compile(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>')
    for _ in range(50):
        inserted = False
        for m in fld_pat.finditer(result):
            b64r = m.group(1)
            b64  = b64r.replace('\r','').replace('\n','').replace(' ','')
            pad  = (4-len(b64)%4)%4
            try:
                dec = base64.b64decode(b64+'='*pad).decode('utf-8', errors='replace').replace('\x00','')
            except: continue
            if '<EndNote>' not in dec: continue
            rns = set(re.findall(r'<RecNum>(\d+)</RecNum>', dec))
            if not any(f'&lt;RecNum&gt;{rn}&lt;/RecNum&gt;' not in result for rn in rns):
                continue
            begin_pos = result.rfind('<w:fldChar w:fldCharType="begin"', 0, m.start())
            if begin_pos < 0: continue
            rs1 = result.rfind('<w:r ', 0, begin_pos)
            rs2 = result.rfind('<w:r>', 0, begin_pos)
            run_start = max(rs1, rs2)
            if run_start < 0: continue
            dec_clean = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', dec)
            esc = dec_clean.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
            clean = (
                f'<w:r><w:rPr><w:noProof/><w:vertAlign w:val="superscript"/></w:rPr>'
                f'<w:fldChar w:fldCharType="begin"/></w:r>'
                f'<w:r><w:instrText xml:space="preserve"> ADDIN EN.CITE {esc}</w:instrText></w:r>'
                f'<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
                f'<w:r><w:rPr><w:noProof/><w:vertAlign w:val="superscript"/></w:rPr>'
                f'<w:t></w:t></w:r>'
                f'<w:r><w:fldChar w:fldCharType="end"/></w:r>'
            )
            result = result[:run_start] + clean + result[run_start:]
            fixes += 1
            inserted = True
            break
        if not inserted: break

    result = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', result)

    # Fix fields missing their separate marker (begin->instrText->end without separate)
    # Word requires: begin -> instrText -> separate -> display -> end
    _SEP_RUN = ('<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
                '<w:r><w:rPr><w:noProof/><w:vertAlign w:val="superscript"/></w:rPr>'
                '<w:t></w:t></w:r>')
    result = re.sub(
        r'(</w:instrText></w:r>)'
        r'((?:(?!fldCharType="separate"|fldCharType="end"|</w:instrText>).)*?)'
        r'(<w:r[^>]*>(?:[^<]|<(?!w:fldChar))*?<w:fldChar w:fldCharType="end"/>)',
        lambda _m: (_m.group(1) + _m.group(2) + _SEP_RUN + _m.group(3)
                    if 'fldCharType="separate"' not in _m.group(2) else _m.group(0)),
        result, flags=re.DOTALL
    )

    # Deduplicate w:id attributes — paragraph copying causes ID clashes
    # that make Word refuse to open the file ("unreadable content")
    from collections import Counter as _Counter
    _all_ids = re.findall(r'\bw:id="(\d+)"', result)
    if _all_ids:
        _max_id  = max(int(x) for x in _all_ids)
        _next_id = [_max_id + 1]; _seen_ids = set()
        def _fix_wid(m):
            v = m.group(2)
            if v in _seen_ids:
                nw = str(_next_id[0]); _next_id[0] += 1
                return f'{m.group(1)}{nw}{m.group(3)}'
            _seen_ids.add(v); return m.group(0)
        result = re.sub(r'(w:id=")(\d+)(")', _fix_wid, result)

    # Remove orphaned plain-text citation numbers left from merges
    # (depth-0 pure-digit/comma runs in paragraphs that have EN.CITE fields)
    # We do this at XML level to avoid needing python-docx here
    import re as _re
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    try:
        from lxml import etree as _etree
        from docx import Document as _DocX
        import io as _io
        # Quick pass: remove depth-0 digit-only runs in cite paragraphs
        _root = _etree.fromstring(result.encode('utf-8'))
        _orphans_removed = 0
        for _para in _root.findall(f'.//{{{_W}}}p'):
            _has_cite = any('EN.CITE' in (_i.text or '')
                            for _i in _para.findall(f'.//{{{_W}}}instrText'))
            if not _has_cite: continue
            _fd = 0; _to_del = []
            for _run in list(_para):
                _tag = _run.tag.split('}')[-1] if '}' in _run.tag else _run.tag
                if _tag != 'r': continue
                for _fc in _run.findall(f'.//{{{_W}}}fldChar'):
                    _ft = _fc.get(f'{{{_W}}}fldCharType', '')
                    if _ft == 'begin': _fd += 1
                    elif _ft == 'end': _fd = max(0, _fd - 1)
                if _fd > 0: continue
                _t = _run.find(f'{{{_W}}}t')
                _txt = (_t.text or '') if _t is not None else ''
                if _txt.strip() and _re.match(r'^[\d,;\s]+$', _txt.strip()):
                    _to_del.append(_run)
            for _run in _to_del:
                _para.remove(_run); _orphans_removed += 1
        if _orphans_removed:
            result = _etree.tostring(_root, xml_declaration=True,
                                     encoding='UTF-8', standalone=True).decode('utf-8')
            fixes += _orphans_removed
    except Exception: pass  # non-critical; don't break the main fix

    # Fix garbled figure numbers: merge sometimes concatenates chapter numbers
    # resulting in "33" + "34" across runs -> shows as "Fig. 3334.X"
    # (This is handled in safe_merge_documents for full merges,
    #  but also apply here for standalone repairs)

    return result, fixes


def remove_orphan_superscripts(docx_bytes):
    """
    Remove orphaned plain-text citation numbers left over after merges.

    Merges leave behind three types of plain-text citation remnants in paragraphs
    that also have working EndNote field codes:

    1. Superscript runs (vertAlign=superscript) outside field codes — the most
       obvious: numbers displayed as superscripts next to field codes.

    2. Style-based superscript runs (rStyle=citsup or rStyle=sup) outside field
       codes — same problem but using Word character styles instead of direct
       formatting.

    3. Depth-0 plain digit runs anywhere in the paragraph — numbers that are not
       superscripted but are purely citation digits/commas sitting outside field
       codes in a paragraph that has working EN.CITE fields. These are harder to
       detect visually but still produce duplicate citation display.

    All three patterns: only removes runs containing ONLY digits, commas,
    semicolons, and spaces. Never removes runs with other content.

    Returns (fixed_bytes, n_removed)
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    doc = Document(io.BytesIO(docx_bytes))
    removed = 0

    for para in doc.paragraphs:
        has_field = any(
            'EN.CITE' in (instr.text or '')
            for instr in para._p.findall(f'.//{{{W}}}instrText')
        )
        if not has_field:
            continue

        field_depth = 0
        to_remove   = []

        for run in list(para._p):
            tag = run.tag.split('}')[-1] if '}' in run.tag else run.tag
            if tag != 'r': continue

            for fc in run.findall(f'.//{{{W}}}fldChar'):
                ft = fc.get(f'{{{W}}}fldCharType', '')
                if ft == 'begin':   field_depth += 1
                elif ft == 'end':   field_depth = max(0, field_depth - 1)

            if field_depth > 0:
                continue  # inside a field — leave alone

            t    = run.find(f'{{{W}}}t')
            text = (t.text or '') if t is not None else ''
            if not text.strip():
                continue
            # Must be purely citation-like: digits, commas, semicolons, spaces
            if not re.match(r'^[\d,;\s]+$', text.strip()):
                continue

            # Pattern 1 & 2: superscript by vertAlign or character style
            rpr = run.find(f'{{{W}}}rPr')
            is_super = False
            if rpr is not None:
                va = rpr.find(f'{{{W}}}vertAlign')
                rs = rpr.find(f'{{{W}}}rStyle')
                if va is not None and va.get(f'{{{W}}}val') == 'superscript':
                    is_super = True
                if rs is not None and rs.get(f'{{{W}}}val', '') in ('citsup', 'sup', 'superscript'):
                    is_super = True

            # Pattern 3: plain depth-0 digit run in a para that has cite fields
            # (catches runs that survived the merge without any superscript formatting)
            # Always remove any depth-0 pure-digit run in a cite paragraph.
            to_remove.append(run)

        for run in to_remove:
            para._p.remove(run)
            removed += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), removed


def extract_karol_db_id(xml_bytes):
    try:
        content = xml_bytes.decode('utf-8',errors='replace')
        m = re.search(r'db-id="([a-z0-9]{20,45})"',content,re.IGNORECASE)
        return m.group(1) if m else None
    except: return None

def get_karol_rec_nums(enl_bytes):
    with tempfile.NamedTemporaryFile(suffix='.enl',delete=False) as f:
        f.write(enl_bytes); tmp=f.name
    try:
        conn=sqlite3.connect(tmp); cursor=conn.cursor()
        tbl=_enl_table(cursor)
        cursor.execute(f'SELECT id,author,year,title FROM {tbl} WHERE trash_state=0 OR trash_state IS NULL')
        rows=cursor.fetchall(); conn.close()
        return {str(r[0]):{'id':str(r[0]),'author':r[1] or '','year':str(r[2] or ''),'title':r[3] or ''} for r in rows}
    except: return {}
    finally: os.unlink(tmp)

def check_missing_from_karol(raw_xml, karol_ids):
    all_rns = set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', raw_xml))
    return [rn for rn in all_rns if rn not in karol_ids]

def patch_db_ids(raw_xml, old_ids, new_id):
    result=raw_xml; replaced=0
    for old in old_ids:
        if old!=new_id:
            count=result.count(old); result=result.replace(old,new_id); replaced+=count
    return result, replaced

def build_fixed_docx(original_bytes, fixed_xml):
    with zipfile.ZipFile(io.BytesIO(original_bytes)) as z:
        all_files={n:z.read(n) for n in z.namelist()}
    all_files['word/document.xml']=fixed_xml.encode('utf-8')
    buf=io.BytesIO()
    with zipfile.ZipFile(buf,'w',zipfile.ZIP_DEFLATED) as zout:
        for n,d in all_files.items(): zout.writestr(n,d)
    buf.seek(0); return buf.read()

def extract_traveling_library_xml(docx_bytes):
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        raw=z.read('word/document.xml').decode('utf-8')
    all_cx=[]
    for m in re.findall(r'ADDIN EN\.CITE &lt;EndNote&gt;([\s\S]+?)&lt;/EndNote&gt;',raw):
        all_cx.append(html_module.unescape(f'<EndNote>{m}</EndNote>'))
    for b64r in re.findall(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>',raw):
        b64=b64r.replace('\r','').replace('\n','').replace(' ','')
        pad=(4-len(b64)%4)%4
        try:
            dec=base64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
            if '<EndNote>' in dec: all_cx.append(dec)
        except: pass
    from lxml import etree
    import html as hmod
    traveling={}
    for cx in all_cx:
        try:
            if not cx.startswith('<EndNote>'): cx=f'<EndNote>{cx}</EndNote>'
            root=etree.fromstring(cx.encode('utf-8'))
            for cite in root.findall('.//Cite'):
                rn=cite.findtext('RecNum') or ''
                rec=cite.find('record')
                if rn and rec is not None and rn not in traveling:
                    traveling[rn]=rec
        except: pass
    def gt(elem,tag):
        n=elem.find(f'.//{tag}')
        return ''.join(n.itertext()).strip() if n is not None else ''
    output='<?xml version="1.0" encoding="UTF-8"?>\n<xml>\n  <records>\n'
    for rn,record in sorted(traveling.items(),key=lambda x:int(x[0]) if x[0].isdigit() else 9999):
        rte=record.find('.//ref-type')
        rtn=rte.get('name','Journal Article') if rte is not None else 'Journal Article'
        rtv=rte.text if rte is not None else '17'
        authors=[''.join(a.itertext()).strip() for a in record.findall('.//contributors/authors/author')]
        secauths=[''.join(a.itertext()).strip() for a in record.findall('.//contributors/secondary-authors/author')]
        r=f'    <record>\n      <rec-number>{rn}</rec-number>\n      <ref-type name="{hmod.escape(rtn)}">{rtv}</ref-type>\n'
        if authors:
            r+='      <contributors>\n        <authors>\n'
            for a in authors:
                if a: r+=f'          <author>{hmod.escape(a)}</author>\n'
            r+='        </authors>\n'
            if secauths:
                r+='        <secondary-authors>\n'
                for a in secauths:
                    if a: r+=f'          <author>{hmod.escape(a)}</author>\n'
                r+='        </secondary-authors>\n'
            r+='      </contributors>\n'
        r+='      <titles>\n'
        for tag,xmltag in [('title','title'),('secondary-title','secondary-title'),('tertiary-title','tertiary-title')]:
            v=gt(record,tag)
            if v: r+=f'        <{xmltag}>{hmod.escape(v)}</{xmltag}>\n'
        r+='      </titles>\n'
        for tag,xmltag in [('year','dates><year'),('volume','volume'),('number','number'),
                           ('pages','pages'),('edition','edition'),('publisher','publisher'),
                           ('pub-location','pub-location'),('abstract','abstract')]:
            v=gt(record,tag)
            if v:
                if xmltag=='dates><year': r+=f'      <dates><year>{hmod.escape(v)}</year></dates>\n'
                else: r+=f'      <{xmltag}>{hmod.escape(v)}</{xmltag}>\n'
        kws=[(''.join(k.itertext()).strip()) for k in record.findall('.//keyword')]
        if any(kws):
            r+='      <keywords>\n'
            for kw in kws:
                if kw: r+=f'        <keyword>{hmod.escape(kw)}</keyword>\n'
            r+='      </keywords>\n'
        r+='    </record>\n'
        output+=r
    output+='  </records>\n</xml>\n'
    return output, len(traveling)

def remap_traveling_citations(docx_bytes, enl_bytes):
    from lxml import etree
    import html as hm
    with tempfile.NamedTemporaryFile(suffix='.enl',delete=False) as f:
        f.write(enl_bytes); tmp=f.name
    try:
        conn=sqlite3.connect(tmp); cursor=conn.cursor()
        tbl=_enl_table(cursor)
        cursor.execute(f'SELECT id,author,year,title FROM {tbl} WHERE trash_state=0 OR trash_state IS NULL')
        karol_rows=cursor.fetchall(); conn.close()
    finally: os.unlink(tmp)
    def norm(s): return re.sub(r'[^a-z0-9]','',s.lower()) if s else ''
    def alast(s): return norm(s.split('\r')[0].split('\n')[0].strip().split(',')[0]) if s else ''
    karol_ids=set(str(r[0]) for r in karol_rows)
    by_ay={}
    for row in karol_rows:
        k=(alast(row[1] or ''),norm(str(row[2] or '')))
        by_ay.setdefault(k,[]).append(row)
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        raw=z.read('word/document.xml').decode('utf-8')
        all_files={n:z.read(n) for n in z.namelist()}
    report=[]; remap={}
    for cx_esc in re.findall(r'ADDIN EN\.CITE &lt;EndNote&gt;([\s\S]+?)&lt;/EndNote&gt;',raw):
        cx=hm.unescape(f'<EndNote>{cx_esc}</EndNote>')
        try:
            root=etree.fromstring(cx.encode('utf-8'))
            for cite in root.findall('.//Cite'):
                rn=cite.findtext('RecNum') or ''
                au=cite.findtext('Author') or ''
                yr=cite.findtext('Year') or ''
                if rn in karol_ids or rn in remap: continue
                title=''
                rec=cite.find('record')
                if rec is not None:
                    te=rec.find('.//title')
                    if te is not None: title=''.join(te.itertext())
                key=(alast(au),norm(yr)); cands=by_ay.get(key,[])
                matched=None
                if len(cands)==1: matched=str(cands[0][0])
                elif len(cands)>1:
                    for c in cands:
                        if norm(title)[:25] and norm(str(c[3] or ''))[:25]==norm(title)[:25]:
                            matched=str(c[0]); break
                    if not matched: matched=str(cands[0][0])
                remap[rn]=matched
                report.append({'status':'remapped' if matched else 'not_found',
                               'old_rec_num':rn,'new_rec_num':matched,
                               'author':au,'year':yr,'title':title[:80]})
        except: pass
    if not remap: return docx_bytes, report
    fixed=raw
    for old,new in remap.items():
        if new:
            fixed=fixed.replace(f'&lt;RecNum&gt;{old}&lt;/RecNum&gt;',
                                 f'&lt;RecNum&gt;{new}&lt;/RecNum&gt;')
    def fix_fld(m):
        b64r=m.group(1)
        b64=b64r.replace('\r','').replace('\n','').replace(' ','')
        pad=(4-len(b64)%4)%4
        try:
            dec=base64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
            mod=dec
            for old,new in remap.items():
                if new: mod=mod.replace(f'<RecNum>{old}</RecNum>',f'<RecNum>{new}</RecNum>')
            if mod!=dec:
                nb64=base64.b64encode(mod.encode('utf-8')).decode('ascii')
                wrapped='\r\n'.join(nb64[i:i+76] for i in range(0,len(nb64),76))
                return f'<w:fldData xml:space="preserve">{wrapped}</w:fldData>'
        except: pass
        return m.group(0)
    fixed=re.compile(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>').sub(fix_fld,fixed)
    all_files['word/document.xml']=fixed.encode('utf-8')
    buf=io.BytesIO()
    with zipfile.ZipFile(buf,'w',zipfile.ZIP_DEFLATED) as zout:
        for n,d in all_files.items(): zout.writestr(n,d)
    buf.seek(0)
    return buf.read(), report

def generate_vba_macro(doc_name):
    return f"""' EndNote Citation Relinker — Word VBA Macro
' Generated: {datetime.now():%Y-%m-%d %H:%M}
' HOW TO USE:
'  1. Open "{doc_name}" in Word with your EndNote library open
'  2. Press Alt+F11 → Insert → Module → paste this macro
'  3. Press Alt+F8 → select RelinkAllCitations → Run
Sub RelinkAllCitations()
    Dim oDoc As Document
    Set oDoc = ActiveDocument
    If oDoc Is Nothing Then MsgBox "No document open.", vbCritical: Exit Sub
    Dim sBackup As String
    sBackup = oDoc.Path & "\\{Path(doc_name).stem}_BACKUP_" & Format(Now,"YYYYMMDD_HHMMSS") & ".docx"
    oDoc.SaveAs2 sBackup, wdFormatXMLDocument
    MsgBox "Backup saved: " & sBackup, vbInformation
    On Error Resume Next
    Application.Run "EndNote.UnformatAll"
    If Err.Number <> 0 Then Err.Clear: Application.Run "EndNote.FormatAll": GoTo done
    Err.Clear
    Application.Run "EndNote.FormatAll"
    If Err.Number <> 0 Then Err.Clear: Application.Run "EndNote.UpdateAll"
    done:
    On Error GoTo 0
    oDoc.Save
    MsgBox "Done! Citations re-linked to your library." & vbCrLf & _
           "If any citations are yellow, right-click → Edit Citation → Find.", vbInformation
End Sub
"""

# ─────────────────────────────────────────────────────────────────────────────
# APP 1 LOGIC — CITATION REPAIR
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def parse_endnote_xml_bytes(xml_bytes):
    root=ET.fromstring(xml_bytes); refs=[]
    for rec in root.iter("record"):
        authors=[]
        for a in rec.findall(".//contributors/authors/author"):
            name=" ".join(p.strip() for p in ([a.text or ""]+[c.text or "" for c in a]) if p.strip())
            if name: authors.append(name)
        title=xml_text(rec,".//titles/title")
        journal=xml_text(rec,".//periodical/full-title") or xml_text(rec,".//periodical/abbr-1")
        year=xml_text(rec,".//dates/year"); abstract=xml_text(rec,".//abstract")
        if not title: continue
        corpus=" ".join(filter(None,[title,abstract,journal,year]))
        refs.append(dict(authors=authors,title=title,journal=journal,year=year,corpus=corpus))
    return refs

@st.cache_resource(show_spinner=False)
def build_tfidf(corpora_tuple):
    vec=TfidfVectorizer(ngram_range=(1,2),sublinear_tf=True,max_features=50000)
    mat=vec.fit_transform(list(corpora_tuple)); return vec,mat

def match_sentence(sentence,vec,mat,refs,top_n=TOP_N):
    sv=vec.transform([sentence]); sims=cosine_similarity(sv,mat)[0]
    idx=sims.argsort()[::-1][:top_n]
    return [dict(ref=refs[i],score=float(sims[i])) for i in idx]

def extract_flagged(docx_bytes):
    doc=Document(io.BytesIO(docx_bytes)); flagged=[]
    for pi,para in enumerate(doc.paragraphs):
        text=para.text
        if not text.strip(): continue
        for m in CITATION_MARKERS.finditer(text):
            sents=re.split(r'(?<=[.!?])\s+',text); cum,target=0,text
            for s in sents:
                if cum+len(s)>=m.start(): target=s; break
                cum+=len(s)+1
            flagged.append(dict(para_idx=pi,sentence=target,marker=m.group(),para_text=text))
    return flagged,doc

def author_label(ref):
    aa=ref.get("authors",[]); last=aa[0].split(",")[0].strip().split()[-1] if aa else "Ref"
    return (last+" "+ref.get("year","")).strip()

def insert_superscript(para,marker,label):
    if marker not in para.text: return False
    combined,run_map="",[]
    for run in para.runs:
        s=len(combined); combined+=run.text; run_map.append((s,s+len(run.text),run))
    pos=combined.find(marker)
    if pos==-1: return False
    for (s,e,run) in run_map:
        if s<=pos<e:
            before=run.text[:pos-s]; after=run.text[pos-s+len(marker):]
            run.text=before
            rPr=OxmlElement("w:rPr"); va=OxmlElement("w:vertAlign")
            va.set(qn("w:val"),"superscript"); rPr.append(va)
            nr=OxmlElement("w:r"); nr.append(deepcopy(rPr))
            t=OxmlElement("w:t"); t.text=f"[{label}]"
            t.set("{http://www.w3.org/XML/1998/namespace}space","preserve"); nr.append(t)
            run._r.addnext(nr)
            if after:
                tr=OxmlElement("w:r"); tt=OxmlElement("w:t")
                tt.text=after; tt.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
                tr.append(tt); nr.addnext(tr)
            return True
    return False

def write_repair_report(decisions):
    doc=Document(); doc.add_heading("Citation Repair Report",0)
    doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M}")
    accepted=[d for d in decisions if d["action"]=="accepted"]
    skipped=[d for d in decisions if d["action"]=="skipped"]
    doc.add_paragraph(f"Total:{len(decisions)} | Accepted:{len(accepted)} | Skipped:{len(skipped)}")
    if accepted:
        doc.add_heading("Accepted",1)
        for d in accepted:
            p=doc.add_paragraph(style="List Bullet")
            p.add_run(f"Marker:{d['marker']}\n").bold=True
            p.add_run(f"Context:{d['sentence'][:200]}\n")
            p.add_run(f"Inserted:{fmt_ref(d['ref'])}\nScore:{d['score']:.3f}")
    if skipped:
        doc.add_heading("Skipped",1)
        for d in skipped:
            p=doc.add_paragraph(style="List Bullet")
            r=p.add_run("NEEDS REVIEW\n"); r.font.color.rgb=RGBColor(0xC0,0,0)
            p.add_run(f"Marker:{d['marker']}\nContext:{d['sentence'][:200]}\n")
    return doc_to_bytes(doc)

# ─────────────────────────────────────────────────────────────────────────────
# APP 3 LOGIC — REFERENCE COMPARATOR
# ─────────────────────────────────────────────────────────────────────────────
def load_ref_file(f):
    name=f.name; data=f.read()
    if name.endswith(".xml"):
        refs=parse_endnote_xml_bytes(data); return refs,name
    elif name.endswith(".docx"):
        doc=Document(io.BytesIO(data)); refs=[]; in_refs=False
        pat=re.compile(r'^\s*\d+[\.\)]\s+(.+)')
        for para in doc.paragraphs:
            text=para.text.strip()
            if not text: continue
            if re.match(r'^(references?|bibliography)$',text,re.IGNORECASE): in_refs=True; continue
            if in_refs or pat.match(text):
                in_refs=True; m=pat.match(text); rt=m.group(1) if m else text
                ym=re.search(r'\b(19|20)\d{2}\b',rt)
                refs.append(dict(authors=[],title=rt[:200],journal="",
                                 year=ym.group(0) if ym else "",corpus=rt,id=str(len(refs)+1)))
        return refs,name
    elif name.endswith(".txt"):
        content=data.decode("utf-8",errors="replace"); refs=[]
        for i,block in enumerate(re.split(r'\n{2,}',content)):
            block=block.strip()
            if len(block)<20: continue
            m=re.match(r'^\d+[\.\)]\s+(.*)',block,re.DOTALL); text=m.group(1) if m else block
            ym=re.search(r'\b(19|20)\d{2}\b',text)
            refs.append(dict(authors=[],title=text[:200],journal="",
                             year=ym.group(0) if ym else "",corpus=text,id=str(i)))
        return refs,name
    return [],name

# ─────────────────────────────────────────────────────────────────────────────
# APP 4 LOGIC — DOCUMENT MERGER
# ─────────────────────────────────────────────────────────────────────────────
def analyze_merge_damage(merged_bytes):
    from lxml import etree as _etree
    import base64 as _b64
    with zipfile.ZipFile(io.BytesIO(merged_bytes)) as z:
        raw=z.read('word/document.xml').decode('utf-8')
    total=raw.count('ADDIN EN.CITE')
    working=len(re.findall(r'&lt;EndNote&gt;',raw))
    empty=len(re.findall(r'<w:instrText[^>]*> ADDIN EN\.CITE </w:instrText>',raw))
    begins=raw.count('fldCharType="begin"')
    separates=raw.count('fldCharType="separate"')
    ends=raw.count('fldCharType="end"')
    has_tracked='<w:ins ' in raw or '<w:del ' in raw
    ins_count=raw.count('<w:ins ')
    del_count=raw.count('<w:del ')
    db_ids=list(set(re.findall(r'&lt;key[^&]*db-id=&quot;([^&]+)&quot;',raw)))

    # RecNums in working field codes (what EndNote currently sees)
    working_rns = set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', raw))

    # RecNums locked in broken fldData (invisible to EndNote)
    fld_pat = re.compile(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>')
    fld_rns = set()
    for b64r in fld_pat.findall(raw):
        b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
        pad = (4-len(b64)%4)%4
        try:
            dec = _b64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
            for rn in re.findall(r'<RecNum>(\d+)</RecNum>', dec):
                fld_rns.add(rn)
        except: pass
    hidden_rns = fld_rns - working_rns

    # Count ALL superscript number runs
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    try:
        root = _etree.fromstring(raw.encode('utf-8'))
        super_nums = set()
        for r in root.iter(f'{{{W}}}r'):
            rpr = r.find(f'{{{W}}}rPr')
            if rpr is None: continue
            va = rpr.find(f'{{{W}}}vertAlign')
            if va is None: continue
            if va.get(f'{{{W}}}val') != 'superscript': continue
            texts = []
            for t in r.iter(f'{{{W}}}t'):
                if t.text: texts.append(t.text)
            for t in r.iter(f'{{{W}}}delText'):
                if t.text: texts.append(t.text)
            text = ''.join(texts).strip()
            for part in re.split(r'[,;\s]+', text):
                p = part.strip().rstrip('.')
                if p.isdigit(): super_nums.add(int(p))
    except: super_nums = set()

    # Bibliography entries
    ref_pat = re.compile(r'^\s*(\d+)[\.)\s]\s+')
    try:
        doc = Document(io.BytesIO(merged_bytes))
        bib_nums = set()
        for p in doc.paragraphs:
            m = ref_pat.match(p.text.strip())
            if m: bib_nums.add(int(m.group(1)))
    except: bib_nums = set()

    lost_count  = len(bib_nums - super_nums) if bib_nums else 0
    bib_count   = len(bib_nums)
    cited_count = len(super_nums)

    # Check for w:id duplicates (cause Word to reject file with "unreadable content")
    from collections import Counter as _Ctr
    all_wids = re.findall(r'\bw:id="(\d+)"', raw)
    dup_ids  = len({k for k,v in _Ctr(all_wids).items() if v > 1})

    # Check footnote/comment reference integrity
    fn_refs  = set(re.findall(r'<w:footnoteReference\b[^>]*w:id="(\d+)"', raw))
    com_refs = len(re.findall(r'<w:comment(?:RangeStart|RangeEnd|Reference)\b', raw))

    return dict(raw=raw, total_en=total, with_data=working, empty_cite=empty,
                begins=begins, separates=separates, ends=ends,
                balanced=(begins==separates==ends),
                has_tracked=has_tracked, ins_count=ins_count, del_count=del_count,
                db_ids=db_ids,
                working_rns=working_rns, hidden_rns=hidden_rns,
                endnote_sees=len(working_rns),
                endnote_misses=len(hidden_rns),
                super_nums=super_nums, cited_count=cited_count,
                bib_count=bib_count, lost_in_merge=lost_count,
                dup_ids=dup_ids, fn_refs=fn_refs, orphan_comments=com_refs)

def repair_post_merge_citations(merged_bytes,original_bytes=None):
    analysis=analyze_merge_damage(merged_bytes)
    raw=analysis['raw']
    with zipfile.ZipFile(io.BytesIO(merged_bytes)) as z:
        all_files={n:z.read(n) for n in z.namelist()}
    report=dict(analysis=analysis,steps=[],citations_before=analysis['with_data'])
    fixed=raw
    # Accept tracked changes safely
    if analysis['has_tracked']:
        def rescue_del(m):
            dc=m.group(0)
            if 'ADDIN EN.CITE' in dc:
                runs=re.findall(r'<w:r[^>]*>.*?</w:r>',dc,re.DOTALL)
                cite_runs=[r.replace('<w:delText','<w:t').replace('</w:delText>','</w:t>')
                           for r in runs if any(x in r for x in ['fldChar','instrText','fldData','ADDIN'])]
                if cite_runs: return ''.join(cite_runs)
            return ''
        fixed=re.compile(r'<w:del\b[^>]*>[\s\S]*?</w:del>',re.DOTALL).sub(rescue_del,fixed)
        def accept_ins(m):
            inner=re.sub(r'^<w:ins[^>]*>','',m.group(0)); inner=re.sub(r'</w:ins>$','',inner)
            return inner
        fixed=re.compile(r'<w:ins\b[^>]*>[\s\S]*?</w:ins>',re.DOTALL).sub(accept_ins,fixed)
        report['steps'].append('track_changes_accepted')
    # Restore broken fields
    if analysis['empty_cite']>0:
        fixed,n_fixed=fix_broken_fields(fixed)
        if n_fixed>0: report['steps'].append(f'restored_{n_fixed}_fields')
    # Compare against original
    if original_bytes:
        with zipfile.ZipFile(io.BytesIO(original_bytes)) as z:
            orig_raw=z.read('word/document.xml').decode('utf-8')
        orig_rns=set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;',orig_raw))
        merged_rns=set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;',fixed))
        lost=orig_rns-merged_rns
        report['lost_rec_nums']=list(lost)
        if lost: report['steps'].append(f'{len(lost)}_citations_lost')
    report['citations_after']=len(re.findall(r'&lt;EndNote&gt;',fixed))
    fixed=fixed.replace('\x00','')
    all_files['word/document.xml']=fixed.encode('utf-8')
    buf=io.BytesIO()
    with zipfile.ZipFile(buf,'w',zipfile.ZIP_DEFLATED) as zout:
        for n,d in all_files.items(): zout.writestr(n,d)
    buf.seek(0)
    return buf.read(),report

# ─────────────────────────────────────────────────────────────────────────────
# APP 5 LOGIC
# ─────────────────────────────────────────────────────────────────────────────
# APP 4 LOGIC — SAFE DOCUMENT MERGE
# ─────────────────────────────────────────────────────────────────────────────

def _para_sig(text):
    """Normalised fingerprint for fuzzy paragraph matching."""
    return re.sub(r'[^a-z0-9]', '', text.lower())[:120]


def _para_has_citations(para, W):
    """Return True if a paragraph has any citation content."""
    # EndNote field codes
    for f in para._p.findall(f'.//{{{W}}}instrText'):
        if 'EN.CITE' in (f.text or ''):
            return True
    # Superscript number runs
    for run in para.runs:
        rpr = run._r.find(f'{{{W}}}rPr')
        if rpr is None: continue
        va = rpr.find(f'{{{W}}}vertAlign')
        if va is not None and va.get(f'{{{W}}}val') == 'superscript':
            if any(c.isdigit() for c in run.text):
                return True
    return False


def _extract_cite_runs(para, W):
    """
    Extract citation-related XML elements from a paragraph:
    field begin/instrText/separate/end sequences and superscript runs.
    Returns list of lxml elements.
    """
    cite_elems = []
    p_elem = para._p
    # Walk runs looking for field chars and superscript runs
    in_field = False
    for child in p_elem:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            # Check for fldChar begin
            for fc in child.findall(f'.//{{{W}}}fldChar'):
                ft = fc.get(f'{{{W}}}fldCharType', '')
                if ft == 'begin':
                    in_field = True
                elif ft == 'end':
                    in_field = False
            if in_field:
                cite_elems.append(child)
            else:
                # Check for superscript
                rpr = child.find(f'{{{W}}}rPr')
                if rpr is not None:
                    va = rpr.find(f'{{{W}}}vertAlign')
                    if va is not None and va.get(f'{{{W}}}val') == 'superscript':
                        cite_elems.append(child)
        elif tag in ('bookmarkStart', 'bookmarkEnd'):
            pass  # skip
    return cite_elems


def safe_merge_documents(new_bytes, old_bytes):
    """
    Safely merge a citation-damaged new document with a citation-intact old document.

    1. Fixes any broken citation fields in the old document first (fldData recovery)
    2. Builds a paragraph signature index of the old document
    3. For each paragraph in the new document that lacks citations:
       - Finds the matching paragraph in the old document
       - Copies all citation field code runs and superscript runs across
    4. New-only paragraphs (no match in old) are kept unchanged
    Returns (merged_bytes, report_dict)
    """
    from copy import deepcopy

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Step 1: Fix broken fields in old doc before using it as source
    with zipfile.ZipFile(io.BytesIO(old_bytes)) as z:
        old_raw   = z.read('word/document.xml').decode('utf-8')
        old_files = {n: z.read(n) for n in z.namelist()}
    old_fixed, n_fixed = fix_broken_fields(old_raw)
    old_fixed = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', old_fixed)
    old_files['word/document.xml'] = old_fixed.encode('utf-8')
    old_buf = io.BytesIO()
    with zipfile.ZipFile(old_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for n, d in old_files.items(): zout.writestr(n, d)
    old_buf.seek(0)
    doc_old = Document(old_buf)
    doc_new = Document(io.BytesIO(new_bytes))

    def _has_cites(para):
        for f in para._p.findall(f'.//{{{W}}}instrText'):
            if 'EN.CITE' in (f.text or ''): return True
        for run in para.runs:
            rpr = run._r.find(f'{{{W}}}rPr')
            if rpr is None: continue
            va = rpr.find(f'{{{W}}}vertAlign')
            if va is not None and va.get(f'{{{W}}}val') == 'superscript':
                if any(c.isdigit() for c in run.text): return True
        return False

    def _cite_runs(para):
        """Return all citation-related run elements from a paragraph."""
        runs = para._p.findall(f'{{{W}}}r')
        result = []; in_field = False
        for run in runs:
            is_cite = False
            for fc in run.findall(f'.//{{{W}}}fldChar'):
                ft = fc.get(f'{{{W}}}fldCharType', '')
                if ft == 'begin':   in_field = True;  is_cite = True
                elif ft == 'end':   is_cite = True;   in_field = False
                else:                is_cite = True
            if in_field and not is_cite: is_cite = True
            rpr = run.find(f'{{{W}}}rPr')
            if not is_cite and rpr is not None:
                va = rpr.find(f'{{{W}}}vertAlign')
                if va is not None and va.get(f'{{{W}}}val') == 'superscript':
                    t = run.find(f'{{{W}}}t')
                    if t is not None and any(c.isdigit() for c in (t.text or '')):
                        is_cite = True
            if is_cite: result.append(run)
        return result

    # Build old paragraph index
    old_sig_map = {}
    for pi, para in enumerate(doc_old.paragraphs):
        text = para.text.strip()
        if len(text) < 15: continue
        sig = re.sub(r'[^a-z0-9]', '', text.lower())[:80]
        old_sig_map.setdefault(sig, []).append((pi, para))

    report = dict(total_new_paras=len(doc_new.paragraphs),
                  matched=0, citations_restored=0,
                  already_had_cites=0, unmatched=0,
                  old_fields_fixed=n_fixed, details=[])

    for pi, para_new in enumerate(doc_new.paragraphs):
        text = para_new.text.strip()
        if len(text) < 15: continue

        sig = re.sub(r'[^a-z0-9]', '', text.lower())[:80]
        candidates = old_sig_map.get(sig, [])

        # Partial match fallback (first 50 chars)
        if not candidates and len(sig) > 40:
            for osig, olist in old_sig_map.items():
                if sig[:50] == osig[:50]:
                    candidates = olist; break

        if not candidates:
            report['unmatched'] += 1; continue

        report['matched'] += 1
        _, para_old = candidates[0]

        new_has = _has_cites(para_new)
        old_has = _has_cites(para_old)

        if new_has:
            report['already_had_cites'] += 1; continue
        if not old_has:
            continue

        # Copy citation runs from old → new
        cite_elems = _cite_runs(para_old)
        if not cite_elems: continue

        new_runs = para_new._p.findall(f'{{{W}}}r')
        if not new_runs: continue

        last = new_runs[-1]
        for elem in cite_elems:
            new_elem = deepcopy(elem)
            last.addnext(new_elem)
            last = new_elem

        report['citations_restored'] += 1
        report['details'].append({
            'para_idx':     pi,
            'text_preview': text[:80],
            'cites_added':  len(cite_elems),
        })

    # Fix garbled figure numbers: merge sometimes concatenates chapter numbers
    # e.g. "Fig. 33" + "34" across adjacent runs creates "Fig. 3334.X"
    # Fix by removing the extra chapter number from the second run
    fig_fixes = 0
    for para in doc_new.paragraphs:
        runs = list(para.runs)
        i = 0
        while i < len(runs) - 1:
            t1 = runs[i].text or ''
            t2 = runs[i+1].text or ''
            if t1.endswith('33') and t2.startswith('34'):
                new_t2 = t2[2:]
                t_elem = runs[i+1]._r.find(f'{{{W}}}t')
                if t_elem is not None:
                    t_elem.text = new_t2
                    if not new_t2:
                        runs[i+1]._r.getparent().remove(runs[i+1]._r)
                    fig_fixes += 1
            i += 1
    if fig_fixes:
        report['fig_number_fixes'] = fig_fixes

    # Remove orphaned plain superscripts left by the merge
    buf_check = io.BytesIO(); doc_new.save(buf_check); buf_check.seek(0)
    _, orphans_removed = remove_orphan_superscripts(buf_check.read())
    if orphans_removed:
        # Re-load to apply removal
        buf_check.seek(0)
        clean_bytes, _ = remove_orphan_superscripts(buf_check.read())
        doc_new = Document(io.BytesIO(clean_bytes))
        report['orphan_superscripts_removed'] = orphans_removed

    buf = io.BytesIO()
    doc_new.save(buf)
    buf.seek(0)

    # Post-save: fix duplicate w:id attributes caused by copying paragraph XML
    import zipfile as _zf
    raw_buf = buf.read()
    with _zf.ZipFile(io.BytesIO(raw_buf)) as _z:
        _doc_xml  = _z.read('word/document.xml').decode('utf-8')
        _all_files = {n: _z.read(n) for n in _z.namelist()}
    _all_ids = re.findall(r'\bw:id="(\d+)"', _doc_xml)
    if _all_ids:
        _max_id  = max(int(x) for x in _all_ids)
        _next    = [_max_id + 1]; _seen = set()
        def _fix_id2(m):
            v = m.group(2)
            if v in _seen:
                nw = str(_next[0]); _next[0] += 1
                return f'{m.group(1)}{nw}{m.group(3)}'
            _seen.add(v); return m.group(0)
        _doc_xml = re.sub(r'(w:id=")(\d+)(")', _fix_id2, _doc_xml)
        _all_files['word/document.xml'] = _doc_xml.encode('utf-8')
        _out = io.BytesIO()
        with _zf.ZipFile(_out, 'w', _zf.ZIP_DEFLATED) as _zout:
            for _n, _d in _all_files.items(): _zout.writestr(_n, _d)
        _out.seek(0)
        raw_buf = _out.read()

    return raw_buf, report

# APP 5 — CITATION RENUMBERING
# ─────────────────────────────────────────────────────────────────────────────
def _apply_superscript_mapping(doc, mapping):
    """Apply old->new number mapping to all superscript runs in a document."""
    for para in doc.paragraphs:
        for run in para.runs:
            rpr = run._r.find(qn('w:rPr'))
            if rpr is None: continue
            va = rpr.find(qn('w:vertAlign'))
            if va is None or va.get(qn('w:val')) != 'superscript': continue
            text  = run.text
            sep   = ';' if ';' in text else ','
            parts = re.split(r'[,;]', text)
            new_parts = []
            for p in parts:
                ps = p.strip()
                if ps.isdigit():
                    new_parts.append(str(mapping.get(int(ps), int(ps))))
                else:
                    new_parts.append(p)
            run.text = sep.join(new_parts)


def _renumber_bib_paras(doc, bib, mapping, bib_sorted=None):
    """
    Update leading numbers in bibliography paragraphs.
    If bib_sorted is provided, also physically re-sorts them in that order.
    """
    W     = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body  = doc.element.body
    all_p = list(body.findall(f'{{{W}}}p'))

    # Update leading number text
    for pi, old_num, _ in bib:
        new_num = mapping.get(old_num, old_num)
        para    = doc.paragraphs[pi]
        for run in para.runs:
            if f"{old_num}." in run.text:
                run.text = run.text.replace(f"{old_num}.", f"{new_num}.", 1)
                break

    # Re-sort paragraphs if sort order given
    if bib_sorted and len(bib) > 1:
        anchor = all_p[bib[0][0]]
        for entry in bib_sorted:
            elem = all_p[entry[0]]
            anchor.addprevious(elem)


def renumber_citations_alpha(docx_bytes):
    """
    Renumbers citations alphabetically by first author last name (A=1, B=2...).
    Parses the bibliography, sorts A-Z, assigns new numbers, updates
    inline superscripts, and re-sorts the bibliography.
    Returns (fixed_bytes, {old_num: new_num})
    """
    doc     = Document(io.BytesIO(docx_bytes))
    ref_pat = re.compile(r'^\s*(\d+)\.\s+(.+)')
    bib     = [(pi, int(m.group(1)), para.text.strip())
               for pi, para in enumerate(doc.paragraphs)
               if (m := ref_pat.match(para.text.strip()))]

    if not bib:
        return docx_bytes, {}

    def sort_key(text):
        text       = re.sub(r'^\d+\.\s+', '', text).strip()
        first_auth = re.split(r';|,\s+[A-Z]', text)[0].strip()
        last_name  = first_auth.split(',')[0].strip().split()[-1] if first_auth else text[:20]
        return last_name.lower()

    bib_sorted = sorted(bib, key=lambda x: sort_key(x[2]))
    mapping    = {entry[1]: idx + 1 for idx, entry in enumerate(bib_sorted)}

    _apply_superscript_mapping(doc, mapping)
    _renumber_bib_paras(doc, bib, mapping, bib_sorted=bib_sorted)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read(), mapping


def renumber_citations_appearance(docx_bytes):
    """
    Renumbers citations by order of first appearance in the text (1, 2, 3...).
    Scans inline superscripts left-to-right, assigns new sequential numbers,
    updates all superscripts, and re-orders the bibliography to match.
    Returns (fixed_bytes, {old_num: new_num})
    """
    doc     = Document(io.BytesIO(docx_bytes))
    ref_pat = re.compile(r'^\s*(\d+)\.\s+(.+)')

    # Build appearance-order mapping by scanning superscripts
    seen    = {}
    for para in doc.paragraphs:
        for run in para.runs:
            rpr = run._r.find(qn('w:rPr'))
            if rpr is None: continue
            va  = rpr.find(qn('w:vertAlign'))
            if va is None or va.get(qn('w:val')) != 'superscript': continue
            for ps in re.split(r'[,;]', run.text.strip()):
                if ps.strip().isdigit():
                    num = int(ps.strip())
                    if num not in seen: seen[num] = len(seen) + 1

    if not seen:
        return docx_bytes, {}

    mapping = seen  # old_num -> new_num

    _apply_superscript_mapping(doc, mapping)

    # Re-order bibliography to match appearance order
    bib = [(pi, int(m.group(1)), para.text.strip())
           for pi, para in enumerate(doc.paragraphs)
           if (m := ref_pat.match(para.text.strip()))]

    if bib:
        # Sort bib paragraphs by their new number (appearance order)
        bib_by_new = sorted(bib, key=lambda x: mapping.get(x[1], x[1]))
        _renumber_bib_paras(doc, bib, mapping, bib_sorted=bib_by_new)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read(), mapping

# ─────────────────────────────────────────────────────────────────────────────
# APP 6 LOGIC — FIGURE INVENTORY
# ─────────────────────────────────────────────────────────────────────────────
def scan_figures(docx_bytes):
    doc=Document(io.BytesIO(docx_bytes)); items=[]
    cap_pats=[(re.compile(r'^\s*Fig(?:ure)?\.?\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Figure'),
              (re.compile(r'^\s*Table\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Table'),
              (re.compile(r'^\s*Box\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Box'),
              (re.compile(r'^\s*Plate\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Plate'),
              (re.compile(r'^\s*Video\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Video'),
              (re.compile(r'^\s*Appendix\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Appendix')]
    DN='http://schemas.openxmlformats.org/drawingml/2006/main'
    for pi,para in enumerate(doc.paragraphs):
        text=para.text.strip()
        has_img=para._p.find(f'.//{{{DN}}}blip') is not None
        style=para.style.name if para.style else ''
        itype=None; inum=None
        for pat,pt in cap_pats:
            m=pat.match(text)
            if m: itype=pt; inum=m.group(1); break
        if itype or has_img or 'caption' in style.lower():
            items.append(dict(para_idx=pi,type=itype or ('Image' if has_img else 'Caption'),
                              number=inum or '',caption=text[:250],has_image=has_img,style=style))
    return items

def cross_ref_excel(items, excel_bytes):
    import openpyxl
    try:
        wb=openpyxl.load_workbook(io.BytesIO(excel_bytes)); ws=wb.active
        rows=list(ws.iter_rows(values_only=True))
        if not rows: return [(i,'no_data','') for i in items]
        hdrs=[str(c or '').lower().strip() for c in rows[0]]
        nc=next((i for i,h in enumerate(hdrs) if any(w in h for w in ['num','#'])),None)
        lc=next((i for i,h in enumerate(hdrs) if any(w in h for w in ['name','caption','title','new','label'])),None)
        if lc is None: return [(i,'no_label_col','') for i in items]
        lookup={}
        for row in rows[1:]:
            if not row: continue
            num=str(row[nc] or '').strip() if nc is not None else ''
            label=str(row[lc] or '').strip()
            if num and label: lookup[num.lower()]=label
        results=[]
        for item in items:
            exp=lookup.get(item['number'].lower(),'')
            if not exp: status='not_in_excel'
            elif exp.lower() in item['caption'].lower(): status='match'
            else: status='mismatch'
            results.append((item,status,exp))
        return results
    except Exception as e: return [(i,f'error','') for i in items]

# ─────────────────────────────────────────────────────────────────────────────
# APP 7 LOGIC — PUBMED SEARCH
# ─────────────────────────────────────────────────────────────────────────────
def pubmed_search_full(query,date_from='',date_to='',journal_filter='',max_results=20):
    try:
        term=query
        if date_from and date_to: term+=f' AND {date_from}:{date_to}[dp]'
        elif date_from: term+=f' AND {date_from}:3000[dp]'
        if journal_filter: term+=f' AND "{journal_filter}"[ta]'
        r=requests.get(PUBMED_ESEARCH,params={'db':'pubmed','term':term,
            'retmax':max_results,'retmode':'json','sort':'relevance'},timeout=10)
        ids=r.json().get('esearchresult',{}).get('idlist',[])
        if not ids: return []
        r2=requests.get('https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi',
            params={'db':'pubmed','id':','.join(ids),'retmode':'xml','rettype':'abstract'},timeout=15)
        root=ET.fromstring(r2.content); results=[]
        for art in root.findall('.//PubmedArticle'):
            pmid=art.findtext('.//PMID') or ''
            title=art.findtext('.//ArticleTitle') or ''
            abstract=' '.join(t.text or '' for t in art.findall('.//AbstractText'))
            journal=art.findtext('.//Journal/Title') or art.findtext('.//ISOAbbreviation') or ''
            year=art.findtext('.//PubDate/Year') or (art.findtext('.//PubDate/MedlineDate') or '')[:4]
            volume=art.findtext('.//Volume') or ''; issue=art.findtext('.//Issue') or ''
            pages=art.findtext('.//MedlinePgn') or ''
            doi=next((a.text for a in art.findall('.//ArticleId') if a.get('IdType')=='doi'),'')
            pmc=next((a.text for a in art.findall('.//ArticleId') if a.get('IdType')=='pmc'),'')
            authors=[]
            for auth in art.findall('.//Author')[:6]:
                last=auth.findtext('LastName') or ''; fore=auth.findtext('Initials') or ''
                if last: authors.append(f"{last} {fore}".strip())
            results.append(dict(pmid=pmid,title=title,authors=authors,journal=journal,
                year=year,volume=volume,issue=issue,pages=pages,
                abstract=abstract[:600],doi=doi,
                pubmed_url=f'https://pubmed.ncbi.nlm.nih.gov/{pmid}/',
                pmc_url=f'https://www.ncbi.nlm.nih.gov/pmc/articles/{pmc}/' if pmc else '',
                doi_url=f'https://doi.org/{doi}' if doi else ''))
        return results
    except: return []

def results_to_xml(results):
    import html as hm
    xml='<?xml version="1.0" encoding="UTF-8"?>\n<xml>\n  <records>\n'
    for i,r in enumerate(results,1):
        xml+=f'    <record>\n      <rec-number>{i}</rec-number>\n      <ref-type name="Journal Article">17</ref-type>\n'
        if r['authors']:
            xml+='      <contributors>\n        <authors>\n'
            for a in r['authors']: xml+=f'          <author>{hm.escape(a)}</author>\n'
            xml+='        </authors>\n      </contributors>\n'
        xml+=f'      <titles>\n        <title>{hm.escape(r["title"])}</title>\n        <secondary-title>{hm.escape(r["journal"])}</secondary-title>\n      </titles>\n'
        xml+=f'      <dates><year>{hm.escape(r["year"])}</year></dates>\n'
        for k,t in [('volume','volume'),('issue','number'),('pages','pages')]:
            if r[k]: xml+=f'      <{t}>{hm.escape(r[k])}</{t}>\n'
        if r['abstract']: xml+=f'      <abstract>{hm.escape(r["abstract"])}</abstract>\n'
        if r['doi']: xml+=f'      <electronic-resource-number>{hm.escape(r["doi"])}</electronic-resource-number>\n'
        xml+=f'      <urls><related-urls><url>{hm.escape(r["pubmed_url"])}</url></related-urls></urls>\n'
        xml+='    </record>\n'
    xml+='  </records>\n</xml>\n'; return xml

# ─────────────────────────────────────────────────────────────────────────────
# APP 8 LOGIC — BATCH RENAME
# ─────────────────────────────────────────────────────────────────────────────
def load_rename_pairs(excel_bytes):
    import openpyxl
    try:
        wb=openpyxl.load_workbook(io.BytesIO(excel_bytes)); ws=wb.active
        rows=list(ws.iter_rows(values_only=True))
        if not rows: return []
        hdrs=[str(c or '').lower().strip() for c in rows[0]]
        oc=next((i for i,h in enumerate(hdrs) if any(w in h for w in ['old','current','find','original','from'])),0)
        nc=next((i for i,h in enumerate(hdrs) if any(w in h for w in ['new','replace','final','to','updated'])),1)
        pairs=[]
        for row in rows[1:]:
            if not row or len(row)<=max(oc,nc): continue
            old=str(row[oc] or '').strip(); new=str(row[nc] or '').strip()
            if old and new and old!=new: pairs.append((old,new))
        return pairs
    except: return []

def batch_rename(docx_bytes,pairs,match_case=False,whole_word=False):
    import html as hm
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        raw=z.read('word/document.xml').decode('utf-8')
        all_files={n:z.read(n) for n in z.namelist()}
    report=[]; fixed=raw
    for old,new in pairs:
        oe=hm.escape(old); ne=hm.escape(new)
        flags=0 if match_case else re.IGNORECASE
        pat=re.escape(oe)
        if whole_word: pat=r'(?<![a-zA-Z])'+pat+r'(?![a-zA-Z])'
        count=len(re.findall(pat,fixed,flags=flags))
        if count>0: fixed=re.sub(pat,ne,fixed,flags=flags); report.append({'old':old,'new':new,'count':count,'status':'replaced'})
        else: report.append({'old':old,'new':new,'count':0,'status':'not_found'})
    all_files['word/document.xml']=fixed.encode('utf-8')
    buf=io.BytesIO()
    with zipfile.ZipFile(buf,'w',zipfile.ZIP_DEFLATED) as zout:
        for n,d in all_files.items(): zout.writestr(n,d)
    buf.seek(0); return buf.read(),report

# ─────────────────────────────────────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────────────────────────────────────
defaults=dict(flagged=[],current_idx=0,decisions=[],doc_obj=None,repair_done=False,
              refs=[],vec=None,mat=None,
              fix_stage=1,fix_analysis=None,fix_raw_xml=None,fix_docx_bytes=None,
              fix_after_stage1=None,fix_after_stage2=None,fix_karol_db_id=None,
              fix_karol_rec_nums={},fix_missing_refs=[],fix_doc_name="document.docx",
              comp_result=None,comp_usage={},comp_labels=("",""),comp_refs=([],[]))
for k,v in defaults.items():
    if k not in st.session_state: st.session_state[k]=v

# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
