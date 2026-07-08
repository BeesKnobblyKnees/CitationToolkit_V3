"""
Generalized chapter PDF -> Word converter (tuned for the Tachdjian two-column layout).

convert(pdf_path, chapter_num) -> .docx bytes
  * maintains structure (headings, subheadings, body, Type I/II classification lists)
  * removes figures, tables, boxes, plates (and their captions/legends/footnotes)
  * highlights every in-text Figure/Table/Box/Plate/Video callout in BRIGHT CYAN
  * preserves in-text superscript reference numbers
  * emits a sectioned reference list from the chapter's online-only (.eN) pages

extract_pdf(pdf_path, start, end) -> .pdf bytes  (the chapter's pages as a standalone PDF)
"""
import fitz, re, io, subprocess
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import chapter_pages as cp

CALLOUT = re.compile(
    r'\b(?:e?Figs?\.?|e?Figures?|Tables?|Boxe?s?|Plates?|Videos?)\s*'
    r'\d+\.\d+[A-Za-z]?(?:[\u2013\u2014-][A-Za-z])?'
    r'(?:\s*(?:and|to|through|,|&|[\u2013\u2014-])\s*(?:\d+\.)?\d+[A-Za-z]?(?:[\u2013\u2014-][A-Za-z])?)*',
    re.I)
TYPE = re.compile(r'^\s*(Type\s+[IVXL]+)\b[:.\s]*(.*)$')
def clean(t): return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f\u00ad]', '', t or '')

def extract_pdf(pdf_path, start, end):
    doc = fitz.open(pdf_path); out = fitz.open()
    out.insert_pdf(doc, from_page=start-1, to_page=end-1)
    b = out.tobytes(); out.close(); doc.close(); return b

def _blocks(doc, pi):
    page = doc[pi-1]; W = page.rect.width; H = page.rect.height; mid = W/2
    out = []
    for b in page.get_text("dict")["blocks"]:
        if "lines" not in b: continue
        x0, y0, x1, y1 = b["bbox"]
        if y0 < 78 or y1 > H-45: continue                       # header/footer
        runs = []; sizes = []
        for li, l in enumerate(b["lines"]):
            if li > 0 and runs:
                lt, ls, lf = runs[-1]; rs = lt.rstrip()
                if "\u00ad" in rs[-2:]:
                    runs[-1] = (re.sub(r"[-\u00ad]+$", "", rs), ls, lf)   # soft-hyphen word break
                elif rs.endswith("-"):
                    runs[-1] = (rs, ls, lf)                                # compound hyphen at line end
                else:
                    runs.append((" ", 9.5, 0))                            # normal line break -> space
            for s in l["spans"]:
                if not s["text"]: continue
                runs.append((s["text"], round(s["size"], 1), s["flags"]))
                if s["text"].strip(): sizes.append(round(s["size"], 1))
        if not runs: continue
        dom = max(set(sizes), key=sizes.count) if sizes else 9.5
        bold = any((fl & 16) for _, _, fl in runs)
        fonts = set(s["font"] for l in b["lines"] for s in l["spans"] if s["text"].strip())
        L = []
        for l in b["lines"]:
            lt = "".join(s["text"] for s in l["spans"]).strip()
            if lt: L.append((round(l["bbox"][1]), round(l["bbox"][0]), lt))
        is_list = sum(1 for _, _, t in L if TYPE.match(t)) >= 2
        items = None
        if is_list:
            items = []; cur = None
            for y, x, t in sorted(L):
                m = TYPE.match(t)
                if m:
                    if cur: items.append(cur)
                    cur = [m.group(1), [m.group(2)] if m.group(2).strip() else []]
                else:
                    if cur is None: cur = ["", [t]]
                    else: cur[1].append(t)
            if cur: items.append(cur)
        out.append({"col": (-1 if (x1-x0) > 0.62*W else (0 if (x0+x1)/2 < mid else 1)),
                    "y": y0, "dom": dom, "bold": bold, "fonts": fonts,
                    "helv": bool(fonts) and all("Helvetica" in f for f in fonts),
                    "stone": any("StoneSans" in f for f in fonts),
                    "is_list": is_list, "items": items, "runs": runs})
    out.sort(key=lambda b: (b["col"], b["y"]))
    return out

def convert(pdf_path, chapter_num):
    r = cp.chapter_ranges(pdf_path, chapter_num)
    if not r or not r["text"]["pdf"]:
        raise ValueError("Chapter not found or no text pages")
    (bstart, bend) = r["text"]["pdf"]
    bib = r["bibliography"]["pdf"]
    title = r["title"]
    doc = fitz.open(pdf_path)

    out = Document()
    out.core_properties.title = f"Chapter {chapter_num} - {title}"
    for sec in out.sections:
        sec.page_width = Inches(8.5); sec.page_height = Inches(11)
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    nm = out.styles["Normal"]; nm.font.name = "Times New Roman"; nm.font.size = Pt(12)
    nm.paragraph_format.line_spacing = 2.0
    nm.paragraph_format.space_before = Pt(0); nm.paragraph_format.space_after = Pt(0)
    for hn in ("Heading 1", "Heading 2", "Heading 3", "Title"):
        h = out.styles[hn]; h.font.name = "Times New Roman"; h.font.bold = True
        h.font.size = Pt(14 if hn == "Title" else 12); h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.line_spacing = 2.0
        h.paragraph_format.space_before = Pt(12); h.paragraph_format.space_after = Pt(12)

    def shade(run):
        rpr = run._r.get_or_add_rPr(); hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), "cyan"); rpr.append(hl)
    def add_text(p, text, sup=False, bold=False):
        if sup:
            rr = p.add_run(text); rr.font.superscript = True
            if bold: rr.bold = True
            return
        pos = 0
        for m in CALLOUT.finditer(text):
            if m.start() > pos:
                rr = p.add_run(text[pos:m.start()])
                if bold: rr.bold = True
            rr = p.add_run(m.group(0)); shade(rr)
            if bold: rr.bold = True
            pos = m.end()
        if pos < len(text):
            rr = p.add_run(text[pos:])
            if bold: rr.bold = True
    def add_body(runs, footnotes=None):
        p = out.add_paragraph()
        for t, sz, fl in runs:
            if not t: continue
            sup = bool(fl & 1 or (sz and sz <= 7.2))
            if sup and footnotes and t.strip() in footnotes:
                rr = p.add_run(footnotes[t.strip()]); rr.font.superscript = True; continue
            add_text(p, clean(t), sup, bool(fl & 16))
        return p
    FN = re.compile(r'^\s*([A-Za-z\u2020\u2021\u00b6\u2016#\*])\s*[\u2009\s]*[Rr]eferences?\s+([\d,\s\u2013\u2014-]+)')

    out.add_heading(f"CHAPTER {chapter_num}", level=0)
    out.add_heading(title, level=1)
    def toc_like(t): return len(re.findall(r'\D\s\d{3,4}\b', t)) >= 2
    started = False; prev = None
    for pi in range(bstart, bend+1):
        blocks = _blocks(doc, pi)
        footnotes = {}
        for b in blocks:
            m = FN.match("".join(t for t, _, _ in b["runs"]))
            if b["dom"] < 9.5 and m:
                footnotes[m.group(1)] = re.sub(r'\s+', ' ', m.group(2)).strip().rstrip('.').strip()
        for b in blocks:
            txt = "".join(t for t, _, _ in b["runs"]).strip()
            if not txt: continue
            dom = b["dom"]
            if b["helv"]: continue                               # figure panel labels
            if dom < 9.5 and FN.match(txt): continue             # footnote block -> inlined at marker
            if b["is_list"] and b["items"] and len(b["items"]) >= 2:
                for label, descs in b["items"]:
                    desc = " ".join(d.strip() for d in descs if d.strip())
                    line = (label + ": " + desc) if (label and desc) else (label or desc)
                    if line.strip(): add_text(out.add_paragraph(style="List Bullet"), clean(line))
                prev = "list"; continue
            if dom < 9.0: continue                               # figures/tables/boxes/plates/footnotes
            if re.match(r'^(FIGURE|FIG\.?|TABLE|BOX|PLATE|APPENDIX)\b', txt, re.I): continue
            if dom >= 20: continue                               # big chapter title (added already)
            if dom >= 12.5:                                      # section heading
                ht = clean(txt)
                if ht in ("Chapter Contents",) or ht == title or toc_like(ht) or len(re.sub(r'[^A-Za-z]', '', ht)) < 3:
                    continue
                started = True
                out.add_heading(ht, level=2); prev = "h"; continue
            if not started:
                if toc_like(txt) or dom >= 12: continue
                started = True                                    # fallback: first real body starts the chapter
            if 9.9 <= dom < 12.5 and (b["bold"] or b["stone"]) and len(txt) < 90:
                out.add_heading(clean(txt), level=3); prev = "h"; continue
            starts_bold = next((bool(fl & 16) for t, sz, fl in b["runs"] if t.strip()), False)
            prev_txt = re.sub(r'[\d,\s]+$', '', out.paragraphs[-1].text) if out.paragraphs else ''
            if prev == "body" and out.paragraphs and not starts_bold and not re.search(r'[.!?:)]\s*$', prev_txt):
                add = out.paragraphs[-1]
                for t, sz, fl in b["runs"]:
                    if not t: continue
                    sup = bool(fl & 1 or (sz and sz <= 7.2))
                    if sup and footnotes and t.strip() in footnotes:
                        rr = add.add_run(footnotes[t.strip()]); rr.font.superscript = True; continue
                    pre = (" " if not add.text.endswith(" ") else "")
                    add_text(add, pre + clean(t), sup, bool(fl & 16))
            else:
                add_body(b["runs"], footnotes)
            prev = "body"
    doc.close()

    # ---- bibliography (online-only .eN pages) ----
    if bib:
        out.add_page_break(); out.add_heading("References", level=1)
        txt = subprocess.run(['pdftotext', '-f', str(bib[0]), '-l', str(bib[1]), pdf_path, '-'],
                             capture_output=True, text=True).stdout
        entry = None
        def flush(e):
            if e:
                p = out.add_paragraph(clean(e).strip())
                p.paragraph_format.left_indent = Pt(18); p.paragraph_format.first_line_indent = Pt(-18)
        def is_section(s):
            return (len(s) < 45 and not re.search(r'\d', s) and s[:1].isupper()
                    and len(s.split()) <= 6 and not s.endswith(('.', ',', ';'))
                    and not re.search(r'\b(and|for|with|the|of a)\b.*\b(fracture|injur)', s.lower()))
        for ln in txt.replace('\x07', '').splitlines():
            s = ln.strip()
            if not s or re.match(r'^(CHAPTER|SECTION)\b', s) or re.match(r'^\d{3,4}(\.e\d+)?$', s):
                continue
            m = re.match(r'^(\d{1,3}[a-z]?)\.\s+(.*)', s)
            if m:
                flush(entry); entry = f"{m.group(1)}. {m.group(2)}"
            elif entry is None and is_section(s):
                out.add_heading(clean(s), level=3)
            elif entry is not None:
                entry += " " + s
        flush(entry)

    buf = io.BytesIO(); out.save(buf); return buf.getvalue()
