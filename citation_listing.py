"""
citation_listing.py  -- extract in-text citations in order of appearance.

Reproduces the content of EndNote 21's "Edit & Manage Citations" dialog:
citation groups in document order, each group's formatted display text, and
the references it contains (Author, Year, #RecNum) with a per-reference Count
(total appearances across the whole document) and a Library label.

Three input paths:
  * extract_from_docx_fieldcodes(data)  -- Word .docx WITH live EndNote fields.
        Fully faithful: recovers RecNum/Author/Year from the embedded EndNote
        XML, including the split-blob "monster" citations.
  * extract_numbers_in_order(text)      -- plain text / bare-superscript docx / PDF.
        Degraded: recovers the citation NUMBERS in order only. No RecNums.
        Optionally map numbers -> references via a parsed numbered bibliography.
"""
import io, re, base64, html
from collections import Counter, OrderedDict

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# --------------------------------------------------------------------------- #
#  Field-code path (faithful)                                                  #
# --------------------------------------------------------------------------- #
def _field_blocks(xml):
    """Split a document.xml body into balanced field-code blocks (in order)."""
    runs = re.findall(r'<w:r\b(?:(?!</w:r>).)*?</w:r>', xml, re.DOTALL)
    i, blocks = 0, []
    while i < len(runs):
        if 'fldCharType="begin"' in runs[i]:
            depth, blk, j = 0, [], i
            while j < len(runs):
                depth += (runs[j].count('fldCharType="begin"')
                          - runs[j].count('fldCharType="end"'))
                blk.append(runs[j]); j += 1
                if depth <= 0:
                    break
            blocks.append(''.join(blk)); i = j
        else:
            i += 1
    return blocks


def _block_endnotes(blk):
    """Return the list of complete <EndNote>..</EndNote> strings in one block.

    Handles both storage formats and the split-base64 corruption:
      * escaped XML living directly in <w:instrText> runs;
      * base64 <w:fldData> blobs (EN.CITE / EN.CITE.DATA), including blobs whose
        bytes are split across several chunks by a raw cut mid-tag -- those are
        reassembled by BYTE concatenation (not base64-string concatenation).
    """
    out = []
    # (a) escaped instrText
    instr = html.unescape(''.join(
        re.findall(r'<w:instrText[^>]*>([\s\S]*?)</w:instrText>', blk)))
    out += re.findall(r'<EndNote>[\s\S]*?</EndNote>', instr)
    # (b) fldData -> bytes
    chunks = [''.join(f.split())
              for f in re.findall(r'<w:fldData[^>]*>([\s\S]+?)</w:fldData>', blk)]
    dec = []
    for c in chunks:
        try:
            dec.append(base64.b64decode(c + '=' * ((4 - len(c) % 4) % 4)))
        except Exception:
            dec.append(b'')
    txt = [d.decode('utf-8', 'replace') for d in dec]
    used = [False] * len(chunks)
    seen = set()
    for i, t in enumerate(txt):                      # complete single chunks
        if '<EndNote>' in t and '</EndNote>' in t:
            used[i] = True
            if t not in seen:
                seen.add(t); out.append(t)
    heads = [i for i in range(len(txt))
             if not used[i] and '<EndNote>' in txt[i] and '</EndNote>' not in txt[i]]
    tails = [i for i in range(len(txt))
             if not used[i] and '</EndNote>' in txt[i] and '<EndNote>' not in txt[i]]
    mids = [i for i in range(len(txt))
            if not used[i] and '<EndNote>' not in txt[i] and '</EndNote>' not in txt[i]]
    for h in heads:                                  # reassemble head+mids+tail
        if tails:
            combo = dec[h] + b''.join(dec[m] for m in mids) + dec[tails[0]]
            t = combo.decode('utf-8', 'replace')
            if '<EndNote>' in t and '</EndNote>' in t and t not in seen:
                seen.add(t); out.append(t)
    s2 = set()
    return [x for x in out if not (x in s2 or s2.add(x))]


def _parse_group(en):
    disp = re.search(r'<DisplayText>(.*?)</DisplayText>', en, re.DOTALL)
    dtxt = re.sub(r'<[^>]+>', '', html.unescape(disp.group(1))) if disp else ''
    refs = []
    for c in re.findall(r'<Cite[ >][\s\S]*?</Cite>', en):
        au = re.search(r'<Author>([^<]*)</Author>', c)
        yr = re.search(r'<Year>([^<]*)</Year>', c)
        rn = re.search(r'<RecNum>([^<]*)</RecNum>', c)
        ttl = re.search(r'<titles><title>([^<]*)</title>', c)
        refs.append({
            'author': html.unescape(au.group(1)) if au else '',
            'year': yr.group(1) if yr else '',
            'recnum': rn.group(1) if rn else '',
            'title': html.unescape(ttl.group(1)) if ttl else '',
        })
    return dtxt.strip(), refs


def extract_from_docx_fieldcodes(data):
    """data: bytes of a .docx. Returns dict(groups, counts, totals)."""
    xml = zipfile_read(data, 'word/document.xml')
    groups = []
    for blk in _field_blocks(xml):
        for en in _block_endnotes(blk):
            if '<Cite' not in en:
                continue
            disp, refs = _parse_group(en)
            if refs:
                groups.append({'display': disp, 'refs': refs})
    for i, g in enumerate(groups, 1):
        g['order'] = i
    counts = Counter(r['recnum'] for g in groups for r in g['refs'] if r['recnum'])
    totals = {
        'groups': len(groups),
        'citations': sum(len(g['refs']) for g in groups),
        'references': len(counts),
    }
    return {'groups': groups, 'counts': counts, 'totals': totals}


def has_endnote_fields(data):
    try:
        xml = zipfile_read(data, 'word/document.xml')
    except Exception:
        return False
    return 'ADDIN EN.CITE' in xml


def zipfile_read(data, name):
    import zipfile
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return z.read(name).decode('utf-8')


# --------------------------------------------------------------------------- #
#  Degraded path (PDF / plain text / bare-superscript docx)                    #
# --------------------------------------------------------------------------- #
_NUMRUN = re.compile(r'(?<!\d)(\d{1,3}(?:\s*[,\u2013\u2014-]\s*\d{1,3})*)')

def _expand(token):
    """'23-25' -> [23,24,25];  '191' -> [191]."""
    token = token.replace('\u2013', '-').replace('\u2014', '-')
    out = []
    for part in re.split(r'\s*,\s*', token):
        m = re.match(r'^\s*(\d+)\s*-\s*(\d+)\s*$', part)
        if m:
            a, b = int(m.group(1)), int(m.group(2))
            if 0 < b - a < 200:
                out += list(range(a, b + 1)); continue
        if part.strip().isdigit():
            out.append(int(part.strip()))
    return out


def extract_numbers_in_order(citation_spans):
    """citation_spans: list of raw display strings already isolated as citations
    (e.g. superscript runs).  Returns ordered groups of expanded numbers."""
    groups = []
    for i, span in enumerate(citation_spans, 1):
        nums = _expand(span)
        if nums:
            groups.append({'order': i, 'display': span.strip(), 'numbers': nums})
    counts = Counter(n for g in groups for n in g['numbers'])
    return {'groups': groups, 'counts': counts,
            'totals': {'groups': len(groups),
                       'citations': sum(len(g['numbers']) for g in groups),
                       'references': len(counts)}}


def parse_numbered_bibliography(text):
    """Parse a trailing 'References' list of the form '12. Author A. Title...'.
    Returns {num: 'first line of entry'}.  Best-effort."""
    refs = {}
    for m in re.finditer(r'(?m)^\s*(\d{1,3})[.)]\s+(.{8,200})', text):
        n = int(m.group(1))
        if n not in refs:
            refs[n] = re.sub(r'\s+', ' ', m.group(2)).strip()
    return refs


# --------------------------------------------------------------------------- #
#  Bare-superscript .docx path (no live fields, but explicit superscript runs) #
# --------------------------------------------------------------------------- #
def extract_superscripts_from_docx(data):
    """Recover citation numbers from a .docx whose citations are plain
    superscript runs (vertAlign='superscript' or an rStyle like 'citsup'/'sup').
    Returns the same shape as extract_numbers_in_order()."""
    xml = zipfile_read(data, 'word/document.xml')
    paras = re.findall(r'<w:p\b[\s\S]*?</w:p>', xml)
    spans = []
    for p in paras:
        for r in re.findall(r'<w:r\b[\s\S]*?</w:r>', p):
            rpr = re.search(r'<w:rPr>[\s\S]*?</w:rPr>', r)
            rpr = rpr.group(0) if rpr else ''
            is_sup = ('vertAlign w:val="superscript"' in rpr
                      or re.search(r'<w:rStyle w:val="(?:citsup|sup|Superscript|FootnoteReference)"', rpr))
            if not is_sup:
                continue
            txt = ''.join(re.findall(r'<w:t[^>]*>([\s\S]*?)</w:t>', r))
            txt = html.unescape(txt).strip()
            if txt and re.search(r'\d', txt):
                spans.append(txt)
    return extract_numbers_in_order(spans)


def extract_superscripts_from_pdf(data):
    """Best-effort: pull superscript numeric tokens from a PDF in reading order.
    Requires PyMuPDF (fitz). Superscript detection uses span flags + a smaller
    font size than the surrounding body text. Returns numbers-in-order shape,
    plus a parsed numbered bibliography if one is present at the end."""
    try:
        import fitz  # PyMuPDF
    except Exception:
        raise RuntimeError("PDF support needs PyMuPDF. Install with:  "
                           "pip install pymupdf --break-system-packages")
    doc = fitz.open(stream=data, filetype="pdf")
    spans_all, sizes = [], []
    for page in doc:
        d = page.get_text("dict")
        for blk in d.get("blocks", []):
            for line in blk.get("lines", []):
                for sp in line.get("spans", []):
                    sizes.append(sp["size"])
    body = (sorted(sizes)[len(sizes)//2] if sizes else 10.0)   # median = body size
    full_text = []
    for page in doc:
        full_text.append(page.get_text("text"))
        d = page.get_text("dict")
        for blk in d.get("blocks", []):
            for line in blk.get("lines", []):
                for sp in line.get("spans", []):
                    flags = sp.get("flags", 0)
                    is_super = bool(flags & 1) or sp["size"] <= body * 0.85
                    t = sp["text"].strip()
                    if is_super and t and re.fullmatch(r'[\d,\u2013\u2014\-\s]+', t):
                        spans_all.append(t)
    res = extract_numbers_in_order(spans_all)
    res['bibliography'] = parse_numbered_bibliography('\n'.join(full_text))
    return res


# --------------------------------------------------------------------------- #
#  Normalisers -> uniform group rows {label, count}                            #
# --------------------------------------------------------------------------- #
def normalize_fieldcode(result):
    counts = result['counts']
    groups = []
    for g in result['groups']:
        rows = [{'label': '{}, {} #{}'.format(r['author'] or '\u2014', r['year'], r['recnum']),
                 'count': counts.get(r['recnum'], '')} for r in g['refs']]
        groups.append({'order': g['order'], 'display': g['display'], 'rows': rows})
    return groups, result['totals']


def normalize_numbers(result, bibliography=None):
    counts = result['counts']
    bib = bibliography or {}
    groups = []
    for g in result['groups']:
        rows = []
        for n in g['numbers']:
            label = '#{}'.format(n)
            if n in bib:
                label = '{} \u2014 {}'.format(n, bib[n])
            rows.append({'label': label, 'count': counts.get(n, '')})
        groups.append({'order': g['order'], 'display': g['display'], 'rows': rows})
    return groups, result['totals']


# --------------------------------------------------------------------------- #
#  Document builder (uniform groups)                                           #
# --------------------------------------------------------------------------- #
def build_listing_docx(groups, totals, library_name="Library",
                       source_name="", mode_note="", title="In-Text Citations in Order of Appearance"):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ACCENT, GROUP_FILL = "8B1A1A", "F3ECEC"

    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), fill)
        tcPr.append(sh)

    def grid(table):
        b = OxmlElement('w:tblBorders')
        for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            x = OxmlElement('w:' + e)
            x.set(qn('w:val'), 'single'); x.set(qn('w:sz'), '4')
            x.set(qn('w:space'), '0'); x.set(qn('w:color'), 'CCCCCC'); b.append(x)
        tblPr = table._tbl.tblPr; anchor = None
        for c in tblPr:
            if c.tag in (qn('w:shd'), qn('w:tblLayout'), qn('w:tblLook')):
                anchor = c; break
        anchor.addprevious(b) if anchor is not None else tblPr.append(b)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.8); sec.top_margin = sec.bottom_margin = Inches(0.8)
    doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(10)

    h = doc.add_paragraph(); r = h.add_run(title)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor.from_string(ACCENT)
    if source_name:
        s = doc.add_paragraph(); sr = s.add_run('Source document:  ' + source_name)
        sr.font.size = Pt(9); sr.italic = True
    t = doc.add_paragraph()
    tr = t.add_run('Totals:  {groups} citation groups,  {citations} citations,  '
                   '{references} references     |     Library:  {lib}'.format(lib=library_name, **totals))
    tr.font.size = Pt(9); tr.bold = True
    if mode_note:
        m = doc.add_paragraph(); mr = m.add_run(mode_note)
        mr.font.size = Pt(8.5); mr.italic = True; mr.font.color.rgb = RGBColor.from_string("8B1A1A")
    doc.add_paragraph()

    widths = [Inches(0.5), Inches(4.7), Inches(0.7), Inches(1.6)]
    table = doc.add_table(rows=1, cols=4); table.alignment = WD_TABLE_ALIGNMENT.LEFT
    grid(table)
    hdr = table.rows[0]
    for i, txt in enumerate(['#', 'Citation', 'Count', 'Library']):
        c = hdr.cells[i]; run = c.paragraphs[0].add_run(txt)
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor.from_string("FFFFFF")
        shade(c, ACCENT)
    trPr = hdr._tr.get_or_add_trPr(); th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)

    for g in groups:
        gr = table.add_row()
        gr.cells[0].paragraphs[0].add_run(str(g['order'])).bold = True
        gr.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        merged = gr.cells[1].merge(gr.cells[2]).merge(gr.cells[3])
        disp = g['display'] if g['display'] else '(no display text)'
        run = merged.paragraphs[0].add_run(disp); run.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(ACCENT)
        shade(gr.cells[0], GROUP_FILL); shade(merged, GROUP_FILL)
        for row in g['rows']:
            rr = table.add_row()
            p = rr.cells[1].paragraphs[0]; p.paragraph_format.left_indent = Inches(0.15)
            p.add_run(row['label']).font.size = Pt(9.5)
            cp = rr.cells[2].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.add_run(str(row['count'])).font.size = Pt(9.5)
            rr.cells[3].paragraphs[0].add_run(library_name).font.size = Pt(8.5)

    table.autofit = False; table.allow_autofit = False
    for rw in table.rows:
        for i, w in enumerate(widths):
            rw.cells[i].width = w
    doc.add_paragraph()
    e = doc.add_paragraph()
    er = e.add_run('End of listing  -  {groups} citation groups,  {citations} citations,  '
                   '{references} references.'.format(**totals))
    er.font.size = Pt(9); er.bold = True

    buf = io.BytesIO(); doc.save(buf)
    return _fix_zoom_bytes(buf.getvalue())


def _fix_zoom_bytes(b):
    import zipfile
    zin = zipfile.ZipFile(io.BytesIO(b)); out = io.BytesIO()
    with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            d = zin.read(item.filename)
            if item.filename == 'word/settings.xml':
                d = re.sub(rb'<w:zoom[^>]*/>', b'<w:zoom w:percent="100"/>', d)
            zout.writestr(item, d)
    return out.getvalue()
