"""
Library & Reference Tools — Apps 1, 6, 11 as tabs.
"""
import streamlit as st, zipfile, re, base64, io, html
import pandas as _pd
from pathlib import Path
from shared import *

st.markdown(APP_CSS, unsafe_allow_html=True)
st.markdown('<div class="app-label">A practical group &nbsp; 2 of 5</div>', unsafe_allow_html=True)
st.markdown("## Library & *Reference Tools*")

st.markdown("""<style>
.stButton button { background:#8b1a1a !important; color:#faf7f2 !important; }
.stButton button p, .stButton button span, .stButton button div { color:#faf7f2 !important; font-weight:600 !important; }
.stButton button:hover { background:#6b1212 !important; color:#faf7f2 !important; }
[data-testid="stExpander"] details summary { text-indent:-9999px !important; display:flex !important; align-items:center !important; gap:8px !important; padding:0.75rem 1rem !important; background:var(--surface) !important; cursor:pointer !important; overflow:hidden !important; }
[data-testid="stExpander"] details summary * { text-indent:0 !important; font-size:0.92rem !important; color:#2e2416 !important; font-family:'Source Sans 3',sans-serif !important; font-weight:500 !important; }
[data-testid="stExpander"] details summary svg { width:16px !important; height:16px !important; }
</style>""", unsafe_allow_html=True)

st.markdown("Search PubMed, compare reference lists, and look up library RecNums.")
st.divider()

for _k, _v in [
    ("comp_result",   None),
    ("comp_labels",   None),
    ("comp_usage",    None),
    ("comp_refs",     None),
    ("rn_figures",    []),
    ("rn_scanned",    False),
    ("ck_result",     None),
]:
    if _k not in st.session_state:
        st.session_state[_k] = _v

_t1, _t2, _t3, _t4 = st.tabs(["01 · PubMed Search", "06 · Reference Comparator", "11 · RecNum Inspector", "12 · Citation Checker"])

with _t1:
        st.markdown("## PubMed Literature Search")
        st.markdown('<div class="instruction-box">Search PubMed for relevant articles. Results include abstracts and links to free full text where available (PubMed Central). Export results as EndNote XML to import directly into your EndNote library.</div>',unsafe_allow_html=True)
        query=st.text_input("Search query",placeholder="e.g. arthrogryposis clubfoot Ponseti treatment",help="Supports PubMed syntax: AND, OR, NOT, [MeSH], [ti], [au]")
        col1,col2,col3=st.columns(3)
        with col1: date_from=st.text_input("Year from",placeholder="2019")
        with col2: date_to=st.text_input("Year to",placeholder="2026")
        with col3: max_res=st.slider("Max results",5,50,20)
        journal_filter=st.text_input("Limit to journal (optional)",placeholder="e.g. J Pediatr Orthop")
        if st.button("Search PubMed",type="primary",disabled=not query.strip()):
            with st.spinner(f"Searching PubMed: {query}..."):
                results=pubmed_search_full(query.strip(),date_from.strip(),date_to.strip(),journal_filter.strip(),max_res)
            if not results: st.warning("No results found. Try broadening your search terms.")
            else:
                st.success(f"Found {len(results)} results.")
                xml_export=results_to_xml(results)
                st.download_button(f"⬇ Export all {len(results)} refs as EndNote XML",
                    data=xml_export.encode('utf-8'),file_name="pubmed_results.xml",mime="application/xml")
                st.divider()
                for i,r in enumerate(results,1):
                    authors_str='; '.join(r['authors'][:3])+(' et al.' if len(r['authors'])>3 else '')
                    cit=f"{authors_str} ({r['year']}). *{r['journal']}*"
                    if r['volume']: cit+=f" {r['volume']}"
                    if r['issue']:  cit+=f"({r['issue']})"
                    if r['pages']:  cit+=f":{r['pages']}"
                    with st.expander(f"**{i}.** {r['title'][:100]}{'...' if len(r['title'])>100 else ''}"):
                        st.markdown(cit)
                        lc=st.columns(3)
                        with lc[0]: st.markdown(f"[PubMed]({r['pubmed_url']})")
                        with lc[1]:
                            if r['pmc_url']: st.markdown(f"[Free full text (PMC)]({r['pmc_url']})")
                            elif r['doi_url']: st.markdown(f"[DOI]({r['doi_url']})")
                            else: st.markdown("*No free full text*")
                        with lc[2]:
                            if r['doi']: st.caption(f"DOI: {r['doi']}")
                        if r['abstract']: st.markdown("**Abstract:**"); st.markdown(r['abstract'])
                        single=results_to_xml([r])
                        st.download_button("⬇ Export this ref",data=single.encode('utf-8'),
                            file_name=f"ref_{r['pmid']}.xml",mime="application/xml",key=f"exp_{r['pmid']}")

    # ─────────────────────────────────────────────────────────────────────────────
    # APP 8 UI — BATCH RENAME

with _t2:
        st.markdown("## Reference List Comparator")
        col1,col2=st.columns(2)
        with col1:
            st.markdown("**List A**")
            file_a=st.file_uploader("List A",type=["xml","docx","txt"],key="comp_a",label_visibility="collapsed")
        with col2:
            st.markdown("**List B**")
            file_b=st.file_uploader("List B",type=["xml","docx","txt"],key="comp_b",label_visibility="collapsed")
        st.markdown("**Manuscript (optional)**")
        ms_file=st.file_uploader("Manuscript",type=["docx"],key="comp_ms",label_visibility="collapsed")
        if st.button("Compare Lists",type="primary",disabled=not(file_a and file_b)):
            with st.spinner("Comparing..."):
                refs_a,label_a=load_ref_file(file_a); refs_b,label_b=load_ref_file(file_b)
                if refs_a and refs_b:
                    all_c=[r["corpus"] for r in refs_a]+[r["corpus"] for r in refs_b]
                    vec=TfidfVectorizer(ngram_range=(1,2),sublinear_tf=True,max_features=50000)
                    vec.fit(all_c)
                    emb_a=vec.transform([r["corpus"] for r in refs_a])
                    emb_b=vec.transform([r["corpus"] for r in refs_b])
                    matrix=cosine_similarity(emb_a,emb_b)
                    matched,only_a,fuzzy,matched_b=[],[],[],set()
                    for i,ra in enumerate(refs_a):
                        bj=int(matrix[i].argmax()); bs=float(matrix[i][bj])
                        if bs>=MATCH_THRESHOLD: matched.append((ra,refs_b[bj],bs)); matched_b.add(bj)
                        elif bs>=FUZZY_THRESHOLD: fuzzy.append((ra,refs_b[bj],bs))
                        else: only_a.append(ra)
                    only_b=[refs_b[j] for j in range(len(refs_b)) if j not in matched_b]
                    result=dict(matched=matched,only_in_a=only_a,only_in_b=only_b,fuzzy=fuzzy)
                    usage={}
                    if ms_file:
                        ms_doc=Document(io.BytesIO(ms_file.read()))
                        ms_paras=[p.text for p in ms_doc.paragraphs if p.text.strip()]
                        for ref in only_a+only_b+[x[0] for x in fuzzy]:
                            words=[w for w in ref["title"].split() if len(w)>5][:5]
                            found=[p[:200] for p in ms_paras if sum(1 for w in words if w.lower() in p.lower())>=min(3,len(words))]
                            usage[ref.get("id","")]=found
                    st.session_state.update(dict(comp_result=result,comp_usage=usage,
                        comp_labels=(label_a,label_b),comp_refs=(refs_a,refs_b)))
        if st.session_state.comp_result:
            result=st.session_state.comp_result; usage=st.session_state.comp_usage
            label_a,label_b=st.session_state.comp_labels
            st.divider()
            col1,col2,col3,col4=st.columns(4)
            col1.metric("Matched",len(result["matched"])); col2.metric("Only in A",len(result["only_in_a"]))
            col3.metric("Only in B",len(result["only_in_b"])); col4.metric("Review",len(result["fuzzy"]))
            tabs=st.tabs([f"Only in A ({len(result['only_in_a'])})",f"Only in B ({len(result['only_in_b'])})",
                          f"Review ({len(result['fuzzy'])})",f"Matched ({len(result['matched'])})"])
            def ref_card(ref,color,locs=None):
                loc=""
                if locs: loc=f'<div style="font-size:0.75rem;color:#4fc3f7;margin-top:4px">Found in {len(locs)} paragraph(s)</div>'+"".join(f'<div style="font-size:0.74rem;color:#607080;font-style:italic">{l[:150]}...</div>' for l in locs[:2])
                st.markdown(f'<div class="ref-item {color}">{fmt_ref(ref)}{loc}</div>',unsafe_allow_html=True)
            with tabs[0]:
                if result["only_in_a"]:
                    st.caption(f"In **{label_a}** but not **{label_b}**")
                    for ref in result["only_in_a"]: ref_card(ref,"missing",usage.get(ref.get("id",""),[]))
                else: st.success("None")
            with tabs[1]:
                if result["only_in_b"]:
                    st.caption(f"In **{label_b}** but not **{label_a}**")
                    for ref in result["only_in_b"]: ref_card(ref,"warn",usage.get(ref.get("id",""),[]))
                else: st.success("None")
            with tabs[2]:
                for ra,rb,score in result["fuzzy"]:
                    st.markdown(f'<div class="ref-item warning"><b>[{score:.3f}]</b> A:{fmt_ref(ra,True)}<br>B:{fmt_ref(rb,True)}</div>',unsafe_allow_html=True)
            with tabs[3]:
                for ra,rb,score in result["matched"]:
                    st.markdown(f'<div class="ref-item ok">[{score:.3f}] {fmt_ref(ra,True)}</div>',unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────────────────────────────
    # APP 4 UI — DOCUMENT MERGER

with _t3:
        st.markdown('<div class="app-label">A practical tool &nbsp;11</div>', unsafe_allow_html=True)
        st.markdown("## RecNum Inspector")
        st.markdown('''<div class="instruction-box">
        Reads a Word document that is connected to your EndNote library and shows you every
        reference's RecNum (Record Number), author, year, and title — so you can identify
        the correct RecNums to use when fixing field codes.<br><br>
        <b>Also:</b> if you have field codes using the wrong RecNums (e.g. placeholder RecNums
        500-506 we assigned before knowing your library's actual numbers), paste a remapping
        table and this tool will fix all field codes in one step and download the corrected document.
        </div>''', unsafe_allow_html=True)

        rn_file = st.file_uploader(
            "Word document (.docx) — must be library-connected (have working EN.CITE fields)",
            type=["docx"], key="rn_doc")

        if rn_file:
            raw_bytes = rn_file.read()
            with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
                rn_xml = z.read('word/document.xml').decode('utf-8')

            # Extract all unique RecNum -> author/year/title from instrText
            import html as _html
            rn_data = {}
            for m in re.finditer(
                    r'ADDIN EN\.CITE (&lt;EndNote&gt;[\s\S]+?)(?=</w:instrText>)', rn_xml):
                decoded = _html.unescape(m.group(1))
                rns     = re.findall(r'<RecNum>(\d+)</RecNum>', decoded)
                authors = re.findall(r'<Author>([^<]+)</Author>', decoded)
                years   = re.findall(r'<Year>([^<]+)</Year>', decoded)
                titles  = re.findall(r'<title>([^<]+)</title>', decoded)
                db_ids  = re.findall(r'db-id="([^"]*)"', decoded)
                for i, rn in enumerate(rns):
                    if rn not in rn_data:
                        auth  = authors[i].split(',')[0] if i < len(authors) else (authors[0].split(',')[0] if authors else '?')
                        yr    = years[i] if i < len(years) else (years[0] if years else '?')
                        title = titles[0][:60] if titles else ''
                        db    = db_ids[i] if i < len(db_ids) else (db_ids[0] if db_ids else '')
                        rn_data[rn] = {'author': auth, 'year': yr, 'title': title, 'db_id': db}

            # Also check fldData for hidden refs
            import base64 as _b64
            for b64r in re.findall(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>', rn_xml):
                b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
                pad = (4-len(b64)%4)%4
                try:
                    dec = _b64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
                    if '<EndNote>' not in dec: continue
                    rns    = re.findall(r'<RecNum>(\d+)</RecNum>', dec)
                    auths  = re.findall(r'<Author>([^<]+)</Author>', dec)
                    yrs    = re.findall(r'<Year>([^<]+)</Year>', dec)
                    titles = re.findall(r'<title>([^<]+)</title>', dec)
                    db_ids = re.findall(r'db-id="([^"]*)"', dec)
                    for i, rn in enumerate(rns):
                        if rn not in rn_data:
                            auth  = auths[i].split(',')[0] if i < len(auths) else (auths[0].split(',')[0] if auths else '?')
                            yr    = yrs[i] if i < len(yrs) else (yrs[0] if yrs else '?')
                            title = titles[0][:60] if titles else ''
                            db    = db_ids[i] if i < len(db_ids) else (db_ids[0] if db_ids else '')
                            rn_data[rn] = {'author': auth, 'year': yr, 'title': title,
                                           'db_id': db, 'hidden': True}
                except: pass

            correct_db = '9veea52shxtee2e2wsbpwvf89wz55atsf52s'
            st.metric("Unique RecNums found", len(rn_data))

            # ── Search / filter ───────────────────────────────────────────────────
            search = st.text_input("Search by author, year, or title", placeholder="e.g. Guirguis or 2017")

            rows = []
            for rn, info in sorted(rn_data.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 9999):
                if search:
                    needle = search.lower()
                    haystack = f"{info['author']} {info['year']} {info['title']}".lower()
                    if needle not in haystack:
                        continue
                db_ok = info['db_id'] == correct_db if info['db_id'] else False
                rows.append({
                    'RecNum':  rn,
                    'Author':  info['author'],
                    'Year':    info['year'],
                    'Title':   info['title'][:55],
                    'DB-ID ✓': '✓' if db_ok else ('— empty' if not info['db_id'] else '✗ wrong'),
                    'Hidden':  '⚠ fldData only' if info.get('hidden') else '',
                })

            if rows:
                import pandas as _pd
                df = _pd.DataFrame(rows)
                st.dataframe(df, use_container_width=True, height=400,
                             column_config={
                                 'RecNum':  st.column_config.TextColumn('RecNum', width=80),
                                 'Author':  st.column_config.TextColumn('Author', width=120),
                                 'Year':    st.column_config.TextColumn('Year', width=60),
                                 'Title':   st.column_config.TextColumn('Title'),
                                 'DB-ID ✓': st.column_config.TextColumn('DB-ID', width=80),
                                 'Hidden':  st.column_config.TextColumn('Status', width=110),
                             })
                # Download as CSV
                csv_data = _pd.DataFrame([{
                    'RecNum': r['RecNum'], 'Author': r['Author'],
                    'Year': r['Year'], 'Title': r['Title']} for r in rows])
                st.download_button("⬇ Download RecNum list as CSV",
                                   data=csv_data.to_csv(index=False),
                                   file_name="recnum_list.csv", mime="text/csv")
            else:
                st.info("No matching references found.")

            # ── RecNum remapping ──────────────────────────────────────────────────
            st.divider()
            st.markdown("### Fix wrong RecNums in another document")
            st.markdown(
                "If a document uses placeholder RecNums (e.g. 500, 501…) that don't exist "
                "in your library, enter the remapping below — then upload the document to fix "
                "and download a corrected version."
            )
            st.markdown("**Remapping format** — one per line: `old_recnum -> new_recnum`  "
                        "e.g. `500 -> 79`")

            remap_text = st.text_area("RecNum remapping", height=200,
                                       placeholder="500 -> 79\n501 -> 112\n502 -> 116\n503 -> 413")

            fix_file = st.file_uploader(
                "Document to fix (.docx)", type=["docx"], key="rn_fix_doc")

            if remap_text.strip() and fix_file and st.button("Apply remapping", type="primary"):
                # Parse remap table
                remap = {}
                errors = []
                for line in remap_text.strip().splitlines():
                    line = line.strip()
                    if not line or line.startswith('#'): continue
                    parts = re.split(r'\s*->\s*', line)
                    if len(parts) != 2:
                        errors.append(f"Bad line: '{line}'")
                        continue
                    old_rn, new_rn = parts[0].strip(), parts[1].strip()
                    if not old_rn.isdigit() or not new_rn.isdigit():
                        errors.append(f"Non-numeric RecNum in: '{line}'")
                        continue
                    remap[old_rn] = new_rn

                if errors:
                    for e in errors: st.error(e)
                elif not remap:
                    st.warning("No valid remapping lines found.")
                else:
                    fix_bytes = fix_file.read()
                    with zipfile.ZipFile(io.BytesIO(fix_bytes)) as z:
                        fix_xml   = z.read('word/document.xml').decode('utf-8')
                        fix_files = {n: z.read(n) for n in z.namelist()}

                    applied = []
                    for old_rn, new_rn in remap.items():
                        before = fix_xml.count(f'&lt;RecNum&gt;{old_rn}&lt;/RecNum&gt;')
                        if before == 0:
                            st.warning(f"RecNum {old_rn} not found in document — skipped")
                            continue
                        fix_xml = fix_xml.replace(
                            f'&lt;RecNum&gt;{old_rn}&lt;/RecNum&gt;',
                            f'&lt;RecNum&gt;{new_rn}&lt;/RecNum&gt;')
                        fix_xml = fix_xml.replace(
                            f'&lt;rec-number&gt;{old_rn}&lt;/rec-number&gt;',
                            f'&lt;rec-number&gt;{new_rn}&lt;/rec-number&gt;')
                        fix_xml = re.sub(
                            rf'(&lt;key[^&]*&gt;){old_rn}(&lt;/key&gt;)',
                            rf'\g<1>{new_rn}\2', fix_xml)
                        after = fix_xml.count(f'&lt;RecNum&gt;{new_rn}&lt;/RecNum&gt;')
                        applied.append(f"{old_rn} → {new_rn} ({before} citation(s) updated)")

                    fix_files['word/document.xml'] = fix_xml.encode('utf-8')
                    out_buf = io.BytesIO()
                    with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
                        for n, d in fix_files.items(): zout.writestr(n, d)
                    out_buf.seek(0)
                    fixed_bytes = out_buf.read()

                    new_rns = set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', fix_xml))
                    st.success(f"✓ Applied {len(applied)} remapping(s). Document now has {len(new_rns)} unique RecNums.")
                    for msg in applied:
                        st.markdown(f"  • {msg}")
                    st.download_button(
                        "⬇ Download remapped document",
                        data=fixed_bytes,
                        file_name=Path(fix_file.name).stem + "_remapped.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )

with _t4:
        st.markdown('<div class="app-label">A practical tool &nbsp;12</div>', unsafe_allow_html=True)
        st.markdown("## Citation Checker & Reference Search")
        st.markdown('''<div class="instruction-box">
        <b>Three tools in one:</b>
        <ul style="margin:0.4rem 0 0 1rem">
          <li><b>Citation coverage</b> — checks which bibliography references are actually cited in the chapter body text, and which citations in the body have no matching bibliography entry</li>
          <li><b>Reference search</b> — find every place a specific reference is cited in the document</li>
          <li><b>Orphan finder</b> — shows bibliography entries with no in-text citation, and in-text citations with no bibliography entry</li>
        </ul>
        </div>''', unsafe_allow_html=True)

        _ck_doc = st.file_uploader("Word document (.docx)", type=["docx"], key="ck_doc")

        if _ck_doc:
            import re as _re
            from lxml import etree as _et
            from docx import Document as _CKDoc

            _ck_bytes = _ck_doc.read()
            _ck_doc.seek(0)

            # ── Parse document ────────────────────────────────────────────────────
            @st.cache_data(show_spinner="Reading document...")
            def _parse_ck_doc(_bytes):
                _doc = _CKDoc(io.BytesIO(_bytes))

                # Bibliography entries (numbered paragraphs)
                _ref_pat = _re.compile(r'^\s*(\d+[a-z]?)[\.)]\s+(.+)')
                _bib = {}
                for _p in _doc.paragraphs:
                    _m = _ref_pat.match(_p.text.strip())
                    if _m: _bib[_m.group(1)] = _p.text.strip()

                # In-text citations from EndNote field codes
                with zipfile.ZipFile(io.BytesIO(_bytes)) as _z:
                    _xml = _z.read('word/document.xml').decode('utf-8')

                _field_cited = set(_re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', _xml))

                # In-text citations from superscript runs
                _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                _root = _et.fromstring(_xml.encode('utf-8'))
                _super_cited = set()
                for _r in _root.iter(f'{{{_W}}}r'):
                    _rpr = _r.find(f'{{{_W}}}rPr')
                    if _rpr is None: continue
                    _va = _rpr.find(f'{{{_W}}}vertAlign')
                    if _va is None or _va.get(f'{{{_W}}}val') != 'superscript': continue
                    _t = _r.find(f'{{{_W}}}t')
                    _txt = (_t.text or '') if _t is not None else ''
                    for _part in _re.split(r'[,;\s]+', _txt):
                        _p2 = _part.strip().rstrip('.')
                        if _p2.isdigit(): _super_cited.add(_p2)

                # Paragraph text with their citation numbers (for search)
                _para_cites = []
                for _pi, _p in enumerate(_doc.paragraphs):
                    if not _p.text.strip(): continue
                    _nums = []
                    for _r in _p._p.iter(f'{{{_W}}}r'):
                        _rpr = _r.find(f'{{{_W}}}rPr')
                        if _rpr is None: continue
                        _va  = _rpr.find(f'{{{_W}}}vertAlign')
                        if _va is None or _va.get(f'{{{_W}}}val') != 'superscript': continue
                        _t2 = _r.find(f'{{{_W}}}t')
                        _tt = (_t2.text or '') if _t2 is not None else ''
                        for _pt in _re.split(r'[,;\s]+', _tt):
                            _pp = _pt.strip().rstrip('.')
                            if _pp.isdigit(): _nums.append(_pp)
                    if _nums:
                        _para_cites.append((_pi, _p.text.strip()[:200], _nums))

                return _bib, _super_cited, _field_cited, _para_cites

            _bib, _super_cited, _field_cited, _para_cites = _parse_ck_doc(_ck_bytes)
            _all_cited = _super_cited | {n for n in _field_cited if n.isdigit()}
            _bib_nums  = set(_bib.keys())

            st.caption(f"Found {len(_bib)} bibliography entries · {len(_all_cited)} unique in-text citations")
            st.divider()

            _ck_tab1, _ck_tab2, _ck_tab3 = st.tabs([
                "Citation Coverage",
                "Reference Search",
                "Orphan Finder"
            ])

            # ── Tab 1: Citation Coverage ──────────────────────────────────────────
            with _ck_tab1:
                st.markdown("### Which bibliography entries are cited in the text?")

                _not_cited_in_text = sorted(
                    [n for n in _bib_nums if n not in _all_cited],
                    key=lambda x: int(x) if x.isdigit() else 999
                )
                _cited_no_bib = sorted(
                    [n for n in _all_cited if n not in _bib_nums],
                    key=lambda x: int(x) if x.isdigit() else 999
                )
                _cited_with_bib = sorted(
                    [n for n in _bib_nums if n in _all_cited],
                    key=lambda x: int(x) if x.isdigit() else 999
                )

                _cv1, _cv2, _cv3 = st.columns(3)
                _cv1.metric("Cited in text ✓",       len(_cited_with_bib),  delta=None)
                _cv2.metric("In bib but NOT cited",   len(_not_cited_in_text),
                            delta=f"-{len(_not_cited_in_text)}" if _not_cited_in_text else None,
                            delta_color="inverse" if _not_cited_in_text else "off")
                _cv3.metric("Cited but NOT in bib",   len(_cited_no_bib),
                            delta=f"-{len(_cited_no_bib)}" if _cited_no_bib else None,
                            delta_color="inverse" if _cited_no_bib else "off")

                if _not_cited_in_text:
                    with st.expander(f"In bibliography but NOT cited in text ({len(_not_cited_in_text)})",
                                      expanded=True):
                        st.caption("These references appear in the bibliography but no matching "
                                   "superscript citation was found in the document body.")
                        for _n in _not_cited_in_text:
                            st.markdown(f'<div class="ref-item warning">#{_n} — {_bib[_n][:100]}</div>',
                                        unsafe_allow_html=True)
                else:
                    st.success("✓ All bibliography entries are cited in the text.")

                if _cited_no_bib:
                    with st.expander(f"Cited in text but NOT in bibliography ({len(_cited_no_bib)})",
                                      expanded=True):
                        st.caption("These citation numbers appear as superscripts in the text "
                                   "but have no matching bibliography entry.")
                        for _n in _cited_no_bib:
                            # Find paragraphs where this number is cited
                            _locs = [_txt[:120] for _pi, _txt, _nums in _para_cites if _n in _nums]
                            st.markdown(f'<div class="ref-item error">#{_n} — cited in '
                                        f'{len(_locs)} location(s)</div>', unsafe_allow_html=True)
                            for _l in _locs[:2]:
                                st.markdown(f'<div style="font-size:0.8rem;color:var(--ink-dim);'
                                            f'margin:2px 0 4px 12px;font-style:italic">'
                                            f'…{_l}…</div>', unsafe_allow_html=True)
                else:
                    st.success("✓ All in-text citations have a bibliography entry.")

            # ── Tab 2: Reference Search ───────────────────────────────────────────
            with _ck_tab2:
                st.markdown("### Find where a reference is cited in the document")

                _search_col1, _search_col2 = st.columns([2, 1])
                with _search_col1:
                    _search_query = st.text_input(
                        "Search by reference number or author/keyword",
                        placeholder="e.g.  42  or  Smith  or  fracture fixation",
                        key="ck_search"
                    )
                with _search_col2:
                    _search_mode = st.selectbox("Search type",
                        ["By reference number", "By author/keyword in bibliography"],
                        key="ck_search_mode")

                if _search_query.strip():
                    _q = _search_query.strip()

                    if _search_mode == "By reference number":
                        # Find all paragraphs citing this number
                        _hits = [(pi, txt, nums) for pi, txt, nums in _para_cites if _q in nums]
                        if _hits:
                            st.success(f"Reference #{_q} is cited in {len(_hits)} location(s):")
                            # Show bib entry
                            if _q in _bib:
                                st.markdown(f'<div class="ref-item ok">'
                                            f'<b>Bibliography entry #{_q}:</b> {_bib[_q]}</div>',
                                            unsafe_allow_html=True)
                            for _idx, (_pi, _txt, _nums) in enumerate(_hits):
                                _all_nums_str = ", ".join(f"#{n}" for n in sorted(set(_nums), key=lambda x: int(x) if x.isdigit() else 999))
                                st.markdown(
                                    f'<div class="card" style="margin-bottom:8px;">'
                                    f'<div style="font-size:0.7rem;color:var(--ink-faint);'
                                    f'font-family:monospace;margin-bottom:4px">'
                                    f'Paragraph {_pi+1} · cites {_all_nums_str}</div>'
                                    f'<div style="font-size:0.88rem;color:var(--ink-mid);'
                                    f'line-height:1.6">{_txt[:300]}'
                                    f'{"…" if len(_txt) > 300 else ""}</div>'
                                    f'</div>',
                                    unsafe_allow_html=True)
                        else:
                            # Check if it's in the bib at all
                            if _q in _bib:
                                st.warning(f"Reference #{_q} is in the bibliography but not cited in the text.")
                                st.markdown(f'<div class="ref-item warning">{_bib[_q]}</div>',
                                            unsafe_allow_html=True)
                            else:
                                st.info(f"Reference #{_q} not found in text citations or bibliography.")

                    else:  # By author/keyword in bibliography
                        _q_lower = _q.lower()
                        _bib_hits = [(n, txt) for n, txt in _bib.items() if _q_lower in txt.lower()]
                        if _bib_hits:
                            st.info(f"Found {len(_bib_hits)} bibliography entr{'y' if len(_bib_hits)==1 else 'ies'} matching '{_q}':")
                            for _n, _txt in _bib_hits:
                                _is_cited = _n in _all_cited
                                _color = "ok" if _is_cited else "warning"
                                _status = "✓ cited in text" if _is_cited else "⚠ not cited in text"
                                st.markdown(f'<div class="ref-item {_color}">'
                                            f'<b>#{_n}</b> {_status}<br>{_txt[:120]}</div>',
                                            unsafe_allow_html=True)
                                if _is_cited:
                                    _locs = [(pi, txt) for pi, txt, nums in _para_cites if _n in nums]
                                    if _locs:
                                        with st.expander(f"Show {len(_locs)} citation location(s) for #{_n}"):
                                            for _pi, _ptxt in _locs[:5]:
                                                st.markdown(
                                                    f'<div style="font-size:0.82rem;color:var(--ink-dim);'
                                                    f'margin:4px 0;padding:6px 10px;background:var(--surface);'
                                                    f'border-radius:4px;border-left:2px solid var(--border);">'
                                                    f'Para {_pi+1}: {_ptxt[:250]}…</div>',
                                                    unsafe_allow_html=True)
                        else:
                            st.info(f"No bibliography entries found matching '{_q}'.")

            # ── Tab 3: Orphan Finder ──────────────────────────────────────────────
            with _ck_tab3:
                st.markdown("### Orphan references")
                st.caption("Orphans are mismatches between the bibliography and in-text citations.")

                _or1, _or2 = st.columns(2)
                with _or1:
                    st.markdown("**In bibliography — not cited in text**")
                    if _not_cited_in_text:
                        for _n in _not_cited_in_text:
                            st.markdown(f'<div class="ref-item warning">#{_n} — {_bib[_n][:80]}</div>',
                                        unsafe_allow_html=True)
                    else:
                        st.success("None")

                with _or2:
                    st.markdown("**Cited in text — not in bibliography**")
                    if _cited_no_bib:
                        for _n in _cited_no_bib:
                            _ct = sum(1 for _pi,_txt,_nums in _para_cites if _n in _nums)
                            st.markdown(f'<div class="ref-item error">#{_n} — cited {_ct}x, '
                                        f'no bibliography entry</div>', unsafe_allow_html=True)
                    else:
                        st.success("None")

                if not _not_cited_in_text and not _cited_no_bib:
                    st.success("✓ No orphan references found. All citations match the bibliography.")

