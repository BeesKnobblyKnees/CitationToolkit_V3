"""
context_transplant.py  --  fix a section's citations by matching them to a
"linked twin" of the same passage, by the WORDING around each citation rather
than by any number or bibliography.

Edit-tolerant matching: each citation is compared to the source citations using
the distinctive CONTENT WORDS in a window on both sides of it (rare words count
for more, common words for less), so reworded sentences still match. An
order-aware alignment then uses the sequence of citations to place ones whose
wording was edited, and allows gaps for citations present in only one version.

TARGET = the working section (broken citations: EndNote INVALID errors, bare
superscript numbers, [[REF n]] markers). SOURCE = the linked twin with correct
field codes. Only broken target citations are transplanted by default.
"""
import io, re, math, zipfile, html, difflib, base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ORANGE = "FFC000"; RED = "FF5B5B"; GREENBG = "D9EAD3"
_RUN_RE = re.compile(r"<w:r\b(?:(?!</w:r>).)*?</w:r>", re.DOTALL)
_STOP = set((
    "the a an of to in and or for with on at by from as is are was were be been being this that "
    "these those it its their his her our your not but also into within between during than then "
    "so such can may which who whom whose when where while each both all any more most other some "
    "have has had not no nor only own same too very will just because about above after again "
    "against before below down further here once over under up out off").split())


def _txt(run):
    return html.unescape("".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", run, re.DOTALL)))


def _content_words(text):
    return [w for w in re.findall(r"[a-z]{3,}", text.lower()) if w not in _STOP]


def _is_sup_num(run):
    rpr = re.search(r"<w:rPr>.*?</w:rPr>", run, re.DOTALL)
    rpr = rpr.group(0) if rpr else ""
    sup = ('vertAlign w:val="superscript"' in rpr
           or re.search(r'<w:rStyle w:val="(?:citsup|sup|Superscript|FootnoteReference)"', rpr))
    return bool(sup) and bool(re.fullmatch(r"\s*\d{1,4}\s*", _txt(run)))


def r2c(run):
    return run.count('fldCharType="begin"') - run.count('fldCharType="end"')


def _block_info(block):
    instr = html.unescape("".join(re.findall(r"<w:instrText[^>]*>(.*?)</w:instrText>", block, re.DOTALL)))
    invalid = "INVALID CITATION" in instr
    rec = re.search(r"<RecNum>(\d+)</RecNum>", instr)
    recnum = rec.group(1) if rec else None
    disp = "".join(_txt(r) for r in re.findall(r"<w:r\b.*?</w:r>", block, re.DOTALL)
                   if "fldChar" not in r and "instrText" not in r)
    return disp.strip(), recnum, invalid


def _surname(author):
    m = re.match(r"^([A-Za-z\-\u00C0-\u017F]+)", (author or "").strip())
    return m.group(1).lower() if m else ""


def _toks(s):
    return set(re.findall(r"[a-z]{4,}", (s or "").lower()))


def _block_identity(block):
    """Author surname / year / title tokens / journal tokens from a field code's
    embedded EndNote record. Works even on INVALID codes (the record is intact)."""
    fd = re.search(r"<w:fldData[^>]*>([\s\S]+?)</w:fldData>", block)
    if not fd:
        return None
    b64 = "".join(fd.group(1).split()); pad = (4 - len(b64) % 4) % 4
    try:
        dec = base64.b64decode(b64 + "=" * pad).decode("utf-8", "replace")
    except Exception:
        return None
    au = re.search(r"<Author>([^<]+)</Author>", dec)
    yr = re.search(r"<Year>(\d{4})</Year>", dec)
    ti = re.search(r"<title>([^<]*)</title>", dec)
    jo = re.search(r"<secondary-title>([^<]*)</secondary-title>", dec)
    if not au:
        return None
    return {"sur": _surname(au.group(1)), "year": yr.group(1) if yr else "",
            "ttoks": _toks(html.unescape(ti.group(1)) if ti else ""),
            "jtoks": _toks(html.unescape(jo.group(1)) if jo else "")}


def _id_bonus(ti, si):
    """Identity affinity between a target and a source citation, plus the title
    overlap count (used to tell same-author-same-year siblings apart)."""
    if not ti or not si or not ti["sur"] or not si["sur"]:
        return 0.0, 0
    if ti["sur"] == si["sur"] and ti["year"] and ti["year"] == si["year"]:
        ov = len(ti["ttoks"] & si["ttoks"])
        return 0.10 + 0.05 * min(ov, 6), ov
    if ti["sur"] == si["sur"]:
        return 0.03, 0
    return 0.0, 0


def tokenize(xml):
    """Ordered tokens: ('text', run, text) or ('cite', kind, span_xml, display, broken)."""
    runs = _RUN_RE.findall(xml)
    toks, i = [], 0
    while i < len(runs):
        r = runs[i]
        if 'fldCharType="begin"' in r:
            depth, blk, j = 0, [], i
            while j < len(runs):
                depth += r2c(runs[j]); blk.append(runs[j]); j += 1
                if depth <= 0:
                    break
            block = "".join(blk)
            disp, recnum, invalid = _block_info(block)
            broken = invalid or recnum in (None, "0")
            toks.append(("cite", "field", block, disp, broken)); i = j
        elif _is_sup_num(r):
            grp, nums, j = [r], [_txt(r).strip()], i + 1
            while j < len(runs):
                if _is_sup_num(runs[j]):
                    grp.append(runs[j]); nums.append(_txt(runs[j]).strip()); j += 1
                elif _txt(runs[j]).strip() in {",", ";"}:
                    grp.append(runs[j]); j += 1
                else:
                    break
            toks.append(("cite", "bare", "".join(grp), ", ".join(n for n in nums if n), True)); i = j
        elif "[[REF" in _txt(r):
            toks.append(("cite", "refmarker", r, _txt(r).strip(), True)); i += 1
        else:
            toks.append(("text", r, _txt(r))); i += 1
    return toks


def _windows(toks, w_pre=20, w_post=12):
    """For each token, content-word lists for the prose before/after it."""
    n = len(toks); pre = [None] * n; post = [None] * n
    acc = []
    for k in range(n):
        if toks[k][0] == "cite":
            pre[k] = acc[-w_pre:]
        if toks[k][0] == "text":
            acc = acc + _content_words(toks[k][2])
    acc = []
    for k in range(n - 1, -1, -1):
        if toks[k][0] == "cite":
            post[k] = acc[:w_post]
        if toks[k][0] == "text":
            acc = _content_words(toks[k][2]) + acc
    return pre, post


def _bag(pre, post):
    return set((pre or [])) | set((post or []))


def _wjacc(a, b, idf):
    if not a or not b:
        return 0.0
    inter = a & b
    union = a | b
    si = sum(idf.get(t, 1.0) for t in inter)
    su = sum(idf.get(t, 1.0) for t in union)
    return si / su if su else 0.0


def _align(sim, floor):
    """Order-preserving max-weight alignment with gaps. Returns [(ti, sj)]."""
    n = len(sim); m = len(sim[0]) if n else 0
    if n == 0 or m == 0:
        return []
    dp = [[0.0] * (m + 1) for _ in range(n + 1)]
    move = [[None] * (m + 1) for _ in range(n + 1)]
    for i in range(n - 1, -1, -1):
        for j in range(m - 1, -1, -1):
            best = dp[i + 1][j]; mv = "skipT"
            if dp[i][j + 1] > best:
                best = dp[i][j + 1]; mv = "skipS"
            if sim[i][j] >= floor and sim[i][j] + dp[i + 1][j + 1] > best:
                best = sim[i][j] + dp[i + 1][j + 1]; mv = "match"
            dp[i][j] = best; move[i][j] = mv
    i = j = 0; pairs = []
    while i < n and j < m:
        mv = move[i][j]
        if mv == "match":
            pairs.append((i, j)); i += 1; j += 1
        elif mv == "skipT":
            i += 1
        else:
            j += 1
    return pairs


def transplant(target_bytes, source_bytes, threshold=0.4, replace_all=False, fill_by_order=False):
    src_xml = zipfile.ZipFile(io.BytesIO(source_bytes)).read("word/document.xml").decode("utf-8")
    s_toks = tokenize(src_xml)
    s_pre, s_post = _windows(s_toks)
    sources = []  # (bag, span, display, recnum)
    for k, t in enumerate(s_toks):
        if t[0] == "cite" and t[1] == "field" and not t[4]:
            _, recnum, _ = _block_info(t[2])
            sources.append((_bag(s_pre[k], s_post[k]), t[2], t[3], recnum))

    tz = zipfile.ZipFile(io.BytesIO(target_bytes))
    txml = tz.read("word/document.xml").decode("utf-8")
    t_toks = tokenize(txml)
    t_pre, t_post = _windows(t_toks)
    targets = []  # (tok_index, kind, span, display, broken, bag)
    for k, t in enumerate(t_toks):
        if t[0] == "cite":
            targets.append((k, t[1], t[2], t[3], t[4], _bag(t_pre[k], t_post[k])))

    # idf over every citation window in both documents
    all_bags = [s[0] for s in sources] + [tg[5] for tg in targets]
    N = len(all_bags) or 1
    df = {}
    for bag in all_bags:
        for w in bag:
            df[w] = df.get(w, 0) + 1
    idf = {w: math.log((N + 1) / (c + 0.5)) + 1.0 for w, c in df.items()}

    # similarity matrix (targets x sources) and best-per-target
    sim = [[_wjacc(tg[5], s[0], idf) for s in sources] for tg in targets]

    # identity affinity: pin same-author-same-year citations to the right source
    # by title, even when the surrounding wording is nearly identical
    s_ids = [_block_identity(s[1]) for s in sources]
    t_ids = [_block_identity(tg[2]) if tg[1] == "field" else None for tg in targets]
    ovl = [[0] * len(sources) for _ in targets]
    for i in range(len(targets)):
        if not t_ids[i]:
            continue
        for j in range(len(sources)):
            b, ov = _id_bonus(t_ids[i], s_ids[j])
            sim[i][j] += b; ovl[i][j] = ov

    best = []
    for i in range(len(targets)):
        bj, br = -1, 0.0
        for j in range(len(sources)):
            if sim[i][j] > br:
                br, bj = sim[i][j], j
        best.append((bj, br))

    pairs = _align(sim, floor=0.12)
    matched = {i: (j, sim[i][j]) for i, j in pairs}

    elig_idx = [i for i, tg in enumerate(targets) if (tg[4] or (replace_all and tg[1] == "field"))]

    # order-free fallback for eligible targets the alignment missed
    used = {j for _, (j, _) in matched.items()}
    for i in elig_idx:
        cur = matched.get(i)
        if cur and cur[1] >= threshold:
            continue
        bj, br = -1, 0.0
        for j in range(len(sources)):
            if j in used:
                continue
            if sim[i][j] > br:
                br, bj = sim[i][j], j
        if bj >= 0 and br >= threshold and (not cur or br > cur[1]):
            matched[i] = (bj, br); used.add(bj)

    accepted = {i: v for i, v in matched.items() if v[1] >= threshold}
    method = {i: "context" for i in accepted}
    used = {j for (j, _) in accepted.values()}

    # opt-in: fill leftover eligible targets by their ORDER between confident anchors
    if fill_by_order:
        anchors = sorted(accepted.items())
        bounds = [(-1, -1)] + [(ti, sj) for ti, (sj, _) in anchors] + [(len(targets), len(sources))]
        for (t_lo, s_lo), (t_hi, s_hi) in zip(bounds, bounds[1:]):
            tg_gap = [i for i in elig_idx if t_lo < i < t_hi and i not in accepted]
            s_gap = [j for j in range(s_lo + 1, s_hi) if j not in used]
            for ti, sj in zip(tg_gap, s_gap):
                accepted[ti] = (sj, sim[ti][sj]); used.add(sj); method[ti] = "position"

    # build replacements + report
    replace_at = {}; report = []
    for i, tg in enumerate(targets):
        tok_i, kind, span, disp, broken, _ = tg
        if not (broken or (replace_all and kind == "field")):
            continue
        ctx_words = t_pre[tok_i] or []
        ctx = " ".join(ctx_words[-9:])
        chosen = accepted.get(i)
        if chosen:
            j, r = chosen
            # ambiguity check: chosen source shares (surname,year) with another
            # source and the title didn't clearly separate them
            ambiguous = False
            if t_ids[i] and s_ids[j] and s_ids[j]["sur"]:
                sibs = [jj for jj in range(len(sources))
                        if s_ids[jj] and s_ids[jj]["sur"] == s_ids[j]["sur"]
                        and s_ids[jj]["year"] == s_ids[j]["year"]]
                if len(sibs) > 1:
                    second_ov = max((ovl[i][jj] for jj in sibs if jj != j), default=0)
                    if ovl[i][j] - second_ov < 2:
                        ambiguous = True
            replace_at[tok_i] = sources[j][1]
            report.append({"kind": kind, "target": disp or "(blank)", "context": ctx,
                           "matched": sources[j][2], "recnum": sources[j][3],
                           "ratio": round(r, 2),
                           "method": method.get(i, "context"),
                           "status": "ambiguous" if ambiguous else "transplanted"})
        else:
            report.append({"kind": kind, "target": disp or "(blank)", "context": ctx,
                           "matched": None, "recnum": None, "ratio": round(best[i][1], 2),
                           "method": None, "status": "no match"})

    out_parts = []
    for k, t in enumerate(t_toks):
        if k in replace_at:
            out_parts.append(replace_at[k])
        else:
            out_parts.append(t[1] if t[0] == "text" else t[2])
    new_doc = _reassemble(txml, t_toks, out_parts)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for it in tz.infolist():
            z.writestr(it, new_doc.encode("utf-8") if it.filename == "word/document.xml"
                       else tz.read(it.filename))
    return _fix_zoom(buf.getvalue()), report


def _reassemble(xml, toks, out_parts):
    spans_old = [t[1] if t[0] == "text" else t[2] for t in toks]
    result, pos = [], 0
    for old, new in zip(spans_old, out_parts):
        idx = xml.find(old, pos)
        if idx == -1:
            continue
        result.append(xml[pos:idx]); result.append(new); pos = idx + len(old)
    result.append(xml[pos:])
    return "".join(result)


def _fix_zoom(b):
    zin = zipfile.ZipFile(io.BytesIO(b)); out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for it in zin.infolist():
            d = zin.read(it.filename)
            if it.filename == "word/settings.xml":
                d = re.sub(rb"<w:zoom[^>]*/>", b'<w:zoom w:percent="100"/>', d)
            zout.writestr(it, d)
    return out.getvalue()


def summarize(report):
    inserted = [r for r in report if r["status"] in ("transplanted", "ambiguous")]
    return {"citations": len(report),
            "transplanted": len(inserted),
            "by_context": sum(1 for r in inserted if r.get("method") == "context"),
            "by_position": sum(1 for r in inserted if r.get("method") == "position"),
            "ambiguous": sum(1 for r in report if r["status"] == "ambiguous"),
            "unmatched": sum(1 for r in report if r["status"] == "no match")}


def build_report_docx(report, title="Context Transplant Report"):
    ACCENT = "8B1A1A"
    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), fill); tcPr.append(sh)
    def grid(table):
        b = OxmlElement("w:tblBorders")
        for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
            x = OxmlElement("w:" + e); x.set(qn("w:val"), "single"); x.set(qn("w:sz"), "4")
            x.set(qn("w:space"), "0"); x.set(qn("w:color"), "CCCCCC"); b.append(x)
        tblPr = table._tbl.tblPr; anchor = None
        for c in tblPr:
            if c.tag in (qn("w:shd"), qn("w:tblLayout"), qn("w:tblLook")):
                anchor = c; break
        anchor.addprevious(b) if anchor is not None else tblPr.append(b)

    doc = Document()
    s = doc.sections[0]; s.left_margin = s.right_margin = Inches(0.8); s.top_margin = s.bottom_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    h = doc.add_paragraph(); r = h.add_run(title)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor.from_string(ACCENT)
    sm = summarize(report)
    sp = doc.add_paragraph(); sr = sp.add_run(
        "%d broken citation(s): %d transplanted from the source, %d with no confident match."
        % (sm["citations"], sm["transplanted"], sm["unmatched"]))
    sr.font.size = Pt(9); sr.bold = True
    note = doc.add_paragraph(); nr = note.add_run(
        "Each broken citation was matched to the source by the distinctive words around it "
        "(reworded sentences still match). Transplanted = the source's field code was inserted; "
        "open in Word and Update Citations to renumber. No match = left in place (red).")
    nr.font.size = Pt(8.5); nr.italic = True; nr.font.color.rgb = RGBColor.from_string("666666")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4); grid(table)
    for i, t in enumerate(["Context (run-up words)", "Was", "Now", "Match"]):
        c = table.rows[0].cells[i]; run = c.paragraphs[0].add_run(t)
        run.bold = True; run.font.size = Pt(9.5); run.font.color.rgb = RGBColor.from_string("FFFFFF"); shade(c, ACCENT)
    trPr = table.rows[0]._tr.get_or_add_trPr(); th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)
    for row in report:
        tr = table.add_row()
        tr.cells[0].paragraphs[0].add_run("\u2026" + row["context"]).font.size = Pt(8.5)
        tr.cells[1].paragraphs[0].add_run(row["target"]).font.size = Pt(8.5)
        if row["status"] in ("transplanted", "ambiguous"):
            tr.cells[2].paragraphs[0].add_run("%s  #%s" % (row["matched"], row["recnum"])).font.size = Pt(8.5)
            by_pos = row.get("method") == "position"
            amb = row["status"] == "ambiguous"
            tcPr = tr.cells[2]._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
            sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto")
            sh.set(qn("w:fill"), ORANGE if (by_pos or amb) else GREENBG); tcPr.append(sh)
            if amb:
                label = "same author/year - verify  (%.2f)" % row["ratio"]
            elif by_pos:
                label = "by position  (%.2f) - verify" % row["ratio"]
            else:
                label = "transplanted  (%.2f)" % row["ratio"]
            tr.cells[3].paragraphs[0].add_run(label).font.size = Pt(8.5)
        else:
            tr.cells[2].paragraphs[0].add_run("(left as-is)").font.size = Pt(8.5)
            tcPr = tr.cells[2]._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
            sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), RED); tcPr.append(sh)
            tr.cells[3].paragraphs[0].add_run("no match  (best %.2f)" % row["ratio"]).font.size = Pt(8.5)
    table.autofit = False; table.allow_autofit = False
    for rw in table.rows:
        w = [Inches(3.0), Inches(0.9), Inches(1.9), Inches(1.2)]
        for i, x in enumerate(w):
            rw.cells[i].width = x
    buf = io.BytesIO(); doc.save(buf)
    return _fix_zoom(buf.getvalue())
