"""
Bibliography Relink Module
Relinks a text-edited draft whose citations are bare superscript numbers, using
(a) the draft's own numbered bibliography to learn what each number is, and
(b) a citation-intact "old" document as the source of real EndNote field codes.

Matching is by REFERENCE IDENTITY (first author + year), which is immune to the
renumbering that happens when either document is edited. Each in-text citation
group is resolved to its references; if a whole old field code cites exactly that
set it is transplanted intact, otherwise a field code is assembled from the old
document's real per-reference records (correct nested EN.CITE/EN.CITE.DATA form).
Anything unresolved becomes a yellow [[REF n]] placeholder for manual insertion.

Two citation populations are handled:
  * BODY  - bare superscript numbers (rStyle "citsup"), comma-separated by "sup".
  * FOOTNOTES - "References 1, 86, 119..." lists where the numbers carry rStyle
    "crossref" and are separated by plain ", " runs. Only footnotes whose text
    contains "Reference" are touched, so figure/table cross-refs are never altered.
"""

import re
import io
import zipfile
import base64
import html
import docx
from collections import Counter


def _docxml(b):
    with zipfile.ZipFile(io.BytesIO(b)) as z:
        return z.read("word/document.xml").decode("utf-8")


def _surname(author):
    m = re.match(r'^([A-Za-z\-\u00C0-\u017F]+)', author.strip())
    return m.group(1).lower() if m else None


def _format_ref(rest):
    """Build a disambiguating display: author+title snippet, then the journal
    and year (which sit after the title in Vancouver style and would otherwise
    be truncated away). Year and journal are appended in brackets so they are
    always visible even when the title is long."""
    rest = rest.strip()
    snippet = rest[:80].rstrip()
    if len(rest) > 80:
        snippet += "\u2026"
    journal, year = "", ""
    yrs = list(re.finditer(r'\b(?:18|19|20)\d{2}\b', rest))
    if yrs:
        ym = yrs[-1]
        year = ym.group(0)
        pre = rest[:ym.start()].rstrip().rstrip(".").strip()
        segs = re.split(r'\.\s+', pre)          # journal = last sentence-segment before the year
        if segs:
            cand = segs[-1].strip(" .,")
            if 0 < len(cand) <= 60 and re.search(r'[A-Za-z]', cand):
                journal = cand
    if journal and year:
        return f"{snippet}  [{journal}, {year}]"
    if year:
        return f"{snippet}  [{year}]"
    return snippet


def parse_bibliography(draft_bytes):
    """Map in-text/bibliography number -> (surname, year) and -> display text.
    Display text includes journal + year for disambiguation."""
    d = docx.Document(io.BytesIO(draft_bytes))
    bib, bibtext = {}, {}
    for p in d.paragraphs:
        m = re.match(r'^(\d{1,3})\s+(.*)', p.text.strip())
        if not m:
            continue
        num, rest = int(m.group(1)), m.group(2)
        if len(rest) < 15:
            continue
        yr = re.findall(r'\b(18|19|20)(\d{2})\b', rest)
        am = re.match(
            r'^((?:(?:van|von|der|den|de|del|della|da|di|du|dos|das|la|le|el|'
            r'ten|ter|af|zu|zur|zum|of)\s+)*'
            r'[A-Z\u00C0-\u017F][A-Za-z\u00C0-\u017F\'\-]*)', rest)
        bibtext[num] = _format_ref(rest)
        if yr and am:
            bib[num] = (am.group(1).lower(), yr[-1][0] + yr[-1][1])
    return bib, bibtext


def _parse_bib_full(b):
    """num -> full untruncated reference text (for content disambiguation)."""
    full = {}
    try:
        d = docx.Document(io.BytesIO(b))
    except Exception:
        return full
    for p in d.paragraphs:
        m = re.match(r'^(\d{1,3})\s+(.*)', (p.text or '').strip())
        if m and len(m.group(2)) >= 15:
            full[int(m.group(1))] = m.group(2)
    return full


def _toks(s):
    return set(re.findall(r'[a-z]{4,}', (s or '').lower()))


def _cite_content(cite_xml):
    """Pull title / journal / pages / volume / recnum out of an EN.CITE record
    so two same-author-same-year cites can be told apart."""
    def g(tag):
        m = re.search(r'<%s>([^<]*)</%s>' % (tag, tag), cite_xml)
        return html.unescape(m.group(1)) if m else ''
    return {'xml': cite_xml, 'recnum': g('RecNum'),
            'title': g('title'), 'jour': g('secondary-title') or g('full-title'),
            'pages': re.findall(r'\d+', g('pages')), 'vol': g('volume').strip()}


def _score_cite(cd, hint):
    htoks = _toks(hint); hnums = set(re.findall(r'\d+', hint or ''))
    t = len(_toks(cd['title']) & htoks)
    j = len(_toks(cd['jour']) & htoks)
    vp = (1 if cd['vol'] and cd['vol'] in hnums else 0) + sum(1 for p in cd['pages'] if p in hnums)
    return t * 1.0 + j * 1.0 + min(vp, 2) * 0.5


def _resolve_cite(key, hint, cites_by_id):
    """Pick the right old cite for a citation number. Returns (cite_dict, kind).
    kind 'exact' = confident, 'ambiguous' = indistinguishable (caller flags)."""
    lst = cites_by_id.get(key, [])
    if not lst:
        return None, 'none'
    if len(lst) == 1:
        return lst[0], 'exact'
    scored = sorted(((_score_cite(cd, hint), i, cd) for i, cd in enumerate(lst)),
                    key=lambda x: x[0], reverse=True)
    top, second = scored[0], scored[1]
    if top[0] >= 2.0 and (top[0] - second[0]) >= 1.5:
        return top[2], 'exact'
    if top[0] >= 1.0 and top[0] > second[0]:
        return top[2], 'near'
    # indistinguishable -> still apply the best guess, but mark it for review
    return top[2], 'ambiguous'


def index_old_fieldcodes(old_bytes):
    """Return (fieldxml_by_set, cites_by_id) from the citation-intact old doc.
    cites_by_id maps (surname,year) -> list of cite-dicts (one per distinct
    RecNum), so same-author-same-year papers are kept separate."""
    xml = _docxml(old_bytes)
    runs = re.findall(r"<w:r\b(?:(?!</w:r>).)*?</w:r>", xml, re.DOTALL)
    fieldxml_by_set, cites_by_id = {}, {}
    i = 0
    while i < len(runs):
        if 'fldCharType="begin"' in runs[i]:
            depth, blk, j = 0, [], i
            while j < len(runs):
                depth += runs[j].count('fldCharType="begin"') - runs[j].count('fldCharType="end"')
                blk.append(runs[j]); j += 1
                if depth <= 0:
                    break
            block = "".join(blk)
            fd = re.search(r'<w:fldData[^>]*>([\s\S]+?)</w:fldData>', block)
            if fd:
                b64 = "".join(fd.group(1).split()); pad = (4 - len(b64) % 4) % 4
                try:
                    dec = base64.b64decode(b64 + "=" * pad).decode("utf-8", "replace")
                    ids = set()
                    for c in re.findall(r"<Cite>.*?</Cite>", dec, re.DOTALL):
                        au = re.search(r"<Author>([^<]+)</Author>", c)
                        yr = re.search(r"<Year>(\d{4})</Year>", c)
                        if au and yr:
                            s = _surname(au.group(1))
                            if s:
                                key = (s, yr.group(1))
                                cd = _cite_content(c)
                                lst = cites_by_id.setdefault(key, [])
                                if not any(e['recnum'] == cd['recnum'] and cd['recnum'] for e in lst) \
                                        and not any(e['xml'] == cd['xml'] for e in lst):
                                    lst.append(cd)
                                ids.add(key)
                    if ids:
                        fieldxml_by_set.setdefault(frozenset(ids), block)
                except Exception:
                    pass
            i = j
        else:
            i += 1
    return fieldxml_by_set, cites_by_id


def _strip_disp(c):
    return re.sub(r"<DisplayText>.*?</DisplayText>", "", c, flags=re.DOTALL)


def _build_fieldcode(nums, bib, cites_by_id, bibfull, rstyle="citsup", superscript=True):
    chosen = []; ambiguous = []
    for n in nums:
        key = bib.get(n)
        if key is None:
            return None, []
        cd, kind = _resolve_cite(key, bibfull.get(n, ''), cites_by_id)
        if cd is None:          # truly missing -> placeholder
            return None, []
        if kind == 'ambiguous':
            ambiguous.append(n)
        chosen.append(cd)
    rpr = '<w:rPr><w:rStyle w:val="%s"/></w:rPr>' % rstyle
    face = "superscript" if superscript else "normal"
    cites = []
    for k, cd in enumerate(chosen):
        c = cd['xml']
        if k == 0:
            disp = ('<DisplayText><style face="%s">' % face
                    + ", ".join(str(n) for n in nums) + "</style></DisplayText>")
            c = re.sub(r"(</RecNum>)", r"\1" + disp, _strip_disp(c), count=1)
        else:
            c = _strip_disp(c)
        cites.append(c)
    en = "<EndNote>" + "".join(cites) + "</EndNote>"
    b64 = base64.b64encode(en.encode("utf-8")).decode("ascii")
    show = ", ".join(str(n) for n in nums)
    return ((f'<w:r>{rpr}<w:fldChar w:fldCharType="begin"><w:fldData xml:space="preserve">{b64}</w:fldData></w:fldChar></w:r>'
            f'<w:r>{rpr}<w:instrText xml:space="preserve"> ADDIN EN.CITE </w:instrText></w:r>'
            f'<w:r>{rpr}<w:fldChar w:fldCharType="begin"><w:fldData xml:space="preserve">{b64}</w:fldData></w:fldChar></w:r>'
            f'<w:r>{rpr}<w:instrText xml:space="preserve"> ADDIN EN.CITE.DATA </w:instrText></w:r>'
            f'<w:r>{rpr}</w:r>'
            f'<w:r>{rpr}<w:fldChar w:fldCharType="end"/></w:r>'
            f'<w:r>{rpr}<w:fldChar w:fldCharType="separate"/></w:r>'
            f'<w:r>{rpr}<w:t xml:space="preserve">{show}</w:t></w:r>'
            f'<w:r>{rpr}<w:fldChar w:fldCharType="end"/></w:r>'), ambiguous)


def _placeholder(nums, rstyle="citsup"):
    return ('<w:r><w:rPr><w:rStyle w:val="%s"/><w:highlight w:val="yellow"/></w:rPr>' % rstyle
            + '<w:t xml:space="preserve">[[REF %s]]</w:t></w:r>' % ",".join(str(n) for n in nums))


def _scan_replace(xml, bib, bibtext, fieldxml_by_set, cites_by_id, construct,
                  cite_style, sep_require_style, build_rstyle, build_superscript,
                  bibfull=None):
    """Find consecutive cite-styled number runs, replace each group with a real
    field code (exact transplant / assembled) or a placeholder. Returns
    (new_xml, kinds_counter, placeholders, all_nums)."""
    bibfull = bibfull or {}
    ambig = {k for k, v in cites_by_id.items() if len(v) > 1}
    runs = re.findall(r"<w:r\b(?:(?!</w:r>).)*?</w:r>", xml, re.DOTALL)
    reps, placeholders, all_nums, verify = [], [], [], []
    cur, cur_runs, last_num = [], [], -1

    def flush():
        nonlocal cur, cur_runs, last_num
        if cur:
            span = "".join(cur_runs[:last_num + 1])          # drop trailing separators
            idents = [bib.get(n) for n in cur]
            fs = frozenset(i for i in idents if i)
            # exact whole-block reuse only when no identity is ambiguous
            if all(idents) and fs in fieldxml_by_set and not (fs & ambig):
                reps.append((span, fieldxml_by_set[fs], "exact"))
            elif construct and all(idents) and all(i in cites_by_id for i in idents):
                fc, amb_nums = _build_fieldcode(cur, bib, cites_by_id, bibfull, build_rstyle, build_superscript)
                if fc:
                    reps.append((span, fc, "constructed"))
                    if amb_nums:
                        verify.append((amb_nums[:], " | ".join(bibtext.get(n, "?") for n in amb_nums)))
                else:
                    reps.append((span, _placeholder(cur, build_rstyle), "placeholder"))
                    placeholders.append((cur[:], " | ".join(bibtext.get(n, "?") for n in cur)))
            else:
                reps.append((span, _placeholder(cur, build_rstyle), "placeholder"))
                placeholders.append((cur[:], " | ".join(bibtext.get(n, "?") for n in cur)))
        cur, cur_runs, last_num = [], [], -1

    cite_tok = 'w:val="%s"' % cite_style
    for r in runs:
        tm = re.search(r"<w:t[^>]*>(.*?)</w:t>", r, re.DOTALL)
        txt = html.unescape(tm.group(1)) if tm else ""
        if cite_tok in r and txt.strip().isdigit():
            cur.append(int(txt.strip())); cur_runs.append(r)
            last_num = len(cur_runs) - 1; all_nums.append(int(txt.strip()))
        elif txt.strip() in {",", ";"} and cur and (not sep_require_style or 'w:val="sup"' in r):
            cur_runs.append(r)
        else:
            flush()
    flush()

    out = xml
    for span, repl, _ in reps:
        out = out.replace(span, repl, 1)
    return out, Counter(k for _, _, k in reps), placeholders, all_nums, verify


def _relink_reference_footnotes(foot_xml, bib, bibtext, fieldxml_by_set, cites_by_id, construct, bibfull=None):
    """Relink only footnotes whose text mentions 'Reference', leaving all other
    footnotes (figure/table cross-refs, editorial notes) untouched."""
    out = foot_xml
    kinds, phs, nums, verify = Counter(), [], [], []
    for m in re.finditer(r"<w:footnote\b[^>]*>.*?</w:footnote>", foot_xml, re.DOTALL):
        block = m.group(0)
        vis = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", block))
        if "Reference" not in vis:
            continue
        nb, k, p, n, v = _scan_replace(block, bib, bibtext, fieldxml_by_set, cites_by_id,
                                       construct, cite_style="crossref", sep_require_style=False,
                                       build_rstyle="crossref", build_superscript=False, bibfull=bibfull)
        if nb != block:
            out = out.replace(block, nb, 1)
            kinds += k
            phs += [(ns, "(footnote) " + tx) for ns, tx in p]
            nums += n
            verify += [(ns, "(footnote) " + tx) for ns, tx in v]
    return out, kinds, phs, nums, verify


def relink(draft_bytes, old_bytes, construct=True, bib_source_bytes=None):
    """Return (fixed_bytes, report, placeholders[list of (numbers, ref_text)]).

    bib_source_bytes: optional .docx whose numbered reference list supplies the
    numbering. Use when the draft is a section that doesn't contain its own
    bibliography - point it at the chapter's master reference list."""
    bib, bibtext = parse_bibliography(bib_source_bytes or draft_bytes)
    bibfull = _parse_bib_full(bib_source_bytes or draft_bytes)
    fieldxml_by_set, cites_by_id = index_old_fieldcodes(old_bytes)

    src = zipfile.ZipFile(io.BytesIO(draft_bytes))
    names = set(src.namelist())

    body_xml = src.read("word/document.xml").decode("utf-8")
    body_out, body_kinds, body_ph, body_nums, body_verify = _scan_replace(
        body_xml, bib, bibtext, fieldxml_by_set, cites_by_id, construct,
        cite_style="citsup", sep_require_style=True,
        build_rstyle="citsup", build_superscript=True, bibfull=bibfull)

    foot_out, foot_kinds, foot_ph, foot_nums, foot_verify = None, Counter(), [], [], []
    if "word/footnotes.xml" in names:
        foot_xml = src.read("word/footnotes.xml").decode("utf-8")
        foot_out, foot_kinds, foot_ph, foot_nums, foot_verify = _relink_reference_footnotes(
            foot_xml, bib, bibtext, fieldxml_by_set, cites_by_id, construct, bibfull=bibfull)
        if foot_out == foot_xml:
            foot_out = None  # nothing changed; keep original part verbatim

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for item in src.infolist():
            if item.filename == "word/document.xml":
                z.writestr(item, body_out.encode("utf-8"))
            elif item.filename == "word/footnotes.xml" and foot_out is not None:
                z.writestr(item, foot_out.encode("utf-8"))
            else:
                z.writestr(item, src.read(item.filename))

    cited = set(body_nums) | set(foot_nums)
    verify = body_verify + foot_verify
    report = {
        "locations": sum(body_kinds.values()) + sum(foot_kinds.values()),
        "exact": body_kinds.get("exact", 0) + foot_kinds.get("exact", 0),
        "constructed": body_kinds.get("constructed", 0) + foot_kinds.get("constructed", 0),
        "placeholders": body_kinds.get("placeholder", 0) + foot_kinds.get("placeholder", 0),
        "verify": verify,                       # constructed best-guess (same author/year) to double-check
        "verify_count": sum(len(nz) for nz, _ in verify),
        "unique_refs_cited": len(cited),
        "bibliography_entries": len(bibtext),
        "orphan_bib_entries": len(set(bibtext) - cited),
        # body / footnote split (page reads these with .get, safe if absent)
        "body_locations": sum(body_kinds.values()),
        "footnote_locations": sum(foot_kinds.values()),
        "footnote_exact": foot_kinds.get("exact", 0),
        "footnote_constructed": foot_kinds.get("constructed", 0),
        "footnote_placeholders": foot_kinds.get("placeholder", 0),
    }
    return buf.getvalue(), report, body_ph + foot_ph


if __name__ == "__main__":
    draft = open("lld.docx", "rb").read()
    old = open("lld_linked.docx", "rb").read()
    fixed, report, phs = relink(draft, old)
    print("report:", report)
    print("placeholders:", len(phs), "| footnote placeholders:",
          sum(1 for _, t in phs if t.startswith("(footnote)")))
    open("/tmp/_test_relink.docx", "wb").write(fixed)
    x = zipfile.ZipFile(io.BytesIO(fixed)).read("word/document.xml").decode()
    f = zipfile.ZipFile(io.BytesIO(fixed)).read("word/footnotes.xml").decode()
    print("body begins/ends:", x.count('fldCharType="begin"'), x.count('fldCharType="end"'))
    print("foot begins/ends:", f.count('fldCharType="begin"'), f.count('fldCharType="end"'))


def _parse_bib_full(b):
    """num -> full untruncated reference text, for the bibliography snapshot
    appended to the placeholder list."""
    full = {}
    try:
        d = docx.Document(io.BytesIO(b))
    except Exception:
        return full
    for p in d.paragraphs:
        m = re.match(r'^(\d{1,3})\s+(.*)', (p.text or '').strip())
        if m and len(m.group(2)) >= 15:
            full[int(m.group(1))] = m.group(2)
    return full


def build_placeholders_docx(placeholders, title="Citation Placeholders to Insert", bib_snapshot=None):
    """Render the placeholder list as a readable Word table:
    one row per placeholder, the [[REF #]] bold, each reference on its own line.
    placeholders: list of (nums_list, joined_text) where joined_text is
    ' | '-joined reference displays aligned 1:1 with nums_list."""
    import io as _io
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ACCENT, FILL = "8B1A1A", "F3ECEC"

    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), fill)
        tcPr.append(sh)

    def grid(table):
        b = OxmlElement('w:tblBorders')
        for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            x = OxmlElement('w:' + e)
            x.set(qn('w:val'), 'single'); x.set(qn('w:sz'), '4')
            x.set(qn('w:space'), '0'); x.set(qn('w:color'), 'CCCCCC'); b.append(x)
        tblPr = table._tbl.tblPr; anchor = None
        for c in tblPr:
            if c.tag in (qn('w:shd'), qn('w:tblLayout'), qn('w:tblLook')):
                anchor = c; break
        anchor.addprevious(b) if anchor is not None else tblPr.append(b)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.8); sec.top_margin = sec.bottom_margin = Inches(0.8)
    doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(10)

    h = doc.add_paragraph(); r = h.add_run(title)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor.from_string(ACCENT)
    note = doc.add_paragraph()
    nr = note.add_run("Each row is one citation location that could not be auto-relinked. "
                      "Open the relinked .docx in Word, Update Citations and Bibliography, "
                      "then insert the reference(s) listed for each [[REF #]] marker. "
                      "Journal and year are shown to disambiguate similar authors/titles.")
    nr.font.size = Pt(8.5); nr.italic = True; nr.font.color.rgb = RGBColor.from_string("666666")
    cnt = doc.add_paragraph(); cr = cnt.add_run("%d placeholder location(s)." % len(placeholders))
    cr.font.size = Pt(9); cr.bold = True
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2); grid(table)
    hdr = table.rows[0]
    for i, txt in enumerate(["Placeholder", "Reference(s) to insert at this location"]):
        c = hdr.cells[i]; run = c.paragraphs[0].add_run(txt)
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor.from_string("FFFFFF")
        shade(c, ACCENT)
    trPr = hdr._tr.get_or_add_trPr(); th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)

    for nums, text in placeholders:
        parts = text.split(" | ")
        row = table.add_row()
        # col 1: bold [[REF ...]]
        p = row.cells[0].paragraphs[0]
        br = p.add_run("[[REF %s]]" % ", ".join(str(n) for n in nums))
        br.bold = True; br.font.size = Pt(10); br.font.color.rgb = RGBColor.from_string(ACCENT)
        shade(row.cells[0], FILL)
        # col 2: one reference per line, prefixed with its number when aligned
        cell = row.cells[1]
        cell.paragraphs[0].text = ""
        for i, part in enumerate(parts):
            para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            num = nums[i] if i < len(nums) else "?"
            nrun = para.add_run("%s  " % num); nrun.bold = True; nrun.font.size = Pt(9.5)
            nrun.font.color.rgb = RGBColor.from_string(ACCENT)
            trun = para.add_run(part); trun.font.size = Pt(9.5)

    widths = [Inches(1.6), Inches(5.4)]
    table.autofit = False; table.allow_autofit = False
    for rw in table.rows:
        for i, w in enumerate(widths):
            rw.cells[i].width = w

    # ── bibliography snapshot ─────────────────────────────────────────────── #
    # The numbered reference list relinked against, kept in the same list so the
    # [[REF #]] numbers stay interpretable if the bibliography is later renumbered.
    if bib_snapshot:
        doc.add_paragraph()
        _hb = doc.add_paragraph().add_run("Bibliography used (snapshot)")
        _hb.bold = True; _hb.font.size = Pt(13); _hb.font.color.rgb = RGBColor.from_string(ACCENT)
        _nb = doc.add_paragraph().add_run(
            "The numbered reference list these citations were relinked against, captured at "
            "relink time. Keep it to re-check the [[REF #]] numbers above if the bibliography "
            "is later renumbered.")
        _nb.font.size = Pt(8.5); _nb.italic = True; _nb.font.color.rgb = RGBColor.from_string("666666")
        bt = doc.add_table(rows=1, cols=2); grid(bt)
        for i, t in enumerate(["#", "Reference"]):
            c = bt.rows[0].cells[i]; rn = c.paragraphs[0].add_run(t)
            rn.bold = True; rn.font.size = Pt(10)
            rn.font.color.rgb = RGBColor.from_string("FFFFFF"); shade(c, ACCENT)
        _bp = bt.rows[0]._tr.get_or_add_trPr(); _bh = OxmlElement('w:tblHeader')
        _bh.set(qn('w:val'), 'true'); _bp.append(_bh)
        for num in sorted(bib_snapshot):
            r2 = bt.add_row()
            c0 = r2.cells[0]; rr0 = c0.paragraphs[0].add_run(str(num))
            rr0.bold = True; rr0.font.size = Pt(9.5)
            rr0.font.color.rgb = RGBColor.from_string(ACCENT); shade(c0, FILL)
            rr1 = r2.cells[1].paragraphs[0].add_run(str(bib_snapshot[num])); rr1.font.size = Pt(9)
        bt.autofit = False; bt.allow_autofit = False
        for rw in bt.rows:
            rw.cells[0].width = Inches(0.5); rw.cells[1].width = Inches(6.5)

    buf = _io.BytesIO(); doc.save(buf)
    # fix python-docx zoom (schema-required percent)
    import zipfile
    zin = zipfile.ZipFile(_io.BytesIO(buf.getvalue())); out = _io.BytesIO()
    with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
        for it in zin.infolist():
            d = zin.read(it.filename)
            if it.filename == 'word/settings.xml':
                d = re.sub(rb'<w:zoom[^>]*/>', b'<w:zoom w:percent="100"/>', d)
            zout.writestr(it, d)
    return out.getvalue()
