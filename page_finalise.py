"""
Finalise — 07 · Citation Renumbering, 09 · Document Health Check, 10 · Batch Rename as tabs.
"""
import streamlit as st, zipfile, re, base64, io, html
from pathlib import Path
from shared import *

st.markdown(APP_CSS, unsafe_allow_html=True)
st.markdown('<div class="app-label">A practical group &nbsp; 4 of 5</div>', unsafe_allow_html=True)
st.markdown("## Finalise")
st.markdown("Renumber citations, run a health check, and batch rename.")
st.divider()

for _k, _v in [
    ("finalise_doc",     None),
    ("finalise_result",  None),
    ("health_analysis",  None),
    ("health_result",    None),
    ("rename_doc",       None),
    ("rename_result",    None),
]:
    if _k not in st.session_state:
        st.session_state[_k] = _v

_tab1, _tab2, _tab3 = st.tabs(["07 · Citation Renumbering", "09 · Document Health Check", "10 · Batch Rename"])

with _tab1:
        st.markdown('<div class="app-label">A practical tool &nbsp;07</div>', unsafe_allow_html=True)
        st.markdown("## Citation Renumbering")
        st.markdown('''<div class="instruction-box">
        <b>When to use:</b> After editing, citation superscript numbers and the bibliography
        are out of order. This tool renumbers them consistently using your chosen method.<br><br>
        <b>Before using:</b> Make sure EndNote has already formatted the bibliography so
        citations appear as plain superscript numbers in the text — not live field codes.
        </div>''', unsafe_allow_html=True)

        ren_file = st.file_uploader(
            "Word document (.docx) with formatted citations",
            type=["docx"], key="ren_doc"
        )

        method = st.radio(
            "Numbering method",
            [
                "Alphabetical — A=1, B=2... (sort bibliography by author last name)",
                "Order of appearance — first cited in text = 1, second = 2...",
            ]
        )

        if ren_file:
            if st.button("Renumber citations", type="primary"):
                with st.spinner("Scanning and renumbering..."):
                    raw_bytes = ren_file.read()
                    if "Alphabetical" in method:
                        fixed_bytes, mapping = renumber_citations_alpha(raw_bytes)
                        method_label = "alphabetically (A-Z by author)"
                    else:
                        fixed_bytes, mapping = renumber_citations_appearance(raw_bytes)
                        method_label = "by order of appearance"

                if not mapping:
                    st.warning(
                        "No superscript citation numbers or numbered bibliography found. "
                        "Make sure EndNote has formatted the bibliography first."
                    )
                else:
                    changed = {k: v for k, v in mapping.items() if k != v}
                    st.success(
                        f"Done — renumbered {method_label}. "
                        f"{len(mapping)} unique citations, {len(changed)} numbers changed."
                    )
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Unique citations",  len(mapping))
                    col2.metric("Numbers changed",   len(changed))
                    col3.metric("Already in order",  len(mapping) - len(changed))

                    st.download_button(
                        "⬇ Download renumbered document",
                        data=fixed_bytes,
                        file_name=Path(ren_file.name).stem + "_renumbered.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )

                    if changed:
                        with st.expander(f"Renumbering map ({len(changed)} changes)"):
                            col1, col2 = st.columns(2)
                            col1.markdown("**Old #**")
                            col2.markdown("**New #**")
                            for old in sorted(mapping):
                                new = mapping[old]
                                if old != new:
                                    col1.markdown(str(old))
                                    col2.markdown(str(new))

    # ─────────────────────────────────────────────────────────────────────────────
    # APP 6 UI — FIGURE INVENTORY

with _tab2:
        st.markdown('<div class="app-label">A practical tool &nbsp;09</div>', unsafe_allow_html=True)
        st.markdown("## Document Health Check")
        st.markdown('''<div class="instruction-box">
        Scans your Word document for common problems and presents each issue as a
        click-through review — accept individual fixes or accept all at once.<br><br>
        <b>Checks:</b> orphaned superscripts · garbled figure numbers (3334 → 33) ·
        broken citation fields · duplicate w:id values · orphan comment markup ·
        plain-text bibliography entries that should be EndNote fields
        </div>''', unsafe_allow_html=True)

        health_file = st.file_uploader("Word document (.docx)", type=["docx"], key="health_doc")

        if health_file:
            if st.button("Run health check", type="primary"):
                with st.spinner("Scanning document..."):
                    from lxml import etree as _et
                    from docx import Document as _Doc
                    import base64 as _b64

                    raw_bytes = health_file.read()
                    with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
                        doc_xml   = z.read('word/document.xml').decode('utf-8')
                        all_zfiles = {n: z.read(n) for n in z.namelist()}

                    doc = _Doc(io.BytesIO(raw_bytes))
                    W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

                    issues = []

                    # ── 1. Broken citation fields ──────────────────────────────
                    broken_a = len(re.findall(r'<w:instrText[^>]*> ADDIN EN\.CITE </w:instrText>', doc_xml))
                    broken_b = len(re.findall(r'<w:instrText[^>]*>\s*ADDIN EN\.CITE\.DATA\s*</w:instrText>', doc_xml))
                    if broken_a + broken_b > 0:
                        issues.append({
                            'id': 'broken_fields',
                            'severity': 'error',
                            'title': f'{broken_a + broken_b} broken citation field(s)',
                            'detail': f'{broken_a} empty instrText + {broken_b} ADDIN EN.CITE.DATA — EndNote cannot see these refs, causing bibliography undercount.',
                            'fix_label': 'Fix broken fields',
                            'count': broken_a + broken_b,
                        })

                    # ── 2. Hidden refs in fldData ──────────────────────────────
                    working_rns = set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', doc_xml))
                    fld_rns = set()
                    for b64r in re.findall(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>', doc_xml):
                        b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
                        pad = (4-len(b64)%4)%4
                        try:
                            dec = _b64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
                            for rn in re.findall(r'<RecNum>(\d+)</RecNum>', dec): fld_rns.add(rn)
                        except: pass
                    hidden = fld_rns - working_rns
                    if hidden:
                        issues.append({
                            'id': 'hidden_refs',
                            'severity': 'error',
                            'title': f'{len(hidden)} ref(s) hidden in fldData',
                            'detail': f'RecNums {sorted(int(x) for x in hidden)[:10]}{"..." if len(hidden)>10 else ""} — locked in backup blobs, invisible to EndNote.',
                            'fix_label': 'Restore hidden refs',
                            'count': len(hidden),
                        })

                    # ── 3. Orphaned superscripts ───────────────────────────────
                    orphan_count = 0
                    for para in doc.paragraphs:
                        has_field = any('EN.CITE' in (i.text or '') for i in para._p.findall(f'.//{{{W}}}instrText'))
                        if not has_field: continue
                        fd = 0
                        for run in para._p.findall(f'{{{W}}}r'):
                            for fc in run.findall(f'.//{{{W}}}fldChar'):
                                ft = fc.get(f'{{{W}}}fldCharType','')
                                if ft=='begin': fd+=1
                                elif ft=='end': fd=max(0,fd-1)
                            if fd>0: continue
                            rpr = run.find(f'{{{W}}}rPr')
                            if rpr is None: continue
                            va = rpr.find(f'{{{W}}}vertAlign')
                            if va is None or va.get(f'{{{W}}}val')!='superscript': continue
                            t = run.find(f'{{{W}}}t')
                            text = (t.text or '') if t is not None else ''
                            if text.strip() and re.match(r'[\d,;\s\.]+$', text.strip()):
                                orphan_count += 1
                    if orphan_count:
                        issues.append({
                            'id': 'orphan_superscripts',
                            'severity': 'warning',
                            'title': f'{orphan_count} orphaned superscript(s)',
                            'detail': 'Plain-text citation numbers sitting next to EndNote field codes — visible to reader but invisible to EndNote.',
                            'fix_label': 'Remove orphaned superscripts',
                            'count': orphan_count,
                        })

                    # ── 4. Garbled figure numbers (33+34 split) ────────────────
                    fig_count = 0
                    for para in doc.paragraphs:
                        runs = list(para.runs)
                        for i in range(len(runs)-1):
                            if (runs[i].text or '').endswith('33') and (runs[i+1].text or '').startswith('34'):
                                fig_count += 1
                    if fig_count:
                        issues.append({
                            'id': 'garbled_figs',
                            'severity': 'warning',
                            'title': f'{fig_count} garbled figure number(s)',
                            'detail': 'Chapter numbers got concatenated during merge (e.g. "Fig. 3334.1" instead of "Fig. 33.1"). Each occurrence will have the extra "34" removed.',
                            'fix_label': 'Fix figure numbers',
                            'count': fig_count,
                        })

                    # ── 5. Duplicate w:id values ───────────────────────────────
                    from collections import Counter as _Ctr
                    dup_ids = len({k for k,v in _Ctr(re.findall(r'\bw:id="(\d+)"', doc_xml)).items() if v>1})
                    if dup_ids:
                        issues.append({
                            'id': 'dup_ids',
                            'severity': 'error',
                            'title': f'{dup_ids} duplicate w:id value(s)',
                            'detail': 'Internal ID conflicts caused by copying paragraph XML between documents. Causes Word to show "unreadable content" or repair dialogs.',
                            'fix_label': 'Deduplicate IDs',
                            'count': dup_ids,
                        })

                    # ── 6. Orphan comment markup ───────────────────────────────
                    com_markup = len(re.findall(r'<w:comment(?:RangeStart|RangeEnd|Reference)\b', doc_xml))
                    has_comments_xml = 'word/comments.xml' in all_zfiles
                    if com_markup > 0 and not has_comments_xml:
                        issues.append({
                            'id': 'orphan_comments',
                            'severity': 'warning',
                            'title': f'{com_markup} orphaned comment markup element(s)',
                            'detail': 'Comment reference markers exist in the document body but no comments.xml file exists. Causes Word repair dialog on open.',
                            'fix_label': 'Remove orphaned comment markup',
                            'count': com_markup,
                        })

                    # ── 7. Unbalanced field codes ──────────────────────────────
                    begins = doc_xml.count('fldCharType="begin"')
                    ends   = doc_xml.count('fldCharType="end"')
                    seps   = doc_xml.count('fldCharType="separate"')
                    if begins != ends or begins != seps:
                        issues.append({
                            'id': 'unbalanced_fields',
                            'severity': 'error',
                            'title': f'Unbalanced field markers (begin:{begins}/sep:{seps}/end:{ends})',
                            'detail': 'Field code begin/separate/end markers are mismatched. This causes citation formatting failures and can make Word reject the file.',
                            'fix_label': 'Cannot auto-fix — manual repair needed',
                            'count': abs(begins - ends),
                        })

                    st.session_state['health_issues']    = issues
                    st.session_state['health_raw_bytes'] = raw_bytes
                    st.session_state['health_doc_name']  = health_file.name

            if 'health_issues' in st.session_state and st.session_state['health_issues'] is not None:
                issues    = st.session_state['health_issues']
                raw_bytes = st.session_state['health_raw_bytes']
                doc_name  = st.session_state['health_doc_name']

                if not issues:
                    st.success("✓ No issues found — document is clean.")
                else:
                    errors   = [i for i in issues if i['severity']=='error']
                    warnings = [i for i in issues if i['severity']=='warning']
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Issues found", len(issues))
                    c2.metric("Errors",   len(errors),   delta=str(len(errors))   if errors   else None, delta_color="inverse")
                    c3.metric("Warnings", len(warnings), delta=str(len(warnings)) if warnings else None, delta_color="inverse")

                    # Accept All button
                    fixable = [i for i in issues if 'Cannot' not in i['fix_label']]
                    if st.button(f"✓ Accept all {len(fixable)} fixable issue(s)", type="primary"):
                        st.session_state['health_accept_all'] = True

                    st.divider()
                    st.markdown("### Review each issue")

                    # Per-issue accept checkboxes
                    accepted = {}
                    for issue in issues:
                        fixable_issue = 'Cannot' not in issue['fix_label']
                        icon = '🔴' if issue['severity']=='error' else '🟡'
                        with st.expander(f"{icon} {issue['title']}", expanded=True):
                            st.markdown(issue['detail'])
                            if fixable_issue:
                                accepted[issue['id']] = st.checkbox(
                                    f"Accept fix: {issue['fix_label']}",
                                    value=st.session_state.get('health_accept_all', False),
                                    key=f"accept_{issue['id']}"
                                )
                            else:
                                st.warning(issue['fix_label'])

                    if st.button("Apply accepted fixes"):
                        with st.spinner("Applying fixes..."):
                            with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
                                doc_xml   = z.read('word/document.xml').decode('utf-8')
                                all_zfiles = {n: z.read(n) for n in z.namelist()}

                            applied = []

                            if accepted.get('broken_fields') or accepted.get('hidden_refs'):
                                doc_xml, n = fix_broken_fields(doc_xml)
                                applied.append(f"Fixed {n} broken field(s)")

                            if accepted.get('orphan_comments'):
                                before = len(re.findall(r'<w:comment(?:RangeStart|RangeEnd|Reference)\b', doc_xml))
                                doc_xml = re.sub(r'<w:commentRangeStart\b[^/]*/>', '', doc_xml)
                                doc_xml = re.sub(r'<w:commentRangeEnd\b[^/]*/>', '', doc_xml)
                                doc_xml = re.sub(r'<w:commentReference\b[^/]*/>', '', doc_xml)
                                applied.append(f"Removed {before} orphaned comment element(s)")

                            if accepted.get('dup_ids'):
                                from collections import Counter as _Ctr2
                                _ids = re.findall(r'\bw:id="(\d+)"', doc_xml)
                                _max = max(int(x) for x in _ids)
                                _nxt = [_max+1]; _seen = set()
                                def _fid(m):
                                    v=m.group(2)
                                    if v in _seen:
                                        nw=str(_nxt[0]); _nxt[0]+=1; return f'{m.group(1)}{nw}{m.group(3)}'
                                    _seen.add(v); return m.group(0)
                                doc_xml = re.sub(r'(w:id=")(\d+)(")', _fid, doc_xml)
                                applied.append("Deduplicated w:id values")

                            if accepted.get('garbled_figs'):
                                from docx import Document as _Doc2
                                all_zfiles['word/document.xml'] = doc_xml.encode('utf-8')
                                buf2 = io.BytesIO()
                                with zipfile.ZipFile(buf2,'w',zipfile.ZIP_DEFLATED) as zt:
                                    for nn,dd in all_zfiles.items(): zt.writestr(nn,dd)
                                buf2.seek(0)
                                d2 = _Doc2(buf2)
                                W2 = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                                n_fig = 0
                                for para in d2.paragraphs:
                                    runs = list(para.runs)
                                    for i in range(len(runs)-1):
                                        t1=(runs[i].text or ''); t2=(runs[i+1].text or '')
                                        if t1.endswith('33') and t2.startswith('34'):
                                            nt = t2[2:]
                                            te = runs[i+1]._r.find(f'{{{W2}}}t')
                                            if te is not None:
                                                te.text = nt
                                                if not nt: runs[i+1]._r.getparent().remove(runs[i+1]._r)
                                                n_fig += 1
                                buf3 = io.BytesIO(); d2.save(buf3); buf3.seek(0)
                                with zipfile.ZipFile(buf3) as zt2:
                                    doc_xml = zt2.read('word/document.xml').decode('utf-8')
                                    for nn in zt2.namelist():
                                        all_zfiles[nn] = zt2.read(nn)
                                applied.append(f"Fixed {n_fig} garbled figure number(s)")

                            if accepted.get('orphan_superscripts'):
                                all_zfiles['word/document.xml'] = doc_xml.encode('utf-8')
                                buf4 = io.BytesIO()
                                with zipfile.ZipFile(buf4,'w',zipfile.ZIP_DEFLATED) as zt:
                                    for nn,dd in all_zfiles.items(): zt.writestr(nn,dd)
                                buf4.seek(0)
                                fixed_bytes, n_orp = remove_orphan_superscripts(buf4.read())
                                with zipfile.ZipFile(io.BytesIO(fixed_bytes)) as zt3:
                                    doc_xml = zt3.read('word/document.xml').decode('utf-8')
                                    for nn in zt3.namelist():
                                        all_zfiles[nn] = zt3.read(nn)
                                applied.append(f"Removed {n_orp} orphaned superscript(s)")

                            # Save final
                            all_zfiles['word/document.xml'] = doc_xml.encode('utf-8')
                            out_buf = io.BytesIO()
                            with zipfile.ZipFile(out_buf,'w',zipfile.ZIP_DEFLATED) as zout:
                                for nn,dd in all_zfiles.items(): zout.writestr(nn,dd)
                            out_buf.seek(0)
                            final_bytes = out_buf.read()

                        if applied:
                            for msg in applied:
                                st.success(f"✓ {msg}")
                            st.download_button(
                                "⬇ Download repaired document",
                                data=final_bytes,
                                file_name=Path(doc_name).stem + "_repaired.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="primary"
                            )
                            # Re-run health check on fixed doc
                            if st.button("Re-run health check on repaired document"):
                                st.session_state['health_issues'] = None
                                st.rerun()
                        else:
                            st.info("No fixes were accepted.")

            # Initialise state
            if 'health_issues' not in st.session_state:
                st.session_state['health_issues'] = None
            if 'health_accept_all' not in st.session_state:
                st.session_state['health_accept_all'] = False

with _tab3:
        st.markdown('<div class="app-label">A practical tool &nbsp;10</div>', unsafe_allow_html=True)
        st.markdown("## Batch Rename")
        st.markdown('<div class="instruction-box">Apply bulk find-and-replace across your Word document using an Excel naming sheet. Use for final editing to update chapter names, author names, figure labels, section titles, etc.<br><br><b>Excel format:</b> Two columns — <code>Old Name</code> (find) and <code>New Name</code> (replace). One row per replacement.</div>',unsafe_allow_html=True)
        col1,col2=st.columns(2)
        with col1: ren_doc=st.file_uploader("Word document (.docx)",type=["docx"],key="batch_doc")
        with col2: ren_excel=st.file_uploader("Excel naming sheet (.xlsx)",type=["xlsx","xls"],key="batch_excel")
        col3,col4=st.columns(2)
        with col3: match_case=st.toggle("Match case",value=False)
        with col4: whole_word=st.toggle("Whole word only",value=False)
        pairs=[]
        if ren_excel:
            with st.spinner("Reading naming sheet..."):
                pairs=load_rename_pairs(ren_excel.read())
            if pairs:
                st.success(f"✓ {len(pairs)} rename pairs loaded.")
                with st.expander("Preview"):
                    col1,col2=st.columns(2)
                    col1.markdown("**Find**"); col2.markdown("**Replace**")
                    for old,new in pairs[:20]: col1.markdown(old); col2.markdown(new)
                    if len(pairs)>20: st.caption(f"...and {len(pairs)-20} more")
            else: st.error("Could not read pairs. Check Excel has 'Old Name' and 'New Name' columns.")
        if ren_doc and ren_excel and pairs:
            if st.button("Apply batch rename",type="primary"):
                with st.spinner("Applying..."):
                    fixed_bytes,report=batch_rename(ren_doc.read(),pairs,match_case,whole_word)
                replaced=[r for r in report if r['status']=='replaced']
                not_found=[r for r in report if r['status']=='not_found']
                total_changes=sum(r['count'] for r in replaced)
                col1,col2,col3=st.columns(3)
                col1.metric("Total pairs",len(report)); col2.metric("Replaced",len(replaced)); col3.metric("Not found",len(not_found))
                st.success(f"✓ {total_changes} replacement(s) across {len(replaced)} pairs.")
                st.download_button("⬇ Download renamed document",data=fixed_bytes,
                    file_name=Path(ren_doc.name).stem+"_renamed.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",type="primary")
                if replaced:
                    with st.expander(f"Replacements made ({len(replaced)})"):
                        for r in replaced:
                            st.markdown(f'<div class="ref-item ok"><b>{r["old"]}</b> → {r["new"]} <span style="float:right;font-size:0.75rem">{r["count"]}x</span></div>',unsafe_allow_html=True)
                if not_found:
                    with st.expander(f"Not found ({len(not_found)})"):
                        for r in not_found:
                            st.markdown(f'<div class="ref-item error"><b>{r["old"]}</b> — not found in document</div>',unsafe_allow_html=True)
