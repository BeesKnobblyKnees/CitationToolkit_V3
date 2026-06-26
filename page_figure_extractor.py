"""
Figure Extractor — page module for Citation Toolkit v3.
Extracts figures, tables, boxes, and plates from PDF or Word source files.
Outputs: PDF (one figure per page) AND Word document.
"""
import io, re, os, tempfile
from pathlib import Path

import streamlit as st
from shared import APP_CSS

st.markdown(APP_CSS, unsafe_allow_html=True)
st.markdown('<div class="app-label">A practical group &nbsp; 6 of 6</div>',
            unsafe_allow_html=True)
st.markdown("## Figure *Extractor*")
st.markdown(
    "Extract figures, tables, boxes, and plates from a published PDF or Word file. "
    "Outputs a **PDF** and **Word document** — one figure per page, clean crop, exact caption."
)
st.divider()

# ── dependency check ──────────────────────────────────────────────────────────
missing_deps = []
try:
    import fitz
except ImportError:
    missing_deps.append("PyMuPDF  →  pip3 install pymupdf")
try:
    from PIL import Image, ImageOps
except ImportError:
    missing_deps.append("Pillow  →  pip3 install Pillow")
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import (SimpleDocTemplate, Image as RLImage,
                                    Paragraph, Spacer, PageBreak)
except ImportError:
    missing_deps.append("ReportLab  →  pip3 install reportlab")

if missing_deps:
    st.error("Some packages need to be installed before this tool can run.")
    st.markdown("**Run these commands in Terminal, then restart the app:**")
    for d in missing_deps:
        st.code(f"pip3 install {d.split('→')[1].strip()}", language="bash")
    st.stop()

# ── imports (safe after check) ────────────────────────────────────────────────
import fitz
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import (SimpleDocTemplate, Image as RLImage,
                                 Paragraph, Spacer, PageBreak)

# ── constants ─────────────────────────────────────────────────────────────────
SCAN_DPI   = 200
OUT_DPI    = 300
PAGE_W     = 8.5  * inch
PAGE_H     = 11.0 * inch
MARGIN     = 0.75 * inch
CONTENT_W  = PAGE_W - 2 * MARGIN
CONTENT_H  = PAGE_H - 2 * MARGIN

CAPTION_RE = re.compile(
    r'(?:^|\n)\s*((?:Fig(?:ure)?\.?|Table|Box|Plate|Video)\s*\d+[\w\-\.]*)',
    re.IGNORECASE
)

# ── PDF helpers ───────────────────────────────────────────────────────────────
def pdf_render(pdf_bytes, page_num, dpi=200):
    doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[page_num]
    mat  = fitz.Matrix(dpi/72, dpi/72)
    pix  = page.get_pixmap(matrix=mat, alpha=False)
    img  = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    doc.close()
    return img

def pdf_image_boxes(pdf_bytes, page_num, dpi=200):
    doc   = fitz.open(stream=pdf_bytes, filetype="pdf")
    page  = doc[page_num]
    scale = dpi/72
    boxes = []
    for img_info in page.get_images(full=True):
        for rect in page.get_image_rects(img_info[0]):
            boxes.append((int(rect.x0*scale), int(rect.y0*scale),
                          int(rect.x1*scale), int(rect.y1*scale)))
    doc.close()
    return boxes

def pdf_caption_y(pdf_bytes, page_num, search_text, dpi=200):
    doc   = fitz.open(stream=pdf_bytes, filetype="pdf")
    page  = doc[page_num]
    rects = page.search_for(search_text[:50].strip())
    doc.close()
    if rects:
        return int(rects[0].y0 * dpi/72)
    return None

def pdf_extract_caption(pdf_bytes, page_num, label_text):
    doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[page_num]
    text = page.get_text("text")
    doc.close()
    idx = text.lower().find(label_text.lower()[:30])
    if idx == -1:
        return label_text
    chunk = text[idx:]
    end = re.search(
        r'\n\n|\n(?=(?:Fig(?:ure)?\.?\s*\d|Table\s*\d|Box\s*\d|Plate\s*\d|Video\s*\d))',
        chunk, re.IGNORECASE)
    if end:
        chunk = chunk[:end.start()]
    return re.sub(r'\s{2,}', ' ', re.sub(r'\n+', ' ', chunk)).strip()

def pdf_scan(pdf_bytes, dpi=200, page_filter=None):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total = len(doc); doc.close()
    pages_to_scan = sorted(page_filter) if page_filter else list(range(total))
    figures = []; seen = set()
    bar = st.progress(0, text="Scanning pages…")
    for idx, pn in enumerate(pages_to_scan):
        bar.progress((idx+1)/len(pages_to_scan),
                     text=f"Scanning page {pn+1} ({idx+1}/{len(pages_to_scan)})…")
        doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = doc[pn].get_text("text")
        doc.close()
        for m in CAPTION_RE.finditer(text):
            label = m.group(1).strip()
            norm  = re.sub(r'\s+', ' ', label.upper())
            if norm in seen: continue
            seen.add(norm)
            lm = re.match(r'(Fig(?:ure)?\.?|Table|Box|Plate|Video)\s*(\d+[\w\-\.]*)',
                          label, re.IGNORECASE)
            f_type = lm.group(1).rstrip('.').title() if lm else 'Figure'
            f_num  = lm.group(2) if lm else '?'
            figures.append(dict(
                source='pdf', page_num=pn,
                fig_type=f_type, fig_num=f_num,
                label=f"{f_type} {f_num}",
                caption=pdf_extract_caption(pdf_bytes, pn, label),
                page_img=pdf_render(pdf_bytes, pn, dpi),
                img_boxes=pdf_image_boxes(pdf_bytes, pn, dpi),
                caption_y=pdf_caption_y(pdf_bytes, pn, label, dpi),
                include=True, crop_top=0, crop_bot=0,
            ))
    bar.empty()
    return figures

# ── Word helpers ──────────────────────────────────────────────────────────────
def docx_scan(docx_bytes):
    import zipfile
    from lxml import etree
    doc   = Document(io.BytesIO(docx_bytes))
    paras = doc.paragraphs
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        rels_xml = z.read('word/_rels/document.xml.rels')
        media    = {n for n in z.namelist() if n.startswith('word/media/')}
    rels_root = etree.fromstring(rels_xml)
    rel_map   = {r.get('Id',''): 'word/'+r.get('Target','')
                 for r in rels_root if r.get('Target','').startswith('media/')}
    figures = []; seen = set()
    for pi, para in enumerate(paras):
        p_xml    = para._p.xml
        img_rids = [rid for rid, path in rel_map.items()
                    if rid in p_xml and path in media]
        if not img_rids: continue
        cap_parts = []
        for j in range(pi+1, min(pi+8, len(paras))):
            t = paras[j].text.strip()
            if not t: continue
            if CAPTION_RE.match(t):
                cap_parts.append(t)
                for k in range(j+1, min(j+6, len(paras))):
                    kt = paras[k].text.strip()
                    if not kt or CAPTION_RE.match(kt): break
                    cap_parts.append(kt)
                break
            if len(t) > 20:
                cap_parts.append(t); break
        caption = re.sub(r'\s{2,}', ' ', ' '.join(cap_parts)).strip() or f"Figure {len(figures)+1}"
        lm = CAPTION_RE.match(caption)
        f_type = 'Figure'; f_num = str(len(figures)+1)
        if lm:
            lm2 = re.match(r'(Fig(?:ure)?\.?|Table|Box|Plate|Video)\s*(\d+[\w\-\.]*)',
                           lm.group(1).strip(), re.IGNORECASE)
            if lm2: f_type = lm2.group(1).rstrip('.').title(); f_num = lm2.group(2)
        norm = re.sub(r'\s+',' ', f"{f_type} {f_num}".upper())
        if norm in seen: continue
        seen.add(norm)
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            raw = z.read(rel_map[img_rids[0]])
        try: pil = Image.open(io.BytesIO(raw)).convert('RGB')
        except: continue
        figures.append(dict(
            source='docx', page_num=pi,
            fig_type=f_type, fig_num=f_num,
            label=f"{f_type} {f_num}", caption=caption,
            page_img=pil, img_bytes=raw,
            img_boxes=[], caption_y=None,
            include=True, crop_top=0, crop_bot=0,
        ))
    return figures

# ── cropping ──────────────────────────────────────────────────────────────────
def smart_crop(page_img, img_boxes, caption_y_px, crop_top_pct=0, crop_bot_pct=0, pad=20):
    w, h = page_img.size
    manual_top = int(h * crop_top_pct/100)
    manual_bot = h - int(h * crop_bot_pct/100)
    if img_boxes:
        above = [b for b in img_boxes if b[3] <= (caption_y_px or h)+30] or img_boxes
        x0 = max(0, min(b[0] for b in above)-pad)
        y0 = max(manual_top, min(b[1] for b in above)-pad)
        x1 = min(w, max(b[2] for b in above)+pad)
        y1 = min(manual_bot, (caption_y_px-pad) if caption_y_px else max(b[3] for b in above)+pad)
        return page_img.crop((x0, y0, x1, max(y0+10, y1)))
    elif caption_y_px:
        return page_img.crop((0, manual_top, w, max(manual_top+10, caption_y_px-pad)))
    return page_img.crop((0, manual_top, w, manual_bot))

def get_cropped(fig, pdf_bytes=None):
    if fig['source'] == 'docx':
        img = Image.open(io.BytesIO(fig['img_bytes'])).convert('RGB')
        w, h = img.size
        top = int(h*fig['crop_top']/100); bot = h-int(h*fig['crop_bot']/100)
        if top > 0 or bot < h: img = img.crop((0, top, w, bot))
        return ImageOps.autocontrast(img, cutoff=0.5)
    page_img  = pdf_render(pdf_bytes, fig['page_num'], dpi=OUT_DPI)
    img_boxes = pdf_image_boxes(pdf_bytes, fig['page_num'], dpi=OUT_DPI)
    cap_y = int(fig['caption_y'] * OUT_DPI/SCAN_DPI) if fig['caption_y'] else None
    cropped = smart_crop(page_img, img_boxes, cap_y, fig['crop_top'], fig['crop_bot'])
    cropped = ImageOps.autocontrast(cropped, cutoff=0.5)
    try:
        bbox = ImageOps.invert(cropped.convert('L')).getbbox()
        if bbox:
            bw, bh = cropped.size; p = 15
            cropped = cropped.crop((max(0,bbox[0]-p), max(0,bbox[1]-p),
                                    min(bw,bbox[2]+p), min(bh,bbox[3]+p)))
    except: pass
    return cropped

def img_to_bytes(img):
    buf = io.BytesIO(); img.save(buf, format='PNG', dpi=(OUT_DPI,OUT_DPI)); buf.seek(0)
    return buf.read()

# ── output builders ───────────────────────────────────────────────────────────
def cap_html(text):
    m = re.match(r'^((?:Fig(?:ure)?\.?\s*\d+[\w\-\.]*|Table\s*\d+[\w\-\.]*|'
                 r'Box\s*\d+[\w\-\.]*|Plate\s*\d+[\w\-\.]*|Video\s*\d+[\w\-\.]*)'
                 r'\.?\s*)', text, re.IGNORECASE)
    if m: return f'<b>{m.group(1)}</b>{text[len(m.group(1)):]}'
    return text

def build_pdf(figures, captions, pdf_bytes=None):
    buf = io.BytesIO()
    cap_sty = ParagraphStyle('C', fontName='Times-Roman', fontSize=10,
                              leading=14, alignment=TA_LEFT, spaceAfter=6)
    story = []
    for idx, fig in enumerate([f for f in figures if f.get('include',True)]):
        if idx > 0: story.append(PageBreak())
        try: cropped = get_cropped(fig, pdf_bytes)
        except: continue
        img_w, img_h = cropped.size; asp = img_w/img_h
        cap_text = captions.get(fig['label'], fig['caption'])
        cap_h    = max(1, len(cap_text)//90+1) * 14/72*inch + 0.2*inch
        max_h    = CONTENT_H - cap_h - 0.3*inch
        dw = min(CONTENT_W, img_w/OUT_DPI*inch); dh = dw/asp
        if dh > max_h: dh = max_h; dw = dh*asp
        top_sp = (CONTENT_H - dh - cap_h) / 2
        if top_sp > 0: story.append(Spacer(1, top_sp))
        story.append(RLImage(io.BytesIO(img_to_bytes(cropped)), width=dw, height=dh))
        story.append(Spacer(1, 0.15*inch))
        story.append(Paragraph(cap_html(cap_text), cap_sty))
    SimpleDocTemplate(buf, pagesize=letter,
                      leftMargin=MARGIN, rightMargin=MARGIN,
                      topMargin=MARGIN, bottomMargin=MARGIN).build(story)
    buf.seek(0); return buf.read()

import re
_BAD_XML_CHARS = re.compile('[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\ufdd0-\ufddf\ufffe\uffff]')
def _xml_safe(s):
    if s is None:
        return ''
    if not isinstance(s, str):
        s = str(s)
    return _BAD_XML_CHARS.sub('', s)

import re
_BAD_XML_CHARS = re.compile('[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\ufdd0-\ufddf\ufffe\uffff]')
def _xml_safe(s):
    if s is None:
        return ''
    if not isinstance(s, str):
        s = str(s)
    return _BAD_XML_CHARS.sub('', s)

def build_word(figures, captions, pdf_bytes=None):
    doc     = Document()
    section = doc.sections[0]
    section.page_width  = int(8.5*914400); section.page_height = int(11.0*914400)
    m = int(0.75*914400)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = m
    cw = 8.5-1.5; ch = 11.0-1.5
    for idx, fig in enumerate([f for f in figures if f.get('include',True)]):
        if idx > 0: doc.add_page_break()
        try: cropped = get_cropped(fig, pdf_bytes)
        except: continue
        iw, ih = cropped.size; asp = iw/ih
        cap_text = captions.get(fig['label'], fig['caption'])
        cap_h    = max(1, len(cap_text)//90+1)*14/72+0.2
        max_h    = max(0.5, ch - cap_h - 0.3)
        dw = min(cw, iw/OUT_DPI); dh = dw/asp
        if dh > max_h: dh = max_h; dw = dh*asp
        dw = max(0.5, min(dw, cw))
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(io.BytesIO(img_to_bytes(cropped)), width=Inches(dw))
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cp.paragraph_format.space_before = Pt(4); cp.paragraph_format.space_after = Pt(4)
        cap_text = _xml_safe(cap_text)
        m2 = re.match(r'^((?:Fig(?:ure)?\.?\s*\d+[\w\-\.]*|Table\s*\d+[\w\-\.]*|'
                      r'Box\s*\d+[\w\-\.]*|Plate\s*\d+[\w\-\.]*|Video\s*\d+[\w\-\.]*)'
                      r'\.?\s*)', cap_text, re.IGNORECASE)
        if m2:
            br = cp.add_run(m2.group(1)); br.bold=True; br.font.name='Times New Roman'; br.font.size=Pt(10)
            rr = cp.add_run(cap_text[len(m2.group(1)):]); rr.font.name='Times New Roman'; rr.font.size=Pt(10)
        else:
            r = cp.add_run(cap_text); r.font.name='Times New Roman'; r.font.size=Pt(10)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

# ── session state ─────────────────────────────────────────────────────────────
for k,v in [('fe_figures',[]),('fe_pdf_bytes',None),('fe_src_name',''),
             ('fe_src_type',''),('fe_captions',{}),('fe_scanned',False)]:
    if k not in st.session_state: st.session_state[k] = v

# ── Step 1: Upload ────────────────────────────────────────────────────────────
st.markdown("### Step 1 — Upload source file")
uploaded = st.file_uploader("PDF or Word document", type=["pdf","docx"], key="fe_upload")
c1, c2 = st.columns(2)
with c1:
    quality  = st.selectbox("Scan quality",
                             ["Fast (150 DPI)","Standard (200 DPI)","High (300 DPI)"], index=1)
with c2:
    st.caption("Output: Times New Roman 10pt, 300 DPI, 0.75\" margins")
scan_dpi = {"Fast (150 DPI)":150,"Standard (200 DPI)":200,"High (300 DPI)":300}[quality]

if uploaded:
    src_bytes = uploaded.read()
    src_ext   = Path(uploaded.name).suffix.lower()
    if src_bytes != st.session_state.fe_pdf_bytes:
        st.session_state.fe_pdf_bytes = src_bytes
        st.session_state.fe_src_name  = uploaded.name
        st.session_state.fe_src_type  = src_ext
        st.session_state.fe_figures   = []
        st.session_state.fe_captions  = {}
        st.session_state.fe_scanned   = False

    # ── Page / chapter filter ──────────────────────────────────────────────
    if src_ext == '.pdf':
        _doc = fitz.open(stream=src_bytes, filetype="pdf")
        _total_pages = _doc.page_count
        _doc.close()
        st.caption(f"{uploaded.name} — {_total_pages} pages")

        st.markdown("**Limit scan to specific pages or a chapter** *(optional)*")
        _fc1, _fc2 = st.columns(2)
        with _fc1:
            _page_range = st.text_input(
                "Page range",
                placeholder="e.g. 45-92 or 45,46,47",
                help="Enter a range (45-92) or comma-separated page numbers. Leave blank to scan all pages.",
                key="fe_page_range"
            )
        with _fc2:
            _chapter_kw = st.text_input(
                "Chapter keyword filter",
                placeholder="e.g. Arthrogryposis",
                help="Only extract figures whose caption contains this word. Leave blank for all figures.",
                key="fe_chapter_kw"
            )

        # Parse page range into a set of 0-indexed page numbers
        _page_filter = None
        if _page_range.strip():
            _page_filter = set()
            for _part in _page_range.replace(' ','').split(','):
                if '-' in _part:
                    _a, _b = _part.split('-', 1)
                    try:
                        _page_filter.update(range(int(_a)-1, int(_b)))
                    except ValueError:
                        st.warning(f"Invalid range: {_part}")
                elif _part.isdigit():
                    _page_filter.add(int(_part)-1)
            if _page_filter:
                _valid = {p for p in _page_filter if 0 <= p < _total_pages}
                st.caption(f"Will scan {len(_valid)} of {_total_pages} pages "
                           f"(pages {min(_valid)+1}–{max(_valid)+1})")
                _page_filter = _valid
    else:
        st.caption(f"{uploaded.name} — Word document")
        _page_filter = None
        _chapter_kw  = st.text_input(
            "Chapter keyword filter",
            placeholder="e.g. Arthrogryposis",
            help="Only extract figures whose caption contains this word.",
            key="fe_chapter_kw"
        )

    if st.button("Scan for figures", type="primary", key="fe_scan"):
        with st.spinner("Scanning…"):
            if src_ext == '.pdf':
                figs = pdf_scan(src_bytes, dpi=scan_dpi, page_filter=_page_filter)
            else:
                figs = docx_scan(src_bytes)
        # Apply chapter keyword filter
        _kw = st.session_state.get("fe_chapter_kw","").strip().lower()
        if _kw:
            figs = [f for f in figs if _kw in f["caption"].lower() or _kw in f["label"].lower()]
            st.info(f'Keyword filter "{_kw}" applied — {len(figs)} matching figure(s)')
        st.session_state.fe_figures = figs
        st.session_state.fe_scanned = True
        st.session_state.fe_captions = {}
        if not figs:
            st.warning("No figures found. Check that captions start with Figure / Table / Box / Plate + number.")
        else:
            st.success(f"Found {len(figs)} figure(s). Review below.")

# ── Step 2: Review ────────────────────────────────────────────────────────────
if st.session_state.fe_scanned and st.session_state.fe_figures:
    figs = st.session_state.fe_figures
    st.divider()
    st.markdown(f"### Step 2 — Review {len(figs)} figure(s)")

    type_counts = {}
    for f in figs: type_counts[f['fig_type']] = type_counts.get(f['fig_type'],0)+1
    cols = st.columns(min(len(type_counts),6))
    for i,(t,c) in enumerate(sorted(type_counts.items())): cols[i].metric(t+"s",c)

    bc1, bc2 = st.columns(2)
    with bc1:
        if st.button("Select all", key="fe_selall"):
            for f in figs: f['include']=True; st.rerun()
    with bc2:
        if st.button("Deselect all", key="fe_desel"):
            for f in figs: f['include']=False; st.rerun()

    for i, fig in enumerate(figs):
        included = fig.get('include',True)
        chk_col, content_col = st.columns([1,11])
        with chk_col:
            inc = st.checkbox("", value=included, key=f"fe_inc_{i}")
            if inc != included: figs[i]['include']=inc; st.rerun()
        with content_col:
            st.markdown(f"**{fig['label']}** — "
                        f"{'Page '+str(fig['page_num']+1) if fig['source']=='pdf' else 'Word doc'}")
            with st.expander("Preview & adjust"):
                p1, p2 = st.columns(2)
                with p1:
                    st.caption("Source page")
                    thumb = fig['page_img'].copy(); thumb.thumbnail((280,360)); st.image(thumb)
                with p2:
                    st.caption("Auto-crop preview")
                    try:
                        if fig['source']=='pdf':
                            prev = smart_crop(fig['page_img'], fig['img_boxes'],
                                              fig['caption_y'], fig['crop_top'], fig['crop_bot'])
                        else:
                            prev = fig['page_img'].copy()
                        prev.thumbnail((280,360)); st.image(prev)
                    except Exception as e:
                        st.caption(f"Error: {e}")
                ac1, ac2 = st.columns(2)
                with ac1:
                    top = st.slider("Remove from top %",0,60,fig['crop_top'],1,key=f"fe_top_{i}")
                with ac2:
                    bot = st.slider("Remove from bottom %",0,60,fig['crop_bot'],1,key=f"fe_bot_{i}")
                figs[i]['crop_top']=top; figs[i]['crop_bot']=bot

            current = st.session_state.fe_captions.get(fig['label'], fig['caption'])
            st.markdown("**Caption** *(fix OCR errors only — do not paraphrase)*")
            edited = st.text_area("", value=current, height=85,
                                  key=f"fe_cap_{i}", label_visibility="collapsed")
            if edited != current: st.session_state.fe_captions[fig['label']] = edited

    # ── Step 3: Generate ──────────────────────────────────────────────────────
    included_n = sum(1 for f in figs if f.get('include',True))
    st.divider()
    st.markdown(f"### Step 3 — Generate outputs ({included_n} selected)")

    if included_n == 0:
        st.warning("No figures selected.")
    else:
        if st.button(f"Generate PDF + Word ({included_n} figures)", type="primary", key="fe_gen"):
            pdf_for_crop = st.session_state.fe_pdf_bytes if st.session_state.fe_src_type=='.pdf' else None
            with st.spinner("Cropping and building documents…"):
                try:
                    caps = {k: _xml_safe(v) for k, v in st.session_state.fe_captions.items()}
                    pdf_out  = build_pdf(figs, caps, pdf_bytes=pdf_for_crop)
                    word_out = build_word(figs, caps, pdf_bytes=pdf_for_crop)
                    stem = Path(st.session_state.fe_src_name).stem
                    st.success(f"Done — {included_n} figures.")
                    dl1, dl2 = st.columns(2)
                    with dl1:
                        st.download_button("Download PDF", data=pdf_out,
                            file_name=f"{stem}_figures.pdf", mime="application/pdf",
                            type="primary", use_container_width=True)
                    with dl2:
                        st.download_button("Download Word (.docx)", data=word_out,
                            file_name=f"{stem}_figures.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary", use_container_width=True)
                except Exception as e:
                    st.error(f"Generation failed: {e}"); st.exception(e)

        with st.expander("Download individual figures as PNG"):
            pdf_for_crop = st.session_state.fe_pdf_bytes if st.session_state.fe_src_type=='.pdf' else None
            for i, fig in enumerate(figs):
                if not fig.get('include',True): continue
                c1, c2 = st.columns([4,1])
                with c1: st.markdown(f"**{fig['label']}**")
                with c2:
                    try:
                        png = img_to_bytes(get_cropped(fig, pdf_bytes=pdf_for_crop))
                        st.download_button("PNG", data=png,
                            file_name=f"{fig['label'].replace(' ','_')}.png",
                            mime="image/png", key=f"fe_png_{i}")
                    except Exception as e:
                        st.caption(f"Error: {e}")
