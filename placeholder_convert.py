"""
placeholder_convert.py  --  turn citation placeholders into EndNote temporary
citations ({Author, Year #RecNum}) by resolving them against an .enlx library.

Placeholder kinds:
  * Typed author-year citations, detected either by HIGHLIGHT (a colored marker
    tells the app "this is a citation") or by PATTERN (any "(... 4-digit-year ...)"
    parenthesis). Pattern mode needs no highlighting but only touches a parenthesis
    when at least one reference inside it resolves to the library, so ordinary
    asides like "(termed rebound deformity)" or "(see 2019 update)" are left alone.
  * [[REF n, n, n]] markers (no highlight needed; found by text pattern), each
    number resolved via a numbered reference list (number -> author/year).

Output states (per reference):
  * applied    -> {Author, Year #RecNum}, no highlight
  * suggested  -> surname matched but year/journal differs; original text kept on
                  an ORANGE background, not auto-applied unless apply_near=True
  * unresolved -> not in the library; original text kept on a RED background
"""
import io, re, zipfile, sqlite3, difflib, html, copy
from docx import Document
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.oxml import parse_xml
from lxml import etree
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from citation_bibrelink_module import parse_bibliography
except Exception:
    parse_bibliography = None

ORANGE = "FFC000"   # suggested
RED = "FF5B5B"      # unresolved
_CITE_RE = re.compile(r'\(([^()]*\b(?:18|19|20)\d{2}\b[^()]*)\)')


# ── library ────────────────────────────────────────────────────────────────
def load_library(enlx_bytes):
    zf = zipfile.ZipFile(io.BytesIO(enlx_bytes))
    eni = next((n for n in zf.namelist() if n.endswith('sdb.eni')), None)
    if not eni:
        raise ValueError("No sdb.eni database found inside the .enlx file.")
    raw = zf.read(eni)
    import tempfile
    tmp = tempfile.NamedTemporaryFile(suffix='.eni', delete=False)
    tmp.write(raw); tmp.close()
    con = sqlite3.connect(tmp.name); con.row_factory = sqlite3.Row
    def clean(s): return re.sub(r'<[^>]+>', '', s or '').replace('\r', ' / ').replace('\n', ' ').strip()
    def surnames(au):
        out = set()
        for a in au.split(' / '):
            a = a.strip()
            if a:
                s = a.split(',')[0].strip().lower()
                if len(s) >= 3:
                    out.add(s)
        return out
    lib = []
    for r in con.execute('SELECT id,trash_state,author,year,title,secondary_title,volume,pages FROM refs'):
        if (r['trash_state'] or 0) != 0:
            continue
        au = clean(r['author']); first = au.split(' / ')[0] if au else ''
        sur = first.split(',')[0].strip()
        lib.append({'id': r['id'], 'sur': sur, 'surl': sur.lower(), 'year': str(r['year'] or ''),
                    'jour': clean(r['secondary_title']), 'title': clean(r['title']),
                    'auths': surnames(au),
                    'vol': str(r['volume'] or '').strip(),
                    'pages': re.findall(r'\d+', r['pages'] or '')})
    con.close()
    return lib


def _content_score(rec, hint):
    """How well a library record's title/journal/co-authors/volume/pages match
    the full reference text. Used to separate same-author-same-year papers."""
    hl = (hint or '').lower()
    htoks = set(re.findall(r'[a-z]{4,}', hl))
    hnums = set(re.findall(r'\d+', hint or ''))
    ttoks = set(re.findall(r'[a-z]{4,}', (rec['title'] or '').lower()))
    jtoks = set(re.findall(r'[a-z]{4,}', (rec['jour'] or '').lower()))
    t = len(ttoks & htoks)
    j = len(jtoks & htoks)
    a = len(rec.get('auths', set()) & htoks)
    vp = (1 if rec.get('vol') and rec['vol'] in hnums else 0)
    vp += sum(1 for p in rec.get('pages', []) if p in hnums)
    return t * 1.0 + j * 1.0 + a * 0.6 + min(vp, 2) * 0.5


def match_to_library(surname, year, hint, lib):
    sl = (surname or '').lower()
    # candidates: same year, surname close
    cands = []
    for r in lib:
        if not year or r['year'] != year:
            continue
        ratio = difflib.SequenceMatcher(None, sl, r['surl']).ratio()
        starts = bool(sl) and (r['surl'].startswith(sl[:4]) or sl.startswith(r['surl'][:4]))
        if ratio >= 0.85 or starts:
            cands.append(r)

    if len(cands) == 1:
        return cands[0], 1.0, 'exact'

    if len(cands) > 1:
        scored = sorted(((_content_score(r, hint), r['id'], r) for r in cands),
                        key=lambda x: (x[0], -x[1]), reverse=True)
        top, second = scored[0], scored[1]
        # confident only with real content signal AND a clear margin
        if top[0] >= 2.0 and (top[0] - second[0]) >= 1.5:
            return top[2], top[0], 'exact'
        # some signal but not decisive -> best guess, flagged for verification
        if top[0] >= 1.0 and top[0] > second[0]:
            return top[2], top[0], 'near'
        # indistinguishable from the hint (e.g. typed citation with no title) -> flag
        return cands[0], 0.0, 'ambiguous'

    # no same-year candidate -> looser cross-year fallback
    near = None
    for r in lib:
        ratio = difflib.SequenceMatcher(None, sl, r['surl']).ratio()
        starts = bool(sl) and (r['surl'].startswith(sl[:5]) or sl.startswith(r['surl'][:5]))
        if ratio >= 0.8 or starts:
            jb = 0.3 if re.search(re.escape(r['surl'][:4]), (hint or '').lower()) else 0.0
            score = ratio + (0.2 if starts else 0) + 0.2 + jb
            if near is None or score > near[0]:
                near = (score, r)
    if near and near[0] >= 0.9:
        return near[1], near[0], 'near'
    return None, 0.0, 'none'


def _token(rec):
    return "%s, %s #%d" % (rec['sur'], rec['year'], rec['id'])


def _parse_bib_full(b):
    """num -> full untruncated reference text (for content-based matching)."""
    full = {}
    try:
        doc = Document(io.BytesIO(b))
    except Exception:
        return full
    for p in doc.paragraphs:
        m = re.match(r'^(\d{1,3})\s+(.*)', (p.text or '').strip())
        if m and len(m.group(2)) >= 15:
            full[int(m.group(1))] = m.group(2)
    return full


def _parse_ref_text(text):
    s = text.strip().lstrip('.').lstrip('(').strip()
    ym = re.search(r'\b(?:18|19|20)\d{2}\b', s)
    year = ym.group(0) if ym else ''
    sm = re.match(r"([A-Za-z'\u00C0-\u017F\-]+)", s)
    sur = sm.group(1) if sm else ''
    jh = s.split(year)[-1] if year else s
    return sur, year, jh


def _is_green(run, color='GREEN'):
    try:
        return run.font.highlight_color is not None and color in str(run.font.highlight_color)
    except Exception:
        return False


def _has_shd(run):
    rpr = run._r.find(qn('w:rPr'))
    return rpr is not None and rpr.find(qn('w:shd')) is not None


def _shade_run(run, fill):
    rpr = run._r.get_or_add_rPr()
    for old in rpr.findall(qn('w:shd')):
        rpr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    rpr.append(shd)


def _new_like(paragraph, ref_run, text, fill=None):
    """New run that inherits ref_run's font (minus highlight/shading), optional fill."""
    nr = paragraph.add_run(text)
    nr.font.highlight_color = None
    src = ref_run._r.find(qn('w:rPr'))
    if src is not None:
        ex = nr._r.find(qn('w:rPr'))
        if ex is not None:
            nr._r.remove(ex)
        clone = copy.deepcopy(src)
        for tag in ('w:highlight', 'w:shd'):
            for e in clone.findall(qn(tag)):
                clone.remove(e)
        nr._r.insert(0, clone)
    if fill:
        _shade_run(nr, fill)
    return nr


def _insert_like(paragraph, after_r, ref_run, text, fill=None):
    nr = _new_like(paragraph, ref_run, text, fill)
    after_r.addnext(nr._r)
    return nr._r


_SUP_STYLES = ('citsup', 'sup', 'Superscript', 'FootnoteReference', 'EndNoteBibliography')


def _is_sup_cite(run):
    """A run that is superscript-formatted and whose text is citation numbers."""
    rpr = run._r.find(qn('w:rPr'))
    if rpr is None:
        return False
    va = rpr.find(qn('w:vertAlign'))
    is_sup = va is not None and va.get(qn('w:val')) == 'superscript'
    if not is_sup:
        rs = rpr.find(qn('w:rStyle'))
        is_sup = rs is not None and rs.get(qn('w:val')) in _SUP_STYLES
    if not is_sup:
        return False
    t = run.text or ''
    return bool(re.search(r'\d', t)) and bool(re.fullmatch(r'[\d,;\s\u2013\u2012-]+', t))


def _unsuperscript(run):
    """Drop superscript so an inserted {Author, Year #Rec} reads as normal text
    (EndNote re-applies the superscript style on Update Citations)."""
    rpr = run._r.find(qn('w:rPr'))
    if rpr is None:
        return
    for va in rpr.findall(qn('w:vertAlign')):
        rpr.remove(va)
    for rs in rpr.findall(qn('w:rStyle')):
        if rs.get(qn('w:val')) in _SUP_STYLES:
            rpr.remove(rs)


_SUP_UNITS = {'cm', 'mm', 'm', 'km', 'nm', 'um', 'dm', 'kg', 'g', 'mg', 'ug', 'ng',
              'ml', 'l', 'dl', 'mol', 'mmol', 'kpa', 'mmhg', 'hz', 'db'}
_SUP_VARS = set('xyztnprkmij')


_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
# Tracked-change wrappers whose runs DISAPPEAR when changes are accepted.
_TC_DROP = (_W + 'del', _W + 'moveFrom')


def _effective_runs(para):
    """python-docx's paragraph.runs only returns runs that are DIRECT children of
    <w:p>; runs nested inside tracked-change wrappers (<w:ins>, <w:moveTo>, and
    also <w:del>/<w:moveFrom>) are invisible to it. A citation sitting inside an
    inserted or moved span is therefore never seen by the converter.

    This returns Run objects for every <w:r> in document order that will survive
    an 'accept all changes' - i.e. direct runs plus those inside <w:ins>/<w:moveTo>
    - while skipping runs inside <w:del>/<w:moveFrom> (which would be removed)."""
    out = []
    for r in para._p.iter(_W + 'r'):
        anc = r.getparent()
        drop = False
        while anc is not None and anc is not para._p:
            if anc.tag in _TC_DROP:
                drop = True
                break
            anc = anc.getparent()
        if not drop:
            out.append(Run(r, para))
    return out


_XREF_STYLES = {'crossref', 'crossrefs', 'xref', 'xrefs'}
_CUE_RE = re.compile(r'(?:\bref(?:erence)?s?\b|\bsee\b)\s*:?\s*$', re.I)


def _xref_styled(run):
    """True if a run carries a cross-reference character style (the style Word
    applies to manually-keyed reference numbers like 'References 7, 8, 13')."""
    rpr = run._r.find(qn('w:rPr'))
    if rpr is None:
        return False
    rs = rpr.find(qn('w:rStyle'))
    return rs is not None and (rs.get(qn('w:val')) or '').lower() in _XREF_STYLES


def _is_num_run(run):
    t = run.text or ''
    return bool(re.fullmatch(r'\s*\d{1,3}[a-z]?\s*', t))


def _clear_rstyle(run):
    rpr = run._r.find(qn('w:rPr'))
    if rpr is not None:
        for rs in rpr.findall(qn('w:rStyle')):
            rpr.remove(rs)


def _iter_body_and_table_paras(doc):
    """All paragraphs python-docx's doc.paragraphs misses inside tables, plus the
    body paragraphs, de-duplicated and in a stable order."""
    seen, out = set(), []

    def add(paras):
        for p in paras:
            key = id(p._p)
            if key not in seen:
                seen.add(key); out.append(p)

    def walk_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    add(cell.paragraphs)
                    walk_tables(cell.tables)

    add(doc.paragraphs)
    walk_tables(doc.tables)
    return out


def _looks_like_exponent(prev_text):
    """True if a superscript number is really an exponent/unit power (cm^2, x^2,
    m^3) rather than a citation. Citations follow a word, a year or punctuation."""
    if not prev_text or prev_text[-1].isspace():
        return False                       # space before it -> citation
    tail = re.search(r'([A-Za-z\u00b5]+)$', prev_text)
    if not tail:
        return False                       # preceded by a digit/punctuation -> citation
    tok = tail.group(1).lower()
    return tok in _SUP_UNITS or (len(tok) == 1 and tok in _SUP_VARS)


# ── numbered citation markers: [[REF n]] / [n] / (n) ──────────────────────── #
_NUM_PAT = {
    'refmark': re.compile(r'\[\[REF\s*([\d,\s\u2013-]+)\]\]'),
    'bracket': re.compile(r'\[\s*(\d{1,3}(?:\s*[,\u2013-]\s*\d{1,3})*)\s*\]'),
    'paren':   re.compile(r'\(\s*(\d{1,3}(?:\s*[,\u2013-]\s*\d{1,3})*)\s*\)'),
}
_FMT = {'refmark': '[[REF %s]]', 'bracket': '[%s]', 'paren': '(%s)'}


def _expand_nums(inner):
    out = []
    for part in re.split(r'\s*,\s*', inner.strip()):
        rng = re.match(r'(\d+)\s*[\u2013-]\s*(\d+)$', part)
        if rng:
            out += list(range(int(rng.group(1)), int(rng.group(2)) + 1))
        elif part.isdigit():
            out.append(int(part))
    return out


def _fmt_marker(style, nums):
    return _FMT[style] % ', '.join(str(n) for n in nums)


def _splice_marker(para, s, e, segments):
    """Replace the [s,e) span of a paragraph's text (which may cross runs) with
    the given (text, fill) segments, preserving surrounding runs/formatting."""
    runs = _effective_runs(para)
    texts = [r.text or '' for r in runs]
    pos, spans = 0, []
    for t in texts:
        spans.append((pos, pos + len(t))); pos += len(t)
    ri_s = next((ri for ri, (a, b) in enumerate(spans) if a <= s < b), None)
    ri_e = next((ri for ri, (a, b) in enumerate(spans) if a < e <= b), None)
    if ri_s is None or ri_e is None:
        return
    prefix = texts[ri_s][:s - spans[ri_s][0]]
    suffix = texts[ri_e][e - spans[ri_e][0]:]
    ref = runs[ri_s]
    runs[ri_s].text = prefix
    runs[ri_s].font.highlight_color = None
    for ri in range(ri_s + 1, ri_e + 1):
        runs[ri].text = ''
    anchor = runs[ri_s]._r
    for segtext, fill in segments:
        if segtext == '':
            continue
        anchor = _insert_like(para, anchor, ref, segtext, fill)
    if suffix:
        _insert_like(para, anchor, ref, suffix, None)


# ── conversion ───────────────────────────────────────────────────────────── #
def convert(docx_bytes, enlx_bytes, do_green=True, do_refmarkers=True,
            highlight='GREEN', apply_near=False, bib_source_bytes=None,
            typed_detect='highlight', ref_styles=('refmark',)):
    """typed_detect in {'highlight','pattern','both'} controls how typed
    citations are found (only relevant when do_green=True).
    ref_styles selects which numbered markers to detect: any of
    'refmark' ([[REF n]]), 'bracket' ([n]), 'paren' ((n))."""
    lib = load_library(enlx_bytes)
    bib, bibtext = {}, {}
    bibfull = {}
    if do_refmarkers and parse_bibliography is not None:
        try:
            bib, bibtext = parse_bibliography(bib_source_bytes or docx_bytes)
            bibfull = _parse_bib_full(bib_source_bytes or docx_bytes)
        except Exception:
            bib, bibtext, bibfull = {}, {}, {}

    doc = Document(io.BytesIO(docx_bytes))
    _target_paras = _iter_body_and_table_paras(doc)

    # Footnotes and endnotes live in separate parts python-docx's doc.paragraphs
    # never exposes. Parse those trees now and fold their paragraphs into the scan
    # set so EVERY pass (green, [[REF]], [#], (#), superscript, number lists) runs
    # over them too; the mutated trees are serialised back in after saving.
    _note_trees = {}
    for _kind, _fname in (('footnotes', 'word/footnotes.xml'),
                          ('endnotes', 'word/endnotes.xml')):
        _blob = None
        for _rel in doc.part.rels.values():
            if _kind in _rel.reltype:
                try:
                    _blob = _rel.target_part.blob
                except Exception:
                    _blob = None
                break
        if not _blob:
            continue
        try:
            _tree = parse_xml(_blob)
        except Exception:
            continue
        _note_trees[_fname] = _tree
        _target_paras = _target_paras + [Paragraph(_p, doc)
                                         for _p in _tree.iter(qn('w:p'))]

    report = []
    use_hl = do_green and typed_detect in ('highlight', 'both')
    use_pat = do_green and typed_detect in ('pattern', 'both')

    def classify(key, surname, year, hint):
        # Apply the best-guess record for exact, near AND ambiguous matches.
        # near/ambiguous are still reported as 'suggested' so they show up in
        # the verify list, but they DO get the {Author, Year #RecNum} applied.
        rec, score, kind = match_to_library(surname, year, hint, lib)
        if rec and kind in ('exact', 'near', 'ambiguous'):
            return ('resolved' if kind == 'exact' else 'suggested'), (key, rec), kind
        return 'missing', (key, None), 'none'

    def split_refs(text):
        resolved, suggested, missing, ambig = [], [], [], []
        bucket = {'resolved': resolved, 'suggested': suggested, 'missing': missing}
        for piece in text.split(';'):
            piece = piece.strip()
            if not piece or not re.search(r'\d', piece):
                continue
            sur, year, hint = _parse_ref_text(piece)
            state, item, kind = classify(piece, sur, year, hint)
            bucket[state].append(item)
            if kind == 'ambiguous':
                ambig.append(item[0])
        return resolved, suggested, missing, ambig

    # ---- pass 1: highlighted typed citations ----
    for para in _target_paras:
        runs = _effective_runs(para)
        n = len(runs)
        i = 0
        while i < n:
            run = runs[i]
            txt = run.text or ''

            if use_hl and _is_green(run, highlight):
                members = [i]; j = i + 1
                while j < n:
                    if _is_green(runs[j], highlight):
                        members.append(j); j += 1
                    elif re.fullmatch(r'[\s;,]*', runs[j].text or '') and j + 1 < n and _is_green(runs[j + 1], highlight):
                        members.append(j); j += 1
                    else:
                        break
                green_text = ''.join(runs[k].text for k in members)
                lead = re.match(r'^([.\s]*)', green_text).group(1)
                body = green_text[len(lead):].lstrip('(').rstrip().rstrip(')').rstrip()
                resolved, suggested, missing, ambig = split_refs(body)
                report.append({'kind': 'green', 'orig': green_text.strip(),
                               'resolved': resolved, 'suggested': suggested,
                               'missing': missing, 'ambiguous': ambig})
                first, last = members[0], members[-1]
                if first - 1 >= 0 and (runs[first - 1].text or '').rstrip().endswith('('):
                    pt = runs[first - 1].text
                    runs[first - 1].text = pt[:pt.rstrip().rfind('(')]
                if last + 1 < n and (runs[last + 1].text or '').lstrip().startswith(')'):
                    ft = runs[last + 1].text; pos = ft.find(')')
                    runs[last + 1].text = ft[:pos] + ft[pos + 1:]
                applied = resolved + suggested
                cite = ('{' + '; '.join(_token(r) for _, r in applied) + '}') if applied else ''
                ref_run = runs[members[0]]
                ref_run.text = lead + cite
                ref_run.font.highlight_color = None
                for k in members[1:]:
                    runs[k].text = ''
                    runs[k].font.highlight_color = None
                anchor = ref_run._r
                if missing:
                    anchor = _insert_like(para, anchor, ref_run,
                        ((' ' if cite else '') + '(' + '; '.join(k for k, _ in missing) + ')'), RED)
                i = j
                continue

            i += 1

    # ---- pass 2: pattern-detected typed citations (no highlight) ----
    if use_pat:
        for para in _target_paras:
            for run in list(_effective_runs(para)):
                t = run.text or ''
                if not t or '[[REF' in t or ('{' in t and '#' in t) or _has_shd(run):
                    continue
                matches = list(_CITE_RE.finditer(t))
                if not matches:
                    continue
                segs = []; pos = 0; touched = False
                for mt in matches:
                    resolved, suggested, missing, ambig = split_refs(mt.group(1))
                    if not (resolved or suggested):
                        continue  # not a recognizable citation -> leave it in place
                    touched = True
                    report.append({'kind': 'pattern', 'orig': mt.group(0),
                                   'resolved': resolved, 'suggested': suggested,
                                   'missing': missing, 'ambiguous': ambig})
                    segs.append(('plain', t[pos:mt.start()]))
                    applied = resolved + suggested
                    if applied:
                        segs.append(('plain', '{' + '; '.join(_token(r) for _, r in applied) + '}'))
                    if missing:
                        segs.append((RED, (' ' if applied else '') + '(' + '; '.join(k for k, _ in missing) + ')'))
                    pos = mt.end()
                if not touched:
                    continue
                segs.append(('plain', t[pos:]))
                run.text = segs[0][1]
                anchor = run._r
                for fill, segtext in segs[1:]:
                    anchor = _insert_like(para, anchor, run, segtext,
                                          None if fill == 'plain' else fill)

    # ---- pass 3: numbered markers  [[REF n]] / [n] / (n) ----
    if do_refmarkers and ref_styles:
        pats = [(s, _NUM_PAT[s]) for s in ref_styles if s in _NUM_PAT]
        for para in _target_paras:
            full = ''.join(r.text or '' for r in _effective_runs(para))
            if not full:
                continue
            found = []
            for style, pat in pats:
                for m in pat.finditer(full):
                    found.append((m.start(), m.end(), style, m.group(1)))
            if not found:
                continue
            found.sort()
            chosen, last_end = [], -1
            for s, e, style, inner in found:
                if s >= last_end:
                    chosen.append((s, e, style, inner)); last_end = e
            actions = []
            for s, e, style, inner in chosen:
                nums = _expand_nums(inner)
                if not nums:
                    continue
                in_bib = any(bib.get(num) is not None for num in nums)
                if style != 'refmark' and not in_bib:
                    continue  # bracketed/parenthesised number that isn't a reference
                resolved, suggested, missing, ambig = [], [], [], []
                bucket = {'resolved': resolved, 'suggested': suggested, 'missing': missing}
                for num in nums:
                    sy = bib.get(num); hint = bibfull.get(num) or bibtext.get(num, '')
                    if not sy:
                        missing.append((num, None)); continue
                    st, item, kind = classify(num, sy[0], sy[1], hint)
                    bucket[st].append(item)
                    if kind == 'ambiguous':
                        ambig.append(num)
                report.append({'kind': style, 'orig': full[s:e],
                               'resolved': resolved, 'suggested': suggested,
                               'missing': missing, 'ambiguous': ambig})
                applied = resolved + suggested
                cite = ('{' + '; '.join(_token(r) for _, r in applied) + '}') if applied else ''
                segs = []
                if cite:
                    segs.append((cite, None))
                if missing:
                    segs.append(((' ' if cite else '') + _fmt_marker(style, [k for k, _ in missing]), RED))
                actions.append((s, e, segs))
            for s, e, segs in reversed(actions):
                _splice_marker(para, s, e, segs)

    # ---- pass 4: superscript numbered citations  (Vancouver / JAMA ^1,2) ----
    if do_refmarkers and 'superscript' in (ref_styles or ()):
        for para in _target_paras:
            runs = _effective_runs(para)
            n = len(runs)
            i = 0
            while i < n:
                if not _is_sup_cite(runs[i]):
                    i += 1
                    continue
                members = [i]; j = i + 1
                while j < n:
                    if _is_sup_cite(runs[j]):
                        members.append(j); j += 1
                    elif re.fullmatch(r'[\s,;\u2013\u2012-]+', runs[j].text or '') and j + 1 < n and _is_sup_cite(runs[j + 1]):
                        members.append(j); j += 1
                    else:
                        break
                group_text = ''.join(runs[k].text or '' for k in members)
                nums = _expand_nums(group_text.replace(';', ','))
                prev_text = runs[members[0] - 1].text if members[0] > 0 else ''
                if not nums or _looks_like_exponent(prev_text) or not bib:
                    i = j; continue
                maxnum = max(set(bib) | set(bibtext)) if (bib or bibtext) else 0
                resolved, suggested, missing, ambig, dangling = [], [], [], [], []
                bucket = {'resolved': resolved, 'suggested': suggested, 'missing': missing}
                for num in nums:
                    sy = bib.get(num)
                    if not sy:
                        if num in bibtext:
                            # the reference exists in the list but its surname/year
                            # couldn't be parsed -> unresolved (can't match), NOT dangling.
                            missing.append((num, None))
                        elif num <= maxnum:
                            # a real gap inside the numbered range = dangling citation
                            # (reference deleted, or the in-text number is wrong).
                            dangling.append(num)
                        continue
                    hint = bibfull.get(num) or bibtext.get(num, '')
                    st, item, kind = classify(num, sy[0], sy[1], hint)
                    bucket[st].append(item)
                    if kind == 'ambiguous':
                        ambig.append(num)
                if not (resolved or suggested or missing or dangling):
                    i = j; continue
                report.append({'kind': 'superscript', 'orig': group_text.strip(),
                               'resolved': resolved, 'suggested': suggested,
                               'missing': missing, 'ambiguous': ambig,
                               'dangling': dangling})
                applied = resolved + suggested
                cite = ('{' + '; '.join(_token(r) for _, r in applied) + '}') if applied else ''
                ref_run = runs[members[0]]
                # insert any unresolved + dangling numbers first, still superscript + red
                flagged = [k for k, _ in missing] + dangling
                if flagged:
                    _insert_like(para, ref_run._r, ref_run,
                                 (' ' if cite else '') + ', '.join(str(k) for k in sorted(set(flagged))), RED)
                # the resolved citation becomes normal-baseline {Author, Year #Rec}
                ref_run.text = cite
                _unsuperscript(ref_run)
                ref_run.font.highlight_color = None
                for k in members[1:]:
                    runs[k].text = ''
                i = j

    # ---- pass 5: plain reference-number lists  ("References 7, 8, 13, 17-19") ----
    # Numbers carry a cross-reference character style (or follow a "References"/
    # "See" cue). These are how figure/table reference lists and footnotes cite
    # the bibliography. Handles comma lists and en-dash ranges; resolves each
    # number to the library and replaces the list with EndNote temp citations.
    def _process_numlists(paragraphs):
        if not (do_refmarkers and 'numlist' in (ref_styles or ()) and bib):
            return
        maxnum = max(set(bib) | set(bibtext)) if (bib or bibtext) else 0
        for para in paragraphs:
            runs = _effective_runs(para)
            n = len(runs); i = 0
            while i < n:
                r = runs[i]
                cue = _CUE_RE.search(''.join((runs[k].text or '')
                                             for k in range(max(0, i - 2), i)))
                if not (_is_num_run(r) and (_xref_styled(r) or cue)):
                    i += 1; continue
                seq, members, j = [], [], i
                while j < n:
                    rj = runs[j]; tj = rj.text or ''
                    if _is_num_run(rj) and (_xref_styled(rj) or seq or cue):
                        seq.append(('n', int(re.search(r'\d+', tj).group())))
                        members.append(j); j += 1
                    elif re.fullmatch(r'\s*[\u2013\u2012\u2010-]\s*', tj):
                        seq.append(('dash', None)); members.append(j); j += 1
                    elif re.fullmatch(r'\s*[,;]\s*', tj):
                        seq.append(('sep', None)); members.append(j); j += 1
                    else:
                        break
                while seq and seq[-1][0] != 'n':            # trim trailing separators
                    seq.pop(); members.pop()
                if sum(1 for k, _ in seq if k == 'n') == 0:
                    i = j; continue
                nums, x = [], 0                              # expand ranges
                while x < len(seq):
                    if seq[x][0] == 'n':
                        a = seq[x][1]; nums.append(a)
                        if x + 2 < len(seq) and seq[x + 1][0] == 'dash' and seq[x + 2][0] == 'n':
                            b = seq[x + 2][1]
                            if a < b < a + 50:
                                nums.extend(range(a + 1, b + 1))
                            x += 3; continue
                        x += 1
                    else:
                        x += 1
                resolved, suggested, missing, ambig, dangling = [], [], [], [], []
                bucket = {'resolved': resolved, 'suggested': suggested, 'missing': missing}
                for num in nums:
                    sy = bib.get(num)
                    if not sy:
                        if num in bibtext:
                            missing.append((num, None))
                        elif num <= maxnum:
                            dangling.append(num)
                        continue
                    hint = bibfull.get(num) or bibtext.get(num, '')
                    st, item, kind = classify(num, sy[0], sy[1], hint)
                    bucket[st].append(item)
                    if kind == 'ambiguous':
                        ambig.append(num)
                if not (resolved or suggested or missing or dangling):
                    i = j; continue
                report.append({'kind': 'numlist',
                               'orig': ''.join(runs[k].text or '' for k in members).strip(),
                               'resolved': resolved, 'suggested': suggested,
                               'missing': missing, 'ambiguous': ambig, 'dangling': dangling})
                applied = resolved + suggested
                cite = ('{' + '; '.join(_token(rc) for _, rc in applied) + '}') if applied else ''
                ref_run = runs[members[0]]
                flagged = [k for k, _ in missing] + dangling
                if flagged:
                    _insert_like(para, ref_run._r, ref_run,
                                 (' ' if cite else '') + ', '.join(str(k) for k in sorted(set(flagged))), RED)
                ref_run.text = cite
                _clear_rstyle(ref_run)
                ref_run.font.highlight_color = None
                for k in members[1:]:
                    runs[k].text = ''
                i = j

    _process_numlists(_target_paras)

    # Serialise the footnote/endnote trees that the passes just mutated in place.
    extra_parts = {}
    for _fname, _tree in _note_trees.items():
        _xml = etree.tostring(_tree, encoding='unicode')
        extra_parts[_fname] = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n' + _xml
        ).encode('utf-8')

    buf = io.BytesIO(); doc.save(buf)
    return _fix_zoom(buf.getvalue(), extra_parts), report


def _fix_zoom(b, extra_parts=None):
    extra_parts = extra_parts or {}
    zin = zipfile.ZipFile(io.BytesIO(b)); out = io.BytesIO()
    with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
        for it in zin.infolist():
            if it.filename in extra_parts:
                d = extra_parts[it.filename]
            else:
                d = zin.read(it.filename)
                if it.filename == 'word/settings.xml':
                    d = re.sub(rb'<w:zoom[^>]*/>', b'<w:zoom w:percent="100"/>', d)
            zout.writestr(it, d)
    return out.getvalue()


# ── summary + report doc ─────────────────────────────────────────────────── #
def summarize(report):
    applied = sum(len(r['resolved']) + len(r['suggested']) for r in report)
    verify = sum(len(r['suggested']) for r in report)
    return {
        'placeholders': len(report),
        'resolved_refs': applied,            # everything with a record is applied now
        'applied_refs': applied,
        'verify_refs': verify,               # applied but flagged to double-check
        'suggested_refs': verify,
        'unresolved_refs': sum(len(r['missing']) for r in report),
        'dangling_refs': sum(len(r.get('dangling', [])) for r in report),
        'fully_done': sum(1 for r in report if (r['resolved'] or r['suggested']) and not r['missing'] and not r.get('dangling')),
    }


def build_report_docx(report, title="Placeholder \u2192 EndNote Conversion Report"):
    ACCENT, FILL = "8B1A1A", "F3ECEC"
    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), fill); tcPr.append(sh)
    def grid(table):
        b = OxmlElement('w:tblBorders')
        for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            x = OxmlElement('w:' + e); x.set(qn('w:val'), 'single'); x.set(qn('w:sz'), '4')
            x.set(qn('w:space'), '0'); x.set(qn('w:color'), 'CCCCCC'); b.append(x)
        tblPr = table._tbl.tblPr; anchor = None
        for c in tblPr:
            if c.tag in (qn('w:shd'), qn('w:tblLayout'), qn('w:tblLook')):
                anchor = c; break
        anchor.addprevious(b) if anchor is not None else tblPr.append(b)

    doc = Document()
    s = doc.sections[0]; s.left_margin = s.right_margin = Inches(0.8); s.top_margin = s.bottom_margin = Inches(0.8)
    doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(10)
    h = doc.add_paragraph(); r = h.add_run(title)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor.from_string(ACCENT)
    sm = summarize(report)
    sp = doc.add_paragraph(); _danglingtot = sm.get('dangling_refs', 0)
    _summ = ("%d placeholder(s):  %d reference(s) applied  (%d of them flagged to verify),  %d unresolved (red)."
             % (sm['placeholders'], sm['applied_refs'], sm['verify_refs'], sm['unresolved_refs']))
    if _danglingtot:
        _summ += "  %d dangling citation(s) - cited number with no reference in the list." % _danglingtot
    sr = sp.add_run(_summ)
    sr.font.size = Pt(9); sr.bold = True
    note = doc.add_paragraph(); nr = note.add_run(
        "Everything with a library match was applied as {Author, Year #RecNum}. VERIFY rows were "
        "applied too but are the matcher's best guess - check the ones marked \u201csame author/year\u201d "
        "first, those are the likeliest to need a swap. Unresolved (red) = in your list but not in the "
        "library, add it. DANGLING (red) = the in-text number has no matching entry in the reference "
        "list at all - the reference was likely deleted or the numbering is off; fix the list or remove "
        "the citation.")
    nr.font.size = Pt(8.5); nr.italic = True; nr.font.color.rgb = RGBColor.from_string("666666")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2); grid(table)
    hdr = table.rows[0]
    for i, t in enumerate(["Placeholder", "Result"]):
        c = hdr.cells[i]; run = c.paragraphs[0].add_run(t)
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor.from_string("FFFFFF"); shade(c, ACCENT)
    trPr = hdr._tr.get_or_add_trPr(); th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)

    for row in report:
        tr = table.add_row()
        p = tr.cells[0].paragraphs[0]
        run = p.add_run(row['orig'][:90]); run.bold = True; run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(ACCENT); shade(tr.cells[0], FILL)
        cell = tr.cells[1]; cell.paragraphs[0].text = ""; firstpar = [True]
        def line(text, color="000000", bold=False, fill=None):
            para = cell.paragraphs[0] if firstpar[0] else cell.add_paragraph(); firstpar[0] = False
            rr = para.add_run(text); rr.font.size = Pt(9.5); rr.bold = bold
            rr.font.color.rgb = RGBColor.from_string(color)
            if fill:
                _shade_run(rr, fill)
        for key, rec in row['resolved']:
            line("APPLIED   %s  \u2192  %s, %s #%d  (%s)" % (key, rec['sur'], rec['year'], rec['id'], rec['jour'][:24]), "1A6B2A")
        amb = set(row.get('ambiguous', []))
        for key, rec in row['suggested']:
            tag = "VERIFY (same author/year)" if key in amb else "VERIFY (near match)"
            line("%s  %s  \u2192  %s, %s #%d  (%s)  \u2013 applied, double-check"
                 % (tag, key, rec['sur'], rec['year'], rec['id'], rec['jour'][:24]), "7A5A00", fill=ORANGE)
        for key, rec in row['missing']:
            line("UNRESOLVED  %s  \u2013 add to library" % key, "8B1A1A", bold=True, fill=RED)
        for num in row.get('dangling', []):
            line("DANGLING  %s  \u2013 no reference #%s in the bibliography (citation points to a missing entry)"
                 % (num, num), "8B1A1A", bold=True, fill=RED)
        if not (row['resolved'] or row['suggested'] or row['missing'] or row.get('dangling')):
            line("(no references parsed)", "666666")

    table.autofit = False; table.allow_autofit = False
    for rw in table.rows:
        rw.cells[0].width = Inches(2.0); rw.cells[1].width = Inches(5.0)
    buf = io.BytesIO(); doc.save(buf)
    return _fix_zoom(buf.getvalue())
